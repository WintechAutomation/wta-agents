# -*- coding: utf-8 -*-
"""연구노트 docx에 참고문서 HTML의 base64 이미지 삽입"""
import sys, io, re, base64, os, tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

HTML_DIR = r'C:\MES\wta-agents\reports\MAX\경상연구개발\참고문서-원본'
DOCX_DIR = r'C:\MES\wta-agents\reports\MAX\경상연구개발\v2'

# 매핑: (HTML 파일, docx 파일)
PAIRS = [
    ('연구개발-1-장비물류.html', '연구노트-1-장비물류_new.docx'),
    ('연구개발-2-분말검사.html', '연구노트-2-분말검사.docx'),
    ('연구개발-3-연삭측정제어.html', '연구노트-3-연삭측정제어_new.docx'),
    ('연구개발-4-포장혼입검사.html', '연구노트-4-포장혼입검사_new.docx'),
    ('연구개발-5-호닝신뢰성.html', '연구노트-5-호닝신뢰성_new.docx'),
]

IMG_PATTERN = re.compile(
    r'<img[^>]*src="data:image/(png|jpeg|gif|webp);base64,([^"]+)"[^>]*(?:alt="([^"]*)")?[^>]*/?\s*>',
    re.IGNORECASE
)

# alt 없는 경우 다시 추출
ALT_PATTERN = re.compile(r'alt="([^"]*)"', re.IGNORECASE)


def extract_images(html_path):
    """HTML에서 base64 이미지 추출 → [(img_bytes, ext, alt_text), ...]"""
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    images = []
    for m in IMG_PATTERN.finditer(html):
        ext = m.group(1)
        b64 = m.group(2)
        alt = m.group(3) or ''

        # alt가 비었으면 전체 태그에서 다시 추출
        if not alt:
            full_tag = m.group(0)
            alt_m = ALT_PATTERN.search(full_tag)
            if alt_m:
                alt = alt_m.group(1)

        try:
            img_bytes = base64.b64decode(b64)
            if len(img_bytes) < 500:  # 너무 작은 이미지 무시 (아이콘 등)
                continue
            images.append((img_bytes, ext, alt))
        except Exception:
            continue

    return images


def find_results_cell(doc):
    """'3. 결과 및 데이터' 테이블의 내용 셀 찾기"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '결과 및 데이터' in cell.text or '결과 및' in cell.text:
                    # 다음 행의 셀이 실제 내용 셀
                    table_rows = table.rows
                    for ri, r in enumerate(table_rows):
                        for c in r.cells:
                            if '결과 및 데이터' in c.text:
                                if ri + 1 < len(table_rows):
                                    return table_rows[ri + 1].cells[0]
    return None


def find_experiment_cell(doc):
    """'2. 실험/개발 내용' 테이블의 내용 셀 찾기"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '실험' in cell.text and '개발' in cell.text:
                    table_rows = table.rows
                    for ri, r in enumerate(table_rows):
                        for c in r.cells:
                            if '실험' in c.text and '개발' in c.text:
                                if ri + 1 < len(table_rows):
                                    return table_rows[ri + 1].cells[0]
    return None


def insert_images_to_cell(cell, images, temp_dir):
    """셀에 이미지 삽입"""
    # 구분선 추가
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('\n━━━ 참고 이미지 ━━━')
    run.font.size = Pt(9)
    run.font.bold = True

    inserted = 0
    for i, (img_bytes, ext, alt) in enumerate(images):
        # 임시 파일로 저장
        img_ext = 'jpg' if ext == 'jpeg' else ext
        img_path = os.path.join(temp_dir, f'img_{i}.{img_ext}')
        with open(img_path, 'wb') as f:
            f.write(img_bytes)

        # 이미지 크기에 따라 폭 조정
        img_size_kb = len(img_bytes) / 1024
        if img_size_kb > 100:
            width = Inches(5.5)
        elif img_size_kb > 30:
            width = Inches(4.5)
        else:
            width = Inches(3.5)

        # 캡션 (alt 텍스트가 있으면 사용)
        caption = alt if alt and alt != 'image' else f'그림 {i+1}'
        # 파일명 형태의 alt 정리
        caption = caption.replace('.png', '').replace('.jpg', '').replace('.jpeg', '')
        if caption.startswith('image-'):
            caption = f'그림 {i+1}'

        # 이미지 삽입
        try:
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=width)

            # 캡션
            p_cap = cell.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(f'[{caption}]')
            run_cap.font.size = Pt(8)
            run_cap.font.italic = True

            inserted += 1
        except Exception as e:
            print(f'  이미지 {i} 삽입 실패: {e}')

    return inserted


def process_pair(html_file, docx_file):
    """HTML에서 이미지 추출 → docx에 삽입"""
    html_path = os.path.join(HTML_DIR, html_file)
    docx_path = os.path.join(DOCX_DIR, docx_file)

    if not os.path.exists(html_path):
        print(f'HTML 없음: {html_file}')
        return
    if not os.path.exists(docx_path):
        print(f'DOCX 없음: {docx_file}')
        return

    print(f'\n=== {docx_file} ===')

    # 이미지 추출
    images = extract_images(html_path)
    print(f'  HTML 이미지: {len(images)}개')

    if not images:
        print('  이미지 없음, 건너뜀')
        return

    # docx 로드
    doc = Document(docx_path)

    # 결과 및 데이터 셀 찾기
    results_cell = find_results_cell(doc)
    if not results_cell:
        print('  "결과 및 데이터" 셀 못 찾음!')
        return

    print(f'  결과 셀 텍스트: {results_cell.text[:60]}...')

    # 임시 디렉토리
    with tempfile.TemporaryDirectory() as temp_dir:
        inserted = insert_images_to_cell(results_cell, images, temp_dir)

    # 저장 (_img.docx)
    out_name = docx_file.replace('.docx', '_img.docx')
    out_path = os.path.join(DOCX_DIR, out_name)
    doc.save(out_path)
    print(f'  저장: {out_path} ({inserted}개 이미지 삽입)')


# 실행
for html_file, docx_file in PAIRS:
    process_pair(html_file, docx_file)

print('\n=== 완료 ===')
