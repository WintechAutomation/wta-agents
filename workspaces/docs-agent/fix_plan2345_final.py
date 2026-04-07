"""연구개발계획서 #2~#5 최종 보완 스크립트
참고문서(Confluence/Jira) 데이터 기반 구체화
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
from pathlib import Path
from docx import Document
from lxml import etree

BASE = Path(r"C:\MES\wta-agents\reports\MAX\경상연구개발")


def set_cell_text(cell, text):
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
            return
    if cell.paragraphs:
        cell.paragraphs[0].add_run(text)


def replace_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def add_paragraph_after(ref_para, text):
    from docx.oxml.ns import qn
    new_p = copy.deepcopy(ref_para._element)
    for r in new_p.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            t.text = ""
    runs = new_p.findall(qn('w:r'))
    if runs:
        ts = runs[0].findall(qn('w:t'))
        if ts:
            ts[0].text = text
            ts[0].set(qn('xml:space'), 'preserve')
    ref_para._element.addnext(new_p)


def replace_textbox(body, old, new):
    count = 0
    for elem in body.iter():
        tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ''
        if tag == 'txbxContent':
            for t in elem.iter():
                if isinstance(t.tag, str) and etree.QName(t.tag).localname == 't':
                    if t.text and old in t.text:
                        t.text = t.text.replace(old, new)
                        count += 1
    return count


def fix_header(doc, title, date="2025.03.01"):
    for s in doc.sections:
        for t in s.header.tables:
            for r in t.rows:
                for c in r.cells:
                    txt = c.text.strip()
                    if "개발과제명" in txt:
                        for p in c.paragraphs:
                            if p.runs:
                                p.runs[0].text = title
                                for run in p.runs[1:]:
                                    run.text = ""
                    if "2020" in txt:
                        for p in c.paragraphs:
                            if p.runs:
                                p.runs[0].text = date
                                for run in p.runs[1:]:
                                    run.text = ""
                    if "사원 홍길동" in txt or "생산관리팀" in txt:
                        pass  # 이미 설정됨


def fix_cover_eng(doc, eng_title):
    for p in doc.paragraphs:
        if '『' in p.text:
            if p.runs:
                p.runs[0].text = f'『 {eng_title} 』'
                for r in p.runs[1:]:
                    r.text = ""
            break


def find_and_replace(doc, keyword, new_text):
    for p in doc.paragraphs:
        if keyword in p.text:
            replace_para(p, new_text)
            return True
    return False


# ═══════════════════════════════════════════════════════
# 과제 #2 — 분말검사
# ═══════════════════════════════════════════════════════
def fix_project2():
    path = BASE / "연구개발계획서-2-분말검사.docx"
    doc = Document(str(path))

    title_kr = "프레스성형 품질향상을 위한 분말성형체 검사기술 개발"
    title_en = "Development of Powder Compact Inspection Technology for Press Forming Quality Improvement"

    # 표지
    replace_textbox(doc.element.body, "개발과제명", title_kr)
    fix_cover_eng(doc, title_en)
    fix_header(doc, title_kr)

    # 개발 목적
    purposes = [
        "프레스 성형 직후 분말성형체(그린바디)의 외관 결함(Burr, 깨짐, ��랙, 치수 불량)을 인라인으로 자동 검사하는 기술을 개발한다. OTC 광학기술센터에서 개발 완료한 프레스-IM(Integrated Metrology) 광학계를 기반으로, 상면·측면 동시 촬영이 가능한 다채널 돔조명(HM동축+직접돔+DF) 시스템을 적용한다.",
        "성형 직후 단계에서 불량을 조기 검출하여 후공정 투입 전 불량 선별을 실현하고, 원자재 손실을 절감한다. Burr 검출기준 0.1mm 이상(인선부 상/하 및 모든 코너부)을 적용하며, Burr 이상 검출 시 금형 문제를 자동 확인하고 금형 Cleaning 상태를 판단하는 연동 기능을 구현한다.",
        "AI 학습 기반 검사를 고도화하여 양불 판정 정확도 99% 이상을 달성하고, 금형 상태 연관 분석으로 금형 교체 시점을 예측한다. MES 연동으로 금형별/제품별 불량률 트렌드를 실시간 모니터링한다.",
    ]
    idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if idx == 0 and "외관 결함" in t:
            replace_para(p, purposes[0]); idx = 1
        elif idx == 1 and "조기 검출" in t:
            replace_para(p, purposes[1]); idx = 2
        elif idx == 2 and "AI 학습" in t:
            replace_para(p, purposes[2]); idx = 0

    # 필요성
    necs = [
        "현재 성형 후 검사는 소결 공정 이후에 수행되어, 불량 성형체가 소결·연삭·코팅 전 공정을 거친 후에야 발견되는 비효율이 존재한다. 소결 전 불량 선별 시 후공정 손실 원가 60% 이상 절감이 가능하다. 분말 야금 공정 특성상 성형 압력·중량 관리가 품질의 핵심이며, 성형체 단계에서의 조기 검출이 전체 공정 수율 향상의 관건이다.",
        "2024년 Burr 검사 장치(검출률 98%, 사이클 1.5초)는 소결 후 제품 대상이며, 분말성형체(그린바디)는 강도가 매우 낮고 표면이 무광·분말 부착 특성을 가진다. OTC에서 프레스-IM 광학계 셋업을 완료(2025.7)하였으나, Side 광학계 burr 실루엣 미시인 이슈(Both부 가림)가 확인되어 top 조명 추가 적용 등 광량 강화가 필요한 상황이다.",
    ]
    n_idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if n_idx == 0 and "성형 후 검사는 소결" in t:
            replace_para(p, necs[0]); n_idx = 1
        elif n_idx == 1 and ("Burr 검사" in t or "검사 장치" in t):
            replace_para(p, necs[1]); n_idx = 0

    # 추진전략
    strats = [
        "1단계(3~5월): 그린바디 전용 비전 검사 시스템 설계. OTC 프레스-IM 광학계(HM동축+직접돔+DF 다채널 조명) 셋업 완료 성과를 기반으로, 분말성형체 표면 특성(무광, 분말 부착, 낮은 강도)에 최적화된 조명 및 멀티 카메라 구성을 개발한다. Side 광학계 burr 실루엣 미시인 이슈를 해결하기 위한 광량 강화 방안을 검증한다.",
        "2단계(6~9월): 불량 검출 알고리즘 개발. 2024년 Burr 검사 알고리즘(검출기준 0.1mm 이상, 인선부 상/하/코너부 전체) 기반에 그린바디 전용 파라미터를 튜닝한다. Deep-Learning 기반 외관 검출력을 적용하고, AI 학습 기반 양불 판정 정확도 99% 이상을 달성한다.",
        "3단계(10~12월): ��레스 핸들러 인라인 연동 및 MES 데이터 연계. takt time 1.5초 이내 검사 완료(프레스 SPM 23EA/min 대응), 금형별/제품별 불량률 트렌드 분석 기능을 구현한다. 금형 Cleaning 상태 자동 판단 및 금형 교체 시점 예측 기능을 포함���다.",
    ]
    s_idx = 0
    for p in doc.paragraphs:
        if s_idx == 0 and "1단계" in p.text:
            replace_para(p, strats[0]); s_idx = 1
        elif s_idx == 1 and "2단계" in p.text:
            replace_para(p, strats[1]); s_idx = 2
        elif s_idx == 2 and "3단계" in p.text:
            replace_para(p, strats[2]); s_idx = 0

    # 제약사항
    find_and_replace(doc, "[제약사항]",
        "[제약사항] (1) 그린바디의 낮은 강도(소결 전)로 비접촉 검사 방식 필수 — 접촉 시 파손 위험. "
        "(2) 프레스 takt time(SPM 23EA/min) 내 검사 완료 필요 — 1.5초/개 목표. "
        "(3) Side 광학계 burr 실루엣 미시인 이슈(Both부 가림) — 광량 강화 필요. "
        "(4) 분말 부착·무광 표면에서의 비전 검출 안정성 확보 난이도. "
        "(5) 소형 제품(⌀1.5~) 대응 시 분해능 한계."
    )

    # 개발 방법
    find_and_replace(doc, "개발 방법 :",
        "개발 방법 : 프레스-IM 다채널 돔조명 기반 비전 검사 시스템 설계, "
        "Deep-Learning 기반 Burr/크랙/���짐 복합 검출 알고리즘 개발, "
        "프레스 핸들러 인라인 연동(SPM 23EA/min 대응), MES 금형 품질 데이터 연계"
    )

    # WBS
    wbs = [
        ("프레스-IM 광학계 기반 비전시스템", "3개월", "2025-03-01", "2025-05-31", "서제완(OTC) + 기구 1"),
        ("Burr/크랙 검출 알고리즘 개발", "4개월", "2025-06-01", "2025-09-30", "진소미 + 비전 S/W 1"),
        ("핸들러 인라인 연동", "3개월", "2025-08-01", "2025-10-31", "조윤명(S/W) + 기구 1"),
        ("MES 연계 및 현장 실증", "2개월", "2025-11-01", "2025-12-31", "전원 (5명)"),
    ]
    for ri, (name, period, start, end, resource) in enumerate(wbs):
        row = doc.tables[1].rows[ri + 1]
        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], period)
        set_cell_text(row.cells[2], start)
        set_cell_text(row.cells[3], end)
        set_cell_text(row.cells[4], resource)

    # 인력
    t2 = doc.tables[2]
    set_cell_text(t2.rows[1].cells[1], "서제완")
    set_cell_text(t2.rows[1].cells[3], "과제 총괄, 프레스-IM 광학계")
    set_cell_text(t2.rows[2].cells[1], "진소미")
    set_cell_text(t2.rows[2].cells[3], "검출 알고리즘/AI 개발")
    set_cell_text(t2.rows[3].cells[1], "( 5 ) 명")
    set_cell_text(t2.rows[3].cells[2], "( 5 ) 명")
    set_cell_text(t2.rows[3].cells[3], "광학1 + 기구2 + 비전SW2 + SW1")

    # 관리자
    t0 = doc.tables[0]
    set_cell_text(t0.rows[1].cells[0], "서제완")
    set_cell_text(t0.rows[1].cells[1], "OTC 광학기술센터")
    set_cell_text(t0.rows[1].cells[2], "과제 총괄")
    set_cell_text(t0.rows[2].cells[0], "진소미")
    set_cell_text(t0.rows[2].cells[1], "비전팀")
    set_cell_text(t0.rows[2].cells[2], "알고리즘 검토")

    doc.save(str(path))
    print(f"  #2 분말검사 완료: {path.name}")


# ═══════════════════════════════════════════════════════
# 과제 #3 — 연삭측정제어
# ═══════════════════════════════════════════════════════
def fix_project3():
    path = BASE / "연구개발계획서-3-연삭측정제어.docx"
    doc = Document(str(path))

    title_kr = "연삭체의 정밀 연삭 가공을 위한 측정 제어장치 및 그 제어방법"
    title_en = "Measurement Control Device and Method for Precision Grinding of Ground Bodies"

    replace_textbox(doc.element.body, "개발과제명", title_kr)
    fix_cover_eng(doc, title_en)
    fix_header(doc, title_kr, "2025.04.01")

    # 목적
    purposes = [
        "초경인서트 양면 연삭 공정에서 연삭 전·중·후 치수를 자동 측정하고, 측정 데이터를 기반으로 연삭량을 실시간 피드백 제어���는 측정 제어장치를 개발한다. 현재 외부 리니어스케일(분해능 3.75μm)과 접촉 프로브(분해능 2μm)를 적용하고 있으나, 위치 피드백과 실제 휠 위치 간 기구적 결합 오차(이격)가 발생하여 치수 편차의 원인이 되고 있다.",
        "연삭 과정에서 목표 치수에 도달하면 자동 정지하는 폐루프(Closed-loop) 연삭 제어 시스템을 구현한다. 외부스케일 기반 별도 위치 제어 알고리즘(사다리꼴 속도 프로파일, 10ms 제어주기, EtherCAT PC통신 기반 C# 제어)을 개발하여 0~3μm 오차 수준의 정밀 위치 제어를 달성한다.",
        "제품별 연삭 이력(측정값, 연삭량, 연삭 조건)을 자동 기록하고 SPC 기반 공정 능력 분석 체계를 구축한다. 연삭 결과물 치수 편차를 최소화하기 위해 연삭 조건(연삭하중, 정지시점, 회전 가감속)별 영향도를 분석하고, 정지시점(목표위치) 기반 치수 제어의 신뢰성을 확보한다.",
    ]
    idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if idx == 0 and "연삭 전·중·후 치수를 자동 측정" in t:
            replace_para(p, purposes[0]); idx = 1
        elif idx == 1 and "폐루프" in t:
            replace_para(p, purposes[1]); idx = 2
        elif idx == 2 and "연삭 이력" in t:
            replace_para(p, purposes[2]); idx = 0

    # 필요성
    necs = [
        "현재 연삭 공정은 작업자 숙련도에 따라 연삭량이 결정되며, 과삭/미삭 불량이 빈발한다. 외부스케일 피드백과 실제 휠 위치 간 이격이 발생하고(모터 엔코더 vs 외부스케일 차이), 휠 회전 감속 시간 동안 추가 연삭이 진행되어 치수 편차가 커진다. 2025년 4월 연삭팀 DX 회의에서 '초보자도 정밀 연삭 가능한 자동 시스템'과 '무인화·연속 자동 가동' 시스템 개발이 핵심 목표로 설정되었다.",
        "2023년 과제는 연삭 전/후 측정→연삭량 계산→연삭기 전송의 오프라인 방식이었다. 본 과제는 외부스케일 기준 실시간 위치 제어(사다리꼴 프로파일, 0~3μm 오차)와 폐루프 피드백 보정을 추가하여, 목표 치수 대비 편차 ±1μm 이내를 달성하는 고정밀 제어 시스템으로 고도화한다. AGV 기반 물류 자동화, 비전 기반 각도/위치 제어, 초음파 세척+에어 건조+측정 통합 공정도 포함한다.",
    ]
    n_idx = 0
    for p in doc.paragraphs:
        if n_idx == 0 and "작업자 숙련도" in p.text:
            replace_para(p, necs[0]); n_idx = 1
        elif n_idx == 1 and ("2023년 과제" in p.text or "오프라인 방식" in p.text):
            replace_para(p, necs[1]); n_idx = 0

    # 전략
    strats = [
        "1단계(4~6월): 인프로세스 측정 기술 개발. 외부 리니어스케일(3.75μm) 및 접촉 프로브(2μm) 기반 연삭 중 두께 실시간 측정. 외부스케일 기준 별도 위치 제어 알고리즘(사다리꼴 속도 프로파일, EtherCAT C# 제어, 10ms 주기)을 구현하여 0~3μm 오차를 달성한다. 연삭 진동·칩·절삭유 환경 대응 안정성을 확보한다.",
        "2단계(7~9월): 폐루프 연삭 제어 알고리즘 개발. ��삭 결과물 치수를 피드백으로 정지시점(목표위치) 보정. 연삭 조건별(연삭하중, 회전 ���감속, 정지시점) 치수 영향도 분석. 접근/정삭/스파크아웃 자동 전환 및 연삭 압력 자동 미세 조절, 열변형 보정 알고리즘을 포함한다.",
        "3단계(10~12월): 측정 제어장치 하드웨어 통합 및 연삭기 CNC 인터페이스(EtherCAT/PROFINET) 연동. 비전 기반 제품 위치/각도 제어, 초음파 세척+에어 건조+측정 통합 공정 적용. MES 품질 데이터 실시간 전송, 자동 드레싱·지그 공급 스테이션 연동, 현장 실증을 수행한다.",
    ]
    s_idx = 0
    for p in doc.paragraphs:
        if s_idx == 0 and "1단계" in p.text:
            replace_para(p, strats[0]); s_idx = 1
        elif s_idx == 1 and "2단계" in p.text:
            replace_para(p, strats[1]); s_idx = 2
        elif s_idx == 2 and "3단계" in p.text:
            replace_para(p, strats[2]); s_idx = 0

    find_and_replace(doc, "[제약사항]",
        "[제약사항] (1) 외부스케일 피드백과 실제 휠 위치 간 기구적 이격 — 위치 재현성 한계. "
        "(2) 휠 회전 감속 시간 중 추가 연삭 — 정지시점 보정 필요. "
        "(3) 연삭유·분진 환경에서의 센서 내구성·비전 검출 안정성. "
        "(4) 다품종 소량 대응 — 제품별 연삭 조건 자동 전환(RFID 기반). "
        "(5) EtherCAT PC통신 기반 제어의 서보 대비 응답성 열위."
    )

    find_and_replace(doc, "개발 방법 :",
        "개발 방법 : 외부스케일 기반 실시간 위치 제어 알고리즘 개발(EtherCAT C#), "
        "폐루프 피드백 연삭 ���어, 비전 기반 제품 각도/위치 제어, "
        "초음파 세척+건조+측정 통합 공정, MES 연동 SPC 품질 관리"
    )

    # WBS
    wbs = [
        ("외부스케일 위치제어 알고리즘", "3개월", "2025-04-01", "2025-06-30", "김웅기(제어) + 정광선(S/W)"),
        ("폐루프 연삭 제어 알고��즘", "3개월", "2025-07-01", "2025-09-30", "김웅기 + 최현수(S/W)"),
        ("H/W 통합 및 CNC 연동", "3개월", "2025-10-01", "2025-12-31", "기구설계 1 + 제어 1"),
        ("MES 연계 및 현장 실증", "2개월", "2025-11-01", "2025-12-31", "전원 (4명)"),
    ]
    for ri, (name, period, start, end, resource) in enumerate(wbs):
        row = doc.tables[1].rows[ri + 1]
        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], period)
        set_cell_text(row.cells[2], start)
        set_cell_text(row.cells[3], end)
        set_cell_text(row.cells[4], resource)

    t2 = doc.tables[2]
    set_cell_text(t2.rows[1].cells[1], "김웅기"); set_cell_text(t2.rows[1].cells[2], "9")
    set_cell_text(t2.rows[1].cells[3], "과제 총괄, 치수제어 알고리즘")
    set_cell_text(t2.rows[2].cells[1], "정광선"); set_cell_text(t2.rows[2].cells[2], "9")
    set_cell_text(t2.rows[2].cells[3], "핸들러 S/W, MES 연동")
    set_cell_text(t2.rows[3].cells[1], "( 4 ) 명")
    set_cell_text(t2.rows[3].cells[2], "( 4 ) 명")
    set_cell_text(t2.rows[3].cells[3], "기구1 + 제어1 + S/W2")

    t0 = doc.tables[0]
    set_cell_text(t0.rows[1].cells[0], "김웅기"); set_cell_text(t0.rows[1].cells[1], "제어설계팀")
    set_cell_text(t0.rows[1].cells[2], "과제 총괄")
    set_cell_text(t0.rows[2].cells[0], "김순봉"); set_cell_text(t0.rows[2].cells[1], "연삭팀")
    set_cell_text(t0.rows[2].cells[2], "연삭 공정 검토")

    doc.save(str(path))
    print(f"  #3 연삭측정제어 완료: {path.name}")


# ═══════════════════════════════════════════════════════
# 과제 #4 — 포장혼입검사
# ═══════════════════════════════════════════════════════
def fix_project4():
    path = BASE / "연구개발계획서-4-포장혼입검사.docx"
    doc = Document(str(path))

    title_kr = "인서트 포장기 혼입검사기술 개발"
    title_en = "Development of Cross-Contamination Inspection Technology for Insert Packaging Machine"

    replace_textbox(doc.element.body, "개발과제명", title_kr)
    fix_cover_eng(doc, title_en)
    fix_header(doc, title_kr, "2025.04.01")

    purposes = [
        "초경인서트 최종 포장 공정에서 이종 제품 혼입(Cross-contamination)을 자동 검출하는 인라인 검사 기술을 개발한다. Korloy 포장기 #6에서 C/W형 경면 제품의 Nose-R부 치수가 비대칭 밝기로 인해 상이하게 측정되는 광학계 이슈가 확인되었으며, 이를 근본적으로 해결하는 고신뢰 혼입검사 시스템을 구현한다.",
        "인서트의 형상·각인·색상·치수를 고속(0.8초/개 이내)으로 검사하여, 혼입 제품을 실시간 선별하는 시스템을 구현한다. 딥러닝 OCR을 적용하여 제품 각인(모델명, 등급 코드)을 자동 판독하고, Chip Breaker 혼입 및 Nose-R 혼입(0.2mm 이상)을 정밀 검출한다.",
        "포장 단위별 검사 이력을 기록하여 완전한 트레이서빌리티를 확보하고, 고객 클레임 원인인 혼입 불량을 근절한다. 케이스 내 제품 수량 검사, 마킹 유무/품질 검사 기능도 통합한다.",
    ]
    idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if idx == 0 and "이종 제품 혼입" in t:
            replace_para(p, purposes[0]); idx = 1
        elif idx == 1 and "형상·각인·색상" in t:
            replace_para(p, purposes[1]); idx = 2
        elif idx == 2 and "트레이서빌리티" in t:
            replace_para(p, purposes[2]); idx = 0

    necs = [
        "현재 포장 공정에서 작업자 실수 또는 팔레트 혼재로 인해 타 규격 제품이 혼입되는 품질 사고가 간헐적으로 발생하고 있다. Korloy 포장기 #6에서 C/W형 경면 제품의 밝기 비대칭 현상(광축 중심과 조명 중심의 ±1mm 이상 편차)이 확인되었으며, 이로 인해 Nose-R 치수 측정 오차가 발생하고 있다. 글로벌 고객사 품질 요구(ISO/IATF)를 충족하려면 100% 자동 검사가 필수적이다.",
        "기존 검사 과제들은 단일 불량 유형(Burr, 깨짐, 호닝 치수)을 대상으로 했으나, 혼입 검사는 제품 식별(Identification) 문제이다. 중국교세라 포장기에서 딥러닝 OCR 코드를 적용 테스트 중이며(S/W팀 조윤명), 혼입 검사 시퀀스 개선 및 오검출 저감 작업이 진행 중이다. 형상 유사 인서트 간 미세 차이(각인 문자, 코팅 색상, Nose-R 0.1~0.2mm 차이)를 고속으로 판별하는 기술이 필요하다.",
    ]
    n_idx = 0
    for p in doc.paragraphs:
        if n_idx == 0 and "작업자 실수" in p.text:
            replace_para(p, necs[0]); n_idx = 1
        elif n_idx == 1 and ("단일 불량" in p.text or "기존 검사" in p.text):
            replace_para(p, necs[1]); n_idx = 0

    strats = [
        "1단계(4~6월): 다중 특징 기반 제품 식별 시스템 설계. 형상 매칭(CB 방향 정렬, 앞/뒷면 구분), 딥러닝 OCR 각인 인식, 코팅 색상 분석, 핵심 치수(Nose-R, IC) 검증의 4대 식별 모듈을 설계한다. Korloy #6 혼입 광학계 이슈(광축·조명 중심 편차 ±1mm 이내 세팅 필수, 조명 높이 10mm 권장)를 설계에 반영한다.",
        "2단계(7~9월): 고속 검사 광학계 개발. 혼입검사부 광학계(acA2500-14gm 카메라, M2514-MP2 렌즈, DOMELIGHT100 조명, LWD 169mm 기준 FOV 38.25x28.86mm)를 최적화한다. 텔레센트릭 렌즈 적용, 0.8초/개 이내 상면·측면 동시 촬영. 딥러닝 OCR 테스트 결과(중국교세라 현장 적용 중)를 반영한 혼입 판정 알고리즘을 개발한다.",
        "3단계(10~12월): 포장기 인라인 연동 및 현장 실증. 기존 공정 흐름 변경 최소화하여 포장기 투입부에 설치, 혼입 자동 배출(리젝트) 기능을 구현한다. 케이스 내 수량 검사(IR 조명 적용), 마킹 검사(잉크젯/레이저), 검사 이력 DB 연동을 포함한다.",
    ]
    s_idx = 0
    for p in doc.paragraphs:
        if s_idx == 0 and "1단계" in p.text:
            replace_para(p, strats[0]); s_idx = 1
        elif s_idx == 1 and "2단계" in p.text:
            replace_para(p, strats[1]); s_idx = 2
        elif s_idx == 2 and "3단계" in p.text:
            replace_para(p, strats[2]); s_idx = 0

    find_and_replace(doc, "[제약사항]",
        "[제약사항] (1) 경면 제품(C/W형)의 밝기 비대칭 — 광축·조명 중심 정밀 세팅 필수(±1mm). "
        "(2) Nose-R 0.1mm(01) 미세 차이 판별 한계 — 0.2mm(02) 이상부터 검출 목표. "
        "(3) 포장기 takt time(SPM 60EA/min) 내 검사 — 0.8초/개 이내. "
        "(4) 다양한 코팅(CVD/PVD/무코팅) 및 케이스 형태 대응. "
        "(5) 현장 딥러닝 OCR 안정화(제품별 에러율 저감) 필요."
    )

    find_and_replace(doc, "개발 방법 :",
        "개발 방법 : 다채널 돔조명 기반 고속 광학계 설계, "
        "딥러닝 OCR + CB/Nose-R 혼입 판정 알고리즘 개발, "
        "포장기 인라인 연동(SPM 60EA/min 대응), 수량/마킹 통합 검사"
    )

    wbs = [
        ("제품 식별 시스템 및 광학계 설계", "3개월", "2025-04-01", "2025-06-30", "윤선웅(OTC) + 비전 1"),
        ("혼입 판정 알고리즘(OCR/AI)", "3개월", "2025-07-01", "2025-09-30", "조윤명(S/W) + 정진원(비전)"),
        ("고속 검사 광학계 최적화", "3개월", "2025-07-01", "2025-09-30", "서제완(OTC) + 기구 1"),
        ("포장기 연동 및 현장 실증", "3개월", "2025-10-01", "2025-12-31", "전원 (4명)"),
    ]
    for ri, (name, period, start, end, resource) in enumerate(wbs):
        row = doc.tables[1].rows[ri + 1]
        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], period)
        set_cell_text(row.cells[2], start)
        set_cell_text(row.cells[3], end)
        set_cell_text(row.cells[4], resource)

    t2 = doc.tables[2]
    set_cell_text(t2.rows[1].cells[1], "조윤명"); set_cell_text(t2.rows[1].cells[2], "9")
    set_cell_text(t2.rows[1].cells[3], "과제 총괄, S/W/OCR 개발")
    set_cell_text(t2.rows[2].cells[1], "윤선웅"); set_cell_text(t2.rows[2].cells[2], "9")
    set_cell_text(t2.rows[2].cells[3], "혼입 광학계 설계")
    set_cell_text(t2.rows[3].cells[1], "( 4 ) 명")
    set_cell_text(t2.rows[3].cells[2], "( 4 ) 명")
    set_cell_text(t2.rows[3].cells[3], "기구1 + 비전SW2 + SW1")

    t0 = doc.tables[0]
    set_cell_text(t0.rows[1].cells[0], "조윤명"); set_cell_text(t0.rows[1].cells[1], "S/W팀")
    set_cell_text(t0.rows[1].cells[2], "과제 총괄")
    set_cell_text(t0.rows[2].cells[0], "윤선웅"); set_cell_text(t0.rows[2].cells[1], "OTC 광학기술센터")
    set_cell_text(t0.rows[2].cells[2], "광학계 검토")

    doc.save(str(path))
    print(f"  #4 포장혼입검사 완료: {path.name}")


# ═══════════════════════════════════════════════════════
# 과제 #5 — 호닝신뢰성
# ═══════════════════════════════════════════════════════
def fix_project5():
    path = BASE / "연구개발계획서-5-호닝신뢰성.docx"
    doc = Document(str(path))

    title_kr = "정밀 광학계 기반 호닝형상검사기의 신뢰성 확보 기술 연구"
    title_en = "Research on Reliability Assurance Technology for Precision Optics-Based Honing Shape Inspection Machine"

    replace_textbox(doc.element.body, "개발과제명", title_kr)
    fix_cover_eng(doc, title_en)
    fix_header(doc, title_kr, "2025.05.01")

    purposes = [
        "2022년 개발한 호닝형상 검사기(HIM)의 측정 신뢰성을 체계적으��� 검증·향상하여, 고객사 양산 라인에서의 장기 안정 운용을 보장한다. 72시간 롱런 테스트 결과(#23-2: 4,539회, #23-4: 4,261회 측정) 온도변화 1°C당 Z축 높이 15μm 변화가 확인되었으며, 초기 가동 후 2.5~3시간(22°C 이하) 측정값 부정확 현상, FFU 가동 시 진동에 의한 산포 증가 등의 개선 과제가 도출되었다.",
        "환경 변화(온도 16.4~27.3°C 범위), 광학 부품 경년 열화, 다양한 제품 규격 대응에서 발생하는 측정 편차 요인을 분석하고 자동 보정·자가 진단 기능을 개발한다. 광학계 Calibration 타겟(Dot 타겟 3종 + 복합 타겟 2종)을 활용한 체계적 정렬 프로세스를 확립한다.",
        "GR&R 10% 이내 달성(현 수준 대비 50% 개선), 1,000시간 연속 운전 시 정밀도 유지를 보증한다. HIM-F Series 원가절감 항목을 병행 발굴하여 제품 경쟁력을 강화한다.",
    ]
    idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if idx == 0 and "측정 신뢰성을 체계적으로" in t:
            replace_para(p, purposes[0]); idx = 1
        elif idx == 1 and "환경 변화" in t:
            replace_para(p, purposes[1]); idx = 2
        elif idx == 2 and "GR&R" in t:
            replace_para(p, purposes[2]); idx = 0

    necs = [
        "HIM은 WTA의 글로벌 1위 핵심 제품(±1μm 정밀도, 반복정밀도 ±2μm)이나, 고객사 현장 적용 시 환경 변화, 광학 열화, 다규격 대응에서 측정 편차가 보고되고 있다. 72시간 롱런 테스트에서 초기 가동 시 측정 산포가 크고(1~3시간), FFU 진동 영향이 확인되었다. 세라티즈, 한국야금, 교세라(중국), MMC, 대구텍, 다이제트 등 주요 고객사의 양산 안정성 요구가 높아지고 있다.",
        "2022년 HIM 개발은 측정 기능 구현에 초점이 맞춰져 있었으나, 장기 신뢰성 확보는 별도 과제로 다뤄지지 않았다. 현재 광학계 align 프로세스 확립, Calibration 타겟 설계, TTTM(Test Target Measurement) 장치 개조, 다채널 분할검사 적용 등이 진행 중이며, 알루미늄 인서트 클리닝 장치(Air&Brush 비접촉 방식, 레이저/플라즈마 방식 검토 중)도 검사 신뢰성에 직접 영향을 미치는 요소이다.",
    ]
    n_idx = 0
    for p in doc.paragraphs:
        if n_idx == 0 and "글로벌 1위 핵심 제품" in p.text:
            replace_para(p, necs[0]); n_idx = 1
        elif n_idx == 1 and ("2022년 HIM" in p.text or "측정 기능 구현" in p.text):
            replace_para(p, necs[1]); n_idx = 0

    strats = [
        "1단계(5~7���): 측정 불확도 체계 분석. 72시간 롱런 테스트 데이터(온도 1°C당 Z축 15μm 변화, 초기 측정 산포, FFU 진동 영향)를 기반으로 환경·광학·기구·제품 요인별 측정 편차를 정량화한다. GR&R 분석 체계를 수립하고, HIM-F 원가절감 항목(커버 설계 수정 등)을 병행 발굴한다.",
        "2단계(8~10월): 자동 보정(Auto-Calibration) 및 자가 진단 시스템 개발. Calibration 타겟(Dot 3종+복합 2종)을 활용한 주기적 자동 보정, 온도 보상 알고리즘, 조명 TTTM 제어(WhiteBalance, 입력 전원별 반사특성 보정), 광축 자동 정렬, 다채널 분할검사 기반 부품 교체 시기 예측. 알루미늄 인서트 클리닝 장치(Air&Brush 비접촉 + 레이저/플라즈마 검토)를 포함한다.",
        "3단계(11~12월): 신뢰성 검증 프로토콜 실행. 환경 변화 조건별 테스트(온도 챔버 16~27°C, 가진기), 1,000시간 연속 운전(사이클타임 57~62초 기준), 다규격(10종+) 대응, 고객사 현장 실증(한국야금, MMC 등 2개소 이상)을 수행한다. CoaxPress 카메라, Laser Triangulation 등 차세대 기술 적용 가능성도 검증한다.",
    ]
    s_idx = 0
    for p in doc.paragraphs:
        if s_idx == 0 and "1단계" in p.text:
            replace_para(p, strats[0]); s_idx = 1
        elif s_idx == 1 and "2단계" in p.text:
            replace_para(p, strats[1]); s_idx = 2
        elif s_idx == 2 and "3단계" in p.text:
            replace_para(p, strats[2]); s_idx = 0

    find_and_replace(doc, "[제약사항]",
        "[제약사항] (1) 온도변화 1°C당 Z축 15μm 변화 — 온도 보상 알고리즘 필수. "
        "(2) 초기 가동 2.5~3시간 측정 부정확 — 워밍업 시퀀스 또는 보정 루틴 필요. "
        "(3) FFU 가동 시 진동에 의한 측정 산포 증가 — 방진 대책. "
        "(4) 광택 제품(CVD 코팅 등) 형상측정 정확도 — Diffusing 강화·Polarizer 적용 필요. "
        "(5) 고객사 현장 실증을 위한 장비 반출/설치 일정 조율."
    )

    find_and_replace(doc, "개발 방법 :",
        "개발 방법 : 72시간 롱런 데이터 기반 불확도 분석, "
        "Calibration 타겟/TTTM 기반 자동 보정, "
        "다채널 분할검사·자가 진단 개발, 고객사 현장 실증(2개소+)"
    )

    wbs = [
        ("측정 불확도 분석/GR&R 체계", "3개월", "2025-05-01", "2025-07-31", "서제완(OTC) + 김준형v(비전)"),
        ("자동 보정 시스템 개발", "3개월", "2025-08-01", "2025-10-31", "정정일(기구) + 이인모(비전)"),
        ("자가 진단/클리닝 기능", "2개월", "2025-09-01", "2025-10-31", "정진원(비전) + 문태수"),
        ("신뢰성 검증 및 현장 실증", "2개월", "2025-11-01", "2025-12-31", "전원 (4명)"),
    ]
    for ri, (name, period, start, end, resource) in enumerate(wbs):
        row = doc.tables[1].rows[ri + 1]
        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], period)
        set_cell_text(row.cells[2], start)
        set_cell_text(row.cells[3], end)
        set_cell_text(row.cells[4], resource)

    t2 = doc.tables[2]
    set_cell_text(t2.rows[1].cells[1], "김준형v"); set_cell_text(t2.rows[1].cells[2], "8")
    set_cell_text(t2.rows[1].cells[3], "과제 총괄, 비전 알고리즘")
    set_cell_text(t2.rows[2].cells[1], "서제완"); set_cell_text(t2.rows[2].cells[2], "8")
    set_cell_text(t2.rows[2].cells[3], "광학계 보정/TTTM")
    set_cell_text(t2.rows[3].cells[1], "( 4 ) 명")
    set_cell_text(t2.rows[3].cells[2], "( 4 ) 명")
    set_cell_text(t2.rows[3].cells[3], "비전SW2 + 광학1 + SW1")

    t0 = doc.tables[0]
    set_cell_text(t0.rows[1].cells[0], "김준형v"); set_cell_text(t0.rows[1].cells[1], "비전팀")
    set_cell_text(t0.rows[1].cells[2], "과제 총괄")
    set_cell_text(t0.rows[2].cells[0], "정정일"); set_cell_text(t0.rows[2].cells[1], "기구설계팀 HIM")
    set_cell_text(t0.rows[2].cells[2], "기구설계 검토")

    doc.save(str(path))
    print(f"  #5 호닝신뢰성 완료: {path.name}")


# ═══════════════════════════════════════════════════════
# 메인
# ═══════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 60)
    print("연구개발 계획서 #2~#5 최종 보완")
    print("=" * 60)

    for fn in [fix_project2, fix_project3, fix_project4, fix_project5]:
        try:
            fn()
        except Exception as e:
            print(f"  [오류] {e}")
            import traceback
            traceback.print_exc()

    # 검증
    print("\n검증:")
    checks = [
        (2, "분말검사", "서제완", "프레스-IM"),
        (3, "연삭측정제어", "김웅기", "외부스케일"),
        (4, "포장혼입검사", "조윤명", "Korloy"),
        (5, "호닝신뢰성", "김준형v", "72시간"),
    ]
    for pid, slug, pm_name, keyword in checks:
        path = BASE / f"연구개발계획서-{pid}-{slug}.docx"
        doc = Document(str(path))
        pm_ok = doc.tables[2].rows[1].cells[1].text.strip() == pm_name
        kw_ok = any(keyword in p.text for p in doc.paragraphs)
        # 표지 확인
        cover_ok = False
        for elem in doc.element.body.iter():
            tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ''
            if tag == 'txbxContent':
                for t in elem.iter():
                    if isinstance(t.tag, str) and etree.QName(t.tag).localname == 't':
                        if t.text and slug not in "분말검사" and len(t.text) > 5:
                            cover_ok = True
        if not cover_ok:
            cover_ok = any("개발과제명" not in p.text and len(p.text) > 10 for p in doc.paragraphs[:7])
        print(f"  #{pid} {slug}: PM={pm_ok}, 내용={kw_ok}")

    print("\n완료")
