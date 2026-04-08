# -*- coding: utf-8 -*-
"""연구개발 HTML → 연구노트 docx 업데이트 스크립트

소스 HTML에서 내용을 추출하여 기존 docx 템플릿의 본문 셀(테이블 2~5)을 업데이트.
이미지는 적당한 사이즈로 축소하여 포함.
"""

import os
import re
import io
import sys
import base64
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

HTML_DIR = os.path.join('C:', os.sep, 'MES', 'wta-agents', 'workspaces', 'docs-agent')
DOCX_DIR = os.path.join('C:', os.sep, 'MES', 'wta-agents', 'reports', 'MAX', '경상연구개발')

PROJECTS = [
    ('장비물류', '장비 무인화운영을 위한 장비 물류 개발'),
    ('분말검사', '분말 비전검사 성능 고도화'),
    ('연삭측정제어', '연삭·측정·제어 통합 정밀도 향상'),
    ('포장혼입검사', '포장 혼입검사 시스템 개선'),
    ('호닝신뢰성', '호닝 형상검사 신뢰성 확보'),
]

# 이미지 최대 폭 (cm)
IMG_MAX_WIDTH_CM = 12

# Confluence 메타데이터 제거: 멀티라인 패턴 (줄바꿈 포함 텍스트에서 제거)
MULTILINE_REMOVE = [
    # "자동 생성됨/생성된" 멘트만 (연구노트 자동생성 안내문)
    re.compile(r'[^\n]*(?:본\s*연구노트는|이\s*문서는)[^\n]*자동\s*생성[^\n]*', re.IGNORECASE),
    # image-meta.json 언급 (줄 전체)
    re.compile(r'[^\n]*image-meta\.json[^\n]*', re.IGNORECASE),
    # 페이지 ID (멀티라인: "페이지 ID" + 줄바꿈 + ": 숫자")
    re.compile(r'페이지\s*ID\s*[:：]?\s*\n?\s*[:：]?\s*\d+', re.IGNORECASE),
    # 원본 링크 (멀티라인: "원본 링크" + URL)
    re.compile(r'원본\s*링크\s*[:：]?\s*\n?\s*[:：]?\s*https?://[^\s\n]+(?:atlassian|confluence)[^\s\n]*', re.IGNORECASE),
    # atlassian URL 단독 행
    re.compile(r'^\s*https?://[^\s]*atlassian[^\s]*\s*$', re.MULTILINE),
    # 이미지 파일명 텍스트 (인라인)
    re.compile(r'image-\d{8}-\d{6}\.\w+'),
]


def clean_text(text):
    """Confluence 메타데이터 텍스트 제거"""
    for pat in MULTILINE_REMOVE:
        text = pat.sub('', text)
    # 연속 빈 줄 정리 (3개 이상 → 2개)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


class HTMLContentExtractor(HTMLParser):
    """HTML에서 텍스트와 이미지를 순서대로 추출"""

    def __init__(self):
        super().__init__()
        self.elements = []  # list of ('text', str) or ('image', bytes)
        self._current_text = []
        self._in_tag = None
        self._skip = False
        self._heading_level = 0

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        if tag in ('h1', 'h2', 'h3', 'h4'):
            self._flush_text()
            self._in_tag = tag
            self._heading_level = int(tag[1])
        elif tag == 'img':
            self._flush_text()
            src = attrs_dict.get('src', '')
            if src.startswith('data:image'):
                # base64 인라인 이미지
                match = re.match(r'data:image/(\w+);base64,(.+)', src)
                if match:
                    fmt = match.group(1)
                    b64data = match.group(2)
                    try:
                        img_bytes = base64.b64decode(b64data)
                        self.elements.append(('image', img_bytes, fmt))
                    except Exception:
                        pass
        elif tag == 'br':
            self._current_text.append('\n')
        elif tag == 'li':
            self._current_text.append('\n· ')
        elif tag == 'p':
            self._current_text.append('\n')
        elif tag == 'tr':
            self._current_text.append('\n')
        elif tag == 'td' or tag == 'th':
            self._current_text.append('\t')
        elif tag == 'style' or tag == 'script':
            self._skip = True
        elif tag == 'strong' or tag == 'b':
            pass  # 볼드 마커 (docx에서 처리)

    def handle_endtag(self, tag):
        if tag in ('h1', 'h2', 'h3', 'h4'):
            text = clean_text(''.join(self._current_text).strip())
            if text:
                prefix = ''
                if self._heading_level == 2:
                    prefix = '\n\n■ '
                elif self._heading_level == 3:
                    prefix = '\n▸ '
                elif self._heading_level == 1:
                    prefix = ''
                self.elements.append(('heading', prefix + text, self._heading_level))
            self._current_text = []
            self._in_tag = None
        elif tag == 'p':
            self._current_text.append('\n')
        elif tag == 'table':
            self._current_text.append('\n')
        elif tag == 'style' or tag == 'script':
            self._skip = False
            self._current_text = []

    def handle_data(self, data):
        if self._skip:
            return
        self._current_text.append(data)

    def _flush_text(self):
        text = ''.join(self._current_text).strip()
        if text:
            # Confluence 메타데이터 제거
            text = clean_text(text)
            if text:
                self.elements.append(('text', text))
        self._current_text = []

    def finalize(self):
        self._flush_text()


def extract_html_sections(html_path):
    """HTML에서 개요, 세부내용을 섹션별로 추출"""
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    # body 내용만 추출
    body_match = re.search(r'<body[^>]*>(.*?)</body>', html, re.DOTALL)
    if body_match:
        body = body_match.group(1)
    else:
        body = html

    # h2 기준으로 섹션 분리
    parts = re.split(r'(?=<h2[^>]*>)', body)

    overview = None
    detail_sections = []

    for part in parts:
        h2_match = re.match(r'<h2[^>]*>(.*?)</h2>', part, re.DOTALL)
        if h2_match:
            title = re.sub(r'<[^>]+>', '', h2_match.group(1)).strip()
            if title == '개요':
                overview = part
            else:
                detail_sections.append((title, part))
        elif not overview and '<h1' in part:
            # h1 영역 (제목) - 스킵
            pass

    return overview, detail_sections


def parse_html_to_elements(html_fragment):
    """HTML 조각을 (type, content) 리스트로 변환"""
    parser = HTMLContentExtractor()
    parser.feed(html_fragment)
    parser.finalize()
    return parser.elements


def clear_cell(cell):
    """셀 내 모든 paragraph 제거 (첫 번째는 유지하되 비움)"""
    for i in range(len(cell.paragraphs) - 1, 0, -1):
        p = cell.paragraphs[i]._element
        p.getparent().remove(p)
    # 첫 번째 paragraph 비우기
    if cell.paragraphs:
        p = cell.paragraphs[0]
        for run in p.runs:
            run._element.getparent().remove(run._element)
        p.text = ''


def set_run_font(run, size_pt=9, font_name='맑은 고딕', bold=False):
    """run 폰트 설정"""
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run.font.bold = bold
    # 한글 폰트 설정
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_text_to_cell(cell, text, size_pt=9, bold=False, is_first=True):
    """셀에 텍스트 추가"""
    if is_first and cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(14)

    # 줄바꿈 처리
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            run = p.add_run()
            run._element.append(run._element.makeelement(qn('w:br'), {}))
        if line.strip():
            run = p.add_run(line)
            set_run_font(run, size_pt=size_pt, bold=bold)

    return p


def add_image_to_cell(cell, img_bytes, fmt='jpeg'):
    """셀에 이미지 추가 (적당한 크기로)"""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    img_stream = io.BytesIO(img_bytes)

    # 이미지 크기 확인 후 적절히 조절
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        w, h = img.size
        # 가로 기준 최대 12cm로 축소
        max_w_px = int(IMG_MAX_WIDTH_CM / 2.54 * 96)  # ~454px
        if w > max_w_px:
            ratio = max_w_px / w
            new_w = int(w * ratio)
            new_h = int(h * ratio)
            img = img.resize((new_w, new_h), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format='JPEG', quality=80)
            img_stream = io.BytesIO(buf.getvalue())
            w, h = new_w, new_h

        # docx 삽입 폭 계산 (최대 12cm)
        width_cm = min(w / 96 * 2.54, IMG_MAX_WIDTH_CM)
        run.add_picture(img_stream, width=Cm(width_cm))
    except ImportError:
        # PIL 없으면 고정 크기로
        run.add_picture(img_stream, width=Cm(10))

    return p


def build_section_text(overview_html, detail_sections_html):
    """HTML 섹션들을 docx 셀별 텍스트+이미지로 구성

    원문 텍스트를 한 글자도 빠짐없이 포함 (요약/축소 금지).
    Confluence 메타데이터(자동생성 멘트, 이미지파일명, 페이지ID, 원본링크)만 제거.

    Returns:
        dict with keys: 'goal', 'method', 'results'
        각 값은 [(type, content), ...] 리스트
    """

    # 개요에서 연구 목표 추출 — 전문 포함
    goal_elements = []
    if overview_html:
        elements = parse_html_to_elements(overview_html)
        for elem in elements:
            if elem[0] in ('text', 'image', 'heading'):
                goal_elements.append(elem)

    # 세부 섹션 — 모든 텍스트+이미지를 빠짐없이 포함
    # method: 핵심내용(텍스트 위주), results: 주요이미지+첨부자료(이미지 위주)
    method_elements = []
    results_elements = []

    for title, section_html in detail_sections_html:
        elements = parse_html_to_elements(section_html)
        current_target = method_elements
        for elem in elements:
            if elem[0] == 'heading':
                text = elem[1]
                if '주요 이미지' in text or '관련 첨부' in text:
                    current_target = results_elements
                elif '핵심 내용' in text:
                    current_target = method_elements
            current_target.append(elem)

    # results에 이미지가 없으면 method의 이미지를 이동
    if not any(e[0] == 'image' for e in results_elements):
        new_method = []
        for elem in method_elements:
            if elem[0] == 'image':
                results_elements.append(elem)
            else:
                new_method.append(elem)
        method_elements = new_method

    return {
        'goal': goal_elements,
        'method': method_elements,
        'results': results_elements,
    }


def write_elements_to_cell(cell, elements, is_content_cell=True):
    """elements 리스트를 docx 셀에 쓰기"""
    clear_cell(cell)
    is_first = True

    for elem in elements:
        if elem[0] == 'text':
            text = elem[1].strip()
            if not text:
                continue
            add_text_to_cell(cell, text, size_pt=9, is_first=is_first)
            is_first = False
        elif elem[0] == 'heading':
            text = elem[1].strip()
            level = elem[2] if len(elem) > 2 else 2
            size = 10 if level <= 2 else 9
            add_text_to_cell(cell, text, size_pt=size, bold=True, is_first=is_first)
            is_first = False
        elif elem[0] == 'image':
            img_bytes = elem[1]
            fmt = elem[2] if len(elem) > 2 else 'jpeg'
            try:
                add_image_to_cell(cell, img_bytes, fmt)
                is_first = False
            except Exception as e:
                print(f'  이미지 삽입 실패: {e}')


def update_docx(project_idx, project_name, project_title):
    """단일 docx 파일 업데이트"""
    html_path = os.path.join(HTML_DIR, f'연구개발-{project_idx}-{project_name}.html')
    docx_path = os.path.join(DOCX_DIR, f'연구노트-{project_idx}-{project_name}.docx')

    print(f'\n=== #{project_idx} {project_name} ===')
    print(f'  HTML: {os.path.basename(html_path)} ({os.path.getsize(html_path):,} bytes)')
    print(f'  DOCX: {os.path.basename(docx_path)} ({os.path.getsize(docx_path):,} bytes)')

    # HTML 파싱
    overview, detail_sections = extract_html_sections(html_path)
    print(f'  개요: {"있음" if overview else "없음"}, 세부섹션: {len(detail_sections)}개')

    # 섹션별 컨텐츠 구성
    content = build_section_text(overview, detail_sections)

    # DOCX 열기
    doc = Document(docx_path)

    # 테이블 2: 연구 목표 (행1이 내용)
    if content['goal']:
        print(f'  테이블2(연구목표): {len(content["goal"])}개 요소')
        write_elements_to_cell(doc.tables[2].rows[1].cells[0], content['goal'])

    # 테이블 3: 실험/개발 내용 (행1이 내용)
    if content['method']:
        print(f'  테이블3(실험/개발): {len(content["method"])}개 요소')
        write_elements_to_cell(doc.tables[3].rows[1].cells[0], content['method'])

    # 테이블 4: 결과 및 데이터 (행1이 내용)
    if content['results']:
        print(f'  테이블4(결과/데이터): {len(content["results"])}개 요소')
        write_elements_to_cell(doc.tables[4].rows[1].cells[0], content['results'])

    # 테이블 5: 고찰 및 분석 - 기존 내용 유지 (HTML에 별도 분석 섹션이 없으면)
    # 기존 내용이 이미 잘 작성되어 있으므로 유지

    # 저장
    doc.save(docx_path)
    new_size = os.path.getsize(docx_path)
    print(f'  저장 완료: {new_size:,} bytes')


def main():
    print('연구노트 docx 업데이트 시작')
    print(f'HTML 소스: {HTML_DIR}')
    print(f'DOCX 대상: {DOCX_DIR}')

    for idx, (name, title) in enumerate(PROJECTS, 1):
        try:
            update_docx(idx, name, title)
        except Exception as e:
            print(f'\n  ERROR #{idx} {name}: {e}')
            import traceback
            traceback.print_exc()

    print('\n\n모든 연구노트 업데이트 완료!')


if __name__ == '__main__':
    main()
