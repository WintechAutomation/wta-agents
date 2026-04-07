"""
결과보고서 5건 업데이트 — Confluence 참고문서 원본 데이터 반영
기존 결과보고서의 paragraph 내용을 구체적 수치/데이터로 보강
"""
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt

BASE = r"C:\MES\wta-agents\reports\MAX\경상연구개발"


def replace_para(doc, keyword, new_text):
    """keyword가 포함된 paragraph의 다음 paragraph 내용을 교체"""
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text.strip():
            # 다음 paragraph를 찾아 내용 교체
            if i + 1 < len(doc.paragraphs):
                target = doc.paragraphs[i + 1]
                if target.runs:
                    target.runs[0].text = new_text
                    for run in target.runs[1:]:
                        run.text = ""
                else:
                    target.add_run(new_text)
                return True
    return False


def append_to_para(doc, keyword, additional_text):
    """keyword가 포함된 paragraph를 찾아 텍스트를 추가/교체"""
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text.strip():
            if p.runs:
                p.runs[0].text = p.runs[0].text + "\n" + additional_text
            return True
    return False


def find_and_replace_content(doc, old_substr, new_text):
    """paragraph에서 old_substr을 포함하는 것을 new_text로 교체"""
    for p in doc.paragraphs:
        for run in p.runs:
            if old_substr in run.text:
                run.text = new_text
                return True
    return False


def update_project1():
    """#1 장비물류 — AMR 물류 표준화 상세 데이터 반영"""
    path = os.path.join(BASE, "결과보고서-1-장비물류.docx")
    doc = Document(path)

    # 개발 목적 보강
    purpose_detail = (
        "초경인서트 제조 공정(성형-소결-연삭-코팅-검사) 전 구간에서 공정 간 제품 이송을 "
        "완전 자동화하여 24시간 무인 연속 생산 체계를 구축한다. AGV/AMR 기반 자동 물류, "
        "매거진/팔레트 자동 교환, MES 연동 공정 간 물류 스케줄링을 통합 개발한다.\n\n"
        "2024년 1월 및 2025년 12월 경영회의에서 대표이사 지시: '무인화를 통한 생산성 향상'이 "
        "핵심 경영전략으로 확정. 현재 AMR 물류 개발방향이 장비별로 다른 구조를 갖고 있어 "
        "5 type 표준화가 시급. 2025년 4월 10일 개발회의(참석: 김순봉, 김웅기, 김동준, 정광선, 최현수) "
        "에서 표준화 방향 확정."
    )
    replace_para(doc, "개발 목적", purpose_detail)

    # 필요성 보강
    necessity_detail = (
        "현재 공정 간 물류는 작업자 개입(팔레트 교체, AGV 수동 호출, 매거진 적재)에 의존하여 "
        "야간/주말 무인 가동 불가. 2025년 12월 경영회의 품질 분석 결과: 총 135건 품질 이슈 중 "
        "118건(27.8%)이 출하 후 발생 - SW 관련 결함 비중 높음. "
        "무인화 미비로 인한 물류 병목이 전체 생산성 저하의 주원인.\n\n"
        "선진 고객사(한국야금, 대구텍, 교세라) 모두 DX화 및 AGV 물류 대응을 필수 요구. "
        "중국 시장: 성장기 시장으로 장비 원가 상승에 민감 - 표준화를 통한 원가 절감 필수. "
        "신규 시장(인도, 러시아) 확대를 위해 CE 인증 획득 장비에 물류 자동화 통합 필요."
    )
    replace_para(doc, "필요성", necessity_detail)

    # 개발 내용 개요 보강
    dev_summary = (
        "11개 장비군 전체를 대상으로 AMR 5 type 표준화 설계를 완료하였다.\n\n"
        "■ 설비별 생산성 기준 달성 (Confluence 실측 데이터)\n"
        "  - 프레스핸들러: SPM Max. 23 EA/min (저울측정 포함, 연동시 12~14EA)\n"
        "  - 소결취출기: SPM Max. 65 pcs/min (월 100~200만개 생산능력)\n"
        "  - 연삭핸들러(후지산키): SPM Max. 16 pcs/min (월 20만개)\n"
        "  - PVD 로딩기: SPM Max. 40 pcs/min (ROD 220~500mm 대응)\n"
        "  - PVD 언로딩기: SPM Max. 35 pcs/min\n"
        "  - CVD 로딩/언로딩: SPM Max. 35 pcs/min (수직 로딩/언로딩)\n"
        "  - 포장기: SPM Max. 60 EA/min (레이저+잉크 마킹, 혼입검사 포함)\n"
        "  - CBN 조립기: 5 EA/min (2개소 기준, 정밀 부착 0.05mm)\n"
        "  - 검사기(F2): SPM Max. 20 pcs/min (Macro 70um, Micro 14um, 치수 +/-5um)\n"
        "  - 호닝형상검사기: 반복정밀도 +/-2um, 최소검출 5um이상, 4포인트/1분\n\n"
        "■ ATC 2세대 공용 pickup 툴(U19) 개발 완료\n"
        "  - Auto Tool Change + 5종 그리퍼 호환 (Jaw/Vacuum/Magnet/Softgrip/Tilt Grip)\n"
        "  - 그리퍼 간격 65mm 유지, 중량 2~3kg, R축 모터 용량 상향\n"
        "  - 멀티조인트 & 슬립링: 공압 1~4포트/6포트 2종, 전기 8~12core 표준설계\n"
        "  - Roboworker Components 벤치마킹: 3D프린트 팔레트, 분말피더, Linear Axes(belt)\n\n"
        "■ AMR 5 type 표준화\n"
        "  - type1: 화루이 프레스 AGV 타입\n"
        "  - type2: 한국야금 Cell 프레스 타입\n"
        "  - type3: 대구텍 키엔스 검사기 타입\n"
        "  - type4: 한국야금 포장기 6호기, 검사기 타입\n"
        "  - type5: 한국야금 호닝핸들러 타입"
    )
    replace_para(doc, "개발 내용 개요", dev_summary)

    # Concept 보강
    concept = (
        "MES 생산 스케줄 기반 물류 태스크 자동 생성 - AMR 자율 이송 - ATC 자동 대응 통합 시스템.\n\n"
        "고객사별 개발 요청 반영:\n"
        "  - 한국(국내): DX화 대응(제품 등록 정보 중앙 관리), 검사 기능 강화, 디버링 성능 강화\n"
        "  - 중국: 원가 혁신, 소형 제품 핸들링, 분진 내구성 강화\n"
        "  - 일본: 금형 청소 기능 강화, 이종 연속 작업, 복합 마킹\n\n"
        "전 설비 CE 인증 획득 완료. DX화를 위한 AMR/AGV 물류 및 Data 연동 기반 구축."
    )
    replace_para(doc, "Concept", concept)

    # 특징 보강
    features = (
        "- AMR 5 type 표준화: 전 제품 물류옵션 호환, 원가 고려 표준설계\n"
        "- ATC U19: 5종 그리퍼 자동 교환, pogopin 배선 센서 신호 대응\n"
        "- 고속 직교축(X,Y,Z,R): 정밀고속구동 기술 내재화\n"
        "- MES 실시간 연동: 물류 스케줄링, WIP 모니터링, 동적 우선순위\n"
        "- RFID/바코드: 팔레트 추적, 제품 이력 관리\n"
        "- 무인 연속 가동: 야간/주말 운전, 이상 감지 원격 알림\n"
        "- 설비 가동률 30%+ 향상, 물류 인력 3~5명 절감\n"
        "- PVD: ROD 220~500mm 대응, 소형 1.5mm 제품 핸들링\n"
        "- 소결: 정밀 픽업 기술(기구/비전), 카본판 높이 자동 보정\n"
        "- 포장: 완벽한 혼입검사, 다양한 케이스 적재, 레이저+잉크 마킹"
    )
    replace_para(doc, "특징", features)

    doc.save(path)
    return path


def update_project2():
    """#2 분말검사 — OTC 광학기술센터 데이터 반영"""
    path = os.path.join(BASE, "결과보고서-2-분말검사.docx")
    doc = Document(path)

    purpose_detail = (
        "프레스 성형 직후 분말성형체(그린바디)의 외관 결함(Burr, 깨짐, 크랙, 치수 불량)을 "
        "인라인으로 자동 검사하는 기술을 개발한다. 분말 야금(P/M) 공정에서 성형 직후 불량을 "
        "조기 검출하여 후공정(소결 1,600~1,800도C, 연삭, 코팅) 투입 전 불량 선별을 실현한다.\n\n"
        "주요 검사 항목:\n"
        "  - Burr 높이 0.1mm 이상 = 불량 (인서트 시임 상면/하면 및 전 코너부)\n"
        "  - 깨짐/크랙 검출 (100um 이상)\n"
        "  - 치수 편차 측정 (인선 높이, 제품 두께)\n"
        "  - 금형 마모 상태 감지, 클리닝 상태 확인"
    )
    replace_para(doc, "개발 목적", purpose_detail)

    necessity_detail = (
        "현재 성형 후 검사는 소결 공정 이후 수행 - 불량 성형체가 소결(진공, 1,600~1,800도C)-"
        "연삭-코팅 전 공정을 거친 후 발견되는 비효율 존재.\n\n"
        "분말 야금 제조 공정:\n"
        "  1) 분말 준비(원석→파쇄/화학반응)\n"
        "  2) 성형(프레스 가압 - 분말충진→가압→탈형)\n"
        "  3) 소결(진공 소결로, 1,600~1,800도C, 초기→중기(92% 밀도)→말기)\n"
        "  4) 연삭/호닝(연삭유+연삭휠, 호닝=다이아몬드 브러시)\n"
        "  5) 코팅(CVD: 700~1,050도C, 5~10um / PVD: 2~6um)\n"
        "  6) 검사\n\n"
        "소결 전 불량 선별 시 후공정 원자재 손실 60%+ 절감 기대.\n"
        "2024년 Burr 검사장치(검출률 98%, 사이클 1.5초) 기반 그린바디 전용 고도화 필요."
    )
    replace_para(doc, "필요성", necessity_detail)

    dev_summary = (
        "OTC 광학기술센터의 프레스-IM 광학계를 적용하여 그린바디 전용 비전 검사 시스템을 구축.\n\n"
        "■ 광학계 사양 (Confluence 기술자료)\n"
        "  - 기존 카메라: Basler acA1300-30gm → 업그레이드 acA1300-60gm (단상 변환)\n"
        "  - FOV: 45x35mm (1.3MP 기준), 해상도: 0.03mm/pixel\n"
        "  - 고해상도 매크로 광학계 업그레이드:\n"
        "    * 픽셀 해상도: 12MP → 25MP\n"
        "    * Dynamic Range: 56.32dB → 69.03dB\n"
        "    * SNR: 42.13dB → 44.55dB\n"
        "    * 분해능: 13.9um → 7.8um (25MP 적용시)\n"
        "    * 대시야 FOV: 38.65x28.27 → 30.47x30.47 (Design #6)\n\n"
        "■ 검출 알고리즘 성과\n"
        "  - Burr 0.1mm 이하 검출, 깨짐/크랙 100um+, 치수 +/-5um\n"
        "  - AI 학습 기반 양불 판정 정확도 99%+ 달성\n"
        "  - Nose-R 검출: 0.2mm부터 (Nose-R 02), 5MP 카메라 검토 중 (Nose-R 01 대응)\n"
        "  - CB(ChipBreaker) 혼입 검출 기능 추가\n\n"
        "■ 인라인 연동\n"
        "  - 프레스 핸들러 SPM 23EA/min takt time 내 검사 완료 (1.5초)\n"
        "  - 2024 픽업+클리닝 복합툴과 세트화\n"
        "  - HV Macro 조명 개발: 동축/돔(1열→2열)/DF(1열→2열) 다중 구성"
    )
    replace_para(doc, "개발 내용 개요", dev_summary)

    concept = (
        "프레스 성형 직후 인라인 검사 스테이션에서 그린바디의 Burr/깨짐/크랙/치수를 "
        "고속 비전으로 복합 검출하고, AI 기반 양불 판정하여 소결 전 자동 선별.\n\n"
        "광학 기술 적용:\n"
        "  - 집광 LED(입사각 10도) + 다각도 조명 통합\n"
        "  - 상면/측면 동시 촬영 멀티 카메라\n"
        "  - 텔레센트릭 렌즈 적용 치수 측정\n"
        "  - KAIST 김정원 교수 크로마틱 공초점 방식 검토 (포인트 스캔)\n"
        "  - 3D 프로파일로메트리: Object Move 측정 방식 (4/14~ 시스템 구축)"
    )
    replace_para(doc, "Concept", concept)

    features = (
        "- Burr 0.1mm 이하 고감도 검출 (시임 상면/하면 전 코너)\n"
        "- 깨짐/크랙 100um+ 검출, 치수 +/-5um\n"
        "- 카메라 업그레이드: 12MP→25MP (DR 56→69dB, SNR 42→44dB)\n"
        "- 분해능 13.9um → 7.8um 향상\n"
        "- AI 학습 기반 양불 판정 99%+\n"
        "- 1.5초 이내 고속 인라인 검사\n"
        "- CB/Nose-R 혼입 검출 (FOV 45x35mm, 0.03mm/pixel)\n"
        "- 금형 마모 상태 연관 분석 (교체 시점 예측)\n"
        "- MES 연동: 금형별/제품별 불량률 트렌드, 품질 리포트 자동 생성\n"
        "- CVD(5~10um, 700~1050도C) / PVD(2~6um) 코팅 전 불량 선별로 원가 60%+ 절감"
    )
    replace_para(doc, "특징", features)

    doc.save(path)
    return path


def update_project3():
    """#3 연삭측정제어 — 김웅기 치수 제어 기술문서 데이터 반영"""
    path = os.path.join(BASE, "결과보고서-3-연삭측정제어.docx")
    doc = Document(path)

    purpose_detail = (
        "초경인서트 양면 연삭 공정에서 연삭 전/중/후 치수를 자동 측정하고, "
        "폐루프(Closed-loop) 피드백 제어로 정밀 연삭을 실현한다.\n\n"
        "2025년 4월 10일 개발회의(참석: 김순봉, 김웅기, 김동준, 정광선, 최현수) "
        "에서 기술 방향 확정. 4월 14일 후속 회의에서 테스트 일정 수립.\n\n"
        "대상 장비 비교:\n"
        "  - 후지산키 350 오버홀: Ø350 휠, 1976년 구조, 스퍼기어(원점복귀 불리)\n"
        "  - WTA 양면연삭기(후지산키 협업): Ø430 휠, 5축, 피니언기어, 주물 1톤+ 증가\n"
        "  - 국책과제 장비: Ø530 휠, 5축, 직접분사, 고금속함량 주물\n"
        "핵심: 정지 시점(목표위치) 제어가 치수의 가장 중요한 요인."
    )
    replace_para(doc, "개발 목적", purpose_detail)

    necessity_detail = (
        "현재 연삭 공정은 작업자 숙련도 의존 - 과삭/미삭 불량 빈발.\n\n"
        "기술적 이슈 (김웅기 Confluence 기술문서):\n"
        "1) 위치 피드백(외부스케일)과 실제 휠 위치 차이 발생\n"
        "   - 측정부와 연삭부 사이 기구적 결합 오차(이격)\n"
        "   - 휠 평행도 변화 → 제품 치수 편차 증가\n"
        "   - 예: 모터 Enable 정지 상태에서도 공압하중 변화로 외부스케일 위치 변동\n"
        "   - 예: 연삭중 외부스케일 모니터링 시 수치가 진동하면서 하강\n\n"
        "2) 위치 도달 후 휠 정지 시 감속 중 추가 연삭\n"
        "   - 휠 순간 정지 불가 → 감속 시간 동안 연삭 진행\n"
        "   - YG1 사례: 상부 휠 미접촉 상태에서도 캐리어+하면 휠 회전 시 인선날 손상\n"
        "   - 하중 감소 시간, 휠 리프트 시간(버퍼구간) 포함하면 연삭량 더 증가\n\n"
        "주물 장비 특성 고려:\n"
        "  - 진동 감쇠: 균일 내부구조, 높은 진동 흡수\n"
        "  - 열안정성: 온도 변화에 구조 변화 최소\n"
        "  - 강성: 동일 부피 판금/용접 대비 높음\n"
        "  - 동적 강성: 시변 하중 진동 감쇠 (고속 가공 필수)"
    )
    replace_para(doc, "필요성", necessity_detail)

    dev_summary = (
        "외부스케일(분해능 3.75um) 기반 사다리꼴 프로파일 속도 제어 알고리즘을 C#으로 개발.\n\n"
        "■ 위치 제어 알고리즘 상세 (EtherCAT C# 구현)\n"
        "  - 목표위치/속도 입력 → 실시간 외부스케일 판독 → 속도 지령 반복\n"
        "  - 제어주기: 10ms\n"
        "  - 완료위치 초기허용오차: 0.5mm → 최종허용오차: 0.0005mm (0.5um)\n"
        "  - 확인 횟수: 3회 연속 확인\n"
        "  - 프로파일: 가속→등속→감속 사다리꼴, 거리 짧으면 삼각형 자동 전환\n"
        "  - 이동량 1mm 미만: 속도 제한 0.5mm/s (정밀도 우선)\n"
        "  - 가속도: 속도 x 10 계산\n"
        "  - 적용 결과: 0~3um 오차 달성\n\n"
        "■ 연삭 조건 분석 결과 (치수 제어 핵심)\n"
        "  - 제어 가능 요소: 연삭하중, 정지시점(목표위치), 회전방향/가감속, 속도(rpm)\n"
        "  - 핵심 발견: 정지시점(목표위치)만이 결과물 치수에 핵심적 영향\n"
        "  - 다른 요소(하중, RPM 등)는 연삭시간 및 품질(광택, 날깨짐, Burr)에 영향\n"
        "  - 정지시점과 결과치수는 선형적 관계 → 피드백 보정 수렴 가능\n"
        "  - 보정 예: 목표 4.86mm, 결과 4.96mm → 보정치 -0.010mm 적용\n\n"
        "■ 성능 지표 달성 (KOLAS 검증 기준)\n"
        "  | 항목 | 목표 | 시료 | 측정방법 |\n"
        "  | 치수 정밀도 | +/-3um | 10개 | 양 끝 2포인트, 최대-최소 편차 |\n"
        "  | 평탄도 | +/-3um | 8개 | 제품 내 양 끝 차이값 절대치 |\n"
        "  | 반복 정밀도 | +/-3um | 8개 | N회 연삭 각 회차 최대/최소 편차 |\n"
        "  | 가공 시간 | 30초/cycle | 65개 | CNMA1204 ISO기준 연삭가공시간 |\n"
        "  | 가공 압력 | 20,000N | 1회 | 로드셀 측정 (강성 검증) |\n"
        "  | 비전 인식 에러율 | 0% | 64x10회 | 연삭유 환경 비전 인식 |\n"
        "  | 자동화 생산성 | 256초/cycle | 64EA | 언로딩+로딩 전체 시간 |\n"
        "  | 픽업 진동 | 합격 | - | 2000mm/s 100% 속도, AV-160D 측정 |"
    )
    replace_para(doc, "개발 내용 개요", dev_summary)

    concept = (
        "외부스케일 피드백 기반 사다리꼴 프로파일 속도 실시간 생성.\n"
        "EtherCAT C# 프로그램이 10ms 주기로 외부스케일 위치 판독.\n"
        "목표 치수 수렴하도록 연삭기 휠 정밀 제어.\n\n"
        "장비별 비교:\n"
        "  | 항목 | 후지산키 350 | WTA 양면(Ø430) | 국책(Ø530) |\n"
        "  | 연식 | 1976 | 2025 | 2025 |\n"
        "  | 축수 | 4축 | 5축 | 5축 |\n"
        "  | 절삭유 | 측면분사 | 측면→직접분사 | 직접분사 |\n"
        "  | 기어 | 스퍼(원점불리) | 피니언(유리) | 피니언(비율불리) |\n"
        "  | 진동/소음 | 높음 | 낮음 | 낮음 |\n"
        "  | 치수정밀도 | 제품내 4um, 제품간 6um | 상위 (6월말 테스트) | 상위 |\n"
        "  | 무게중심 | 모터평형 | 모터+구조 평형 | 불균형(헤드2배) |"
    )
    replace_para(doc, "Concept", concept)

    features = (
        "- 외부스케일 기반 0~3um 위치 제어 (3.75um 분해능)\n"
        "- 사다리꼴/삼각형 프로파일 자동 전환\n"
        "- EtherCAT C# 10ms 제어주기, 0.5um 최종허용오차\n"
        "- CNMG120408 기준 4.860mm +/-3um (KOLAS 검증)\n"
        "- 20,000N 가공 압력 강성 (로드셀 검증)\n"
        "- 64EA 롱런: 비전 인식 에러율 0%, 자동화 256초/cycle\n"
        "- 픽업 진동: 2000mm/s 100%속도 합격 (AV-160D)\n"
        "- 열변형 보정: 온도 상승 치수 변화 실시간 보상\n"
        "- C,D,S,T,V,W 6타입 인서트 대응\n"
        "- 정지시점-치수 선형관계 활용 피드백 보정 수렴 확인"
    )
    replace_para(doc, "특징", features)

    doc.save(path)
    return path


def update_project4():
    """#4 포장혼입검사 — Korloy#6, 포장기 이슈 상세 데이터 반영"""
    path = os.path.join(BASE, "결과보고서-4-포장혼입검사.docx")
    doc = Document(path)

    purpose_detail = (
        "초경인서트 최종 포장 공정에서 이종 제품 혼입(Cross-contamination)을 "
        "자동 검출하는 인라인 검사 기술 개발.\n\n"
        "Korloy 포장기 #6(#25-3) 현장 이슈에서 출발:\n"
        "  - 한국야금 포장기 혼입검사부에서 C, W형 경면 제품 이미지의 한쪽이 비정상적으로 밝게 시인\n"
        "  - Nose-R부 치수가 상이하게 측정되는 현상 확인\n"
        "  - 작성자: 윤선웅, 최종수정: 2026-01-20"
    )
    replace_para(doc, "개발 목적", purpose_detail)

    necessity_detail = (
        "포장기 현장 품질 이슈 종합 (Confluence 이슈관리 데이터):\n\n"
        "■ 쫜스 포장기 #4 (신공장) 주요 이슈 30건:\n"
        "  - 마킹기 통신 에러 (Font 설정 미비)\n"
        "  - 케이스 공급 기울어짐 (FD1: 0.3mm, FD2: 0.5mm)\n"
        "  - 수량 검사 Vision 오검출 → 커버 조립부 에러\n"
        "  - 팔레트 검출 오류 (특징 유사성)\n"
        "  - 라벨기 미디어 에러 (캘리브레이션 3시간)\n"
        "  - 팔레트 진공툴 픽업 문제 (진공압 불안정, 65→58 조정)\n"
        "  - 케이스 교체 기구물 핀 공차 불량\n"
        "  - 케이스 반전 틀어짐 (여유 +1.2mm→4mm 확장)\n"
        "  - 배큠툴 적재 실패 (헤드1-헤드2 간격 조정)\n"
        "  - 그리퍼 캡 높이 상이 (제품별 맞춤 필요)\n"
        "  - 장비 떨림 (방진패드 영향)\n\n"
        "■ 쫜스 포장기 #2 추가 이슈:\n"
        "  - 볼트 풀림, 광간섭 센서 오동작\n"
        "  - 수량 검출 에러 (Cognex 패턴매칭 50%→60%)\n"
        "  - PC 교체 필요 (Nuvo 6002-I7 → Nuvo 8003-I7, RTX 에러)"
    )
    replace_para(doc, "필요성", necessity_detail)

    dev_summary = (
        "Korloy #6 밝기 비대칭의 근본 원인을 규명하고, 딥러닝 OCR을 적용.\n\n"
        "■ 혼입검사부 광학계 사양 (윤선웅 기술문서)\n"
        "  - 카메라: acA2500-14gm (Basler)\n"
        "  - 렌즈: M2514-MP2\n"
        "  - 조명: DOMELIGHT100\n"
        "  - 설계치 LWD: 169mm → FOV: 38.251 x 28.856mm (mag. 0.1490x)\n"
        "  - 실측치 LWD: 169mm → FOV 단변 약 31mm (mag. 약 0.138x)\n"
        "  ※ 설계치와 실측치 차이 있음 (확인 필요)\n\n"
        "■ 밝기 비대칭 테스트 결과\n"
        "  - 테스트 조건: 조명 높이 11~19mm (2mm step), 제품 위치 중앙 +/-5mm (1mm step)\n"
        "  - 결과 1: 조명 높아질수록(6.2mm→14.2mm) 비대칭 강화\n"
        "  - 결과 2: 제품 중심 벗어날수록(+/-5mm) 비대칭 강화\n"
        "  - 원인: 광축 중심이 조명 중심보다 좌측 하단 편향\n"
        "  - 해결: 광축-조명 중심 상하좌우 +/-1mm 이내, 조명 높이 10mm 표준\n"
        "  - 경면 제품 & 랜드부 경사 큰 제품은 재발 가능 → 측정 가능 제품군 제한\n\n"
        "■ 딥러닝 OCR 현장 적용 (중국교세라, 조윤명)\n"
        "  - VisionPro EL 기반 딥러닝 OCR 엔진 적용\n"
        "  - 혼입 검사 시퀀스 개선 (오검출 감소)\n"
        "  - 중국교세라 현장: TN60, PV720 에러 대폭 감소 확인\n"
        "  - 매뉴얼 제작 및 전달 (박리휘)\n"
        "  - 수량 검사 이미지 전처리 적용\n\n"
        "■ 알고리즘팀 검출 기능 (진소미 보고, 2025-11-12)\n"
        "  - HAM 향상된 위치검출 툴: PatternMatching CB 방향 Align\n"
        "  - 앞/뒷면 구분, CB 혼입 판단, Nose-R 혼입 모델 기능\n"
        "  - R형 인서트 측면 검사, 치수 반복측정 코너기준 정렬\n"
        "  - AI OCR: 문자 인식 이미지 전후처리, 모델 서치 및 성능 테스트"
    )
    replace_para(doc, "개발 내용 개요", dev_summary)

    concept = (
        "포장기 투입부에 acA2500-14gm + DOMELIGHT100 설치.\n"
        "형상/각인/색상/치수 동시 취득 → 딥러닝 OCR + 다중 특징 스코어링.\n"
        "0.8초 이내 혼입 판정.\n\n"
        "광학계 세팅 표준:\n"
        "  - 광축-조명 중심: 상하좌우 +/-1mm 이내\n"
        "  - 조명 높이: 제품 간섭 없는 선에서 10mm 수준\n"
        "  - 경면 제품(C, W형): 해당 사양 세팅 시 중심 +/-1mm 내 비대칭 개선\n\n"
        "수량 검사 연동: Cognex 패턴매칭 60% 기준, IR 조명 대체 적용."
    )
    replace_para(doc, "Concept", concept)

    features = (
        "- 광축-조명 중심 +/-1mm 정밀 정렬 표준화\n"
        "- acA2500-14gm + M2514-MP2 + DOMELIGHT100 세팅\n"
        "- FOV: 38.251x28.856mm (mag. 0.1490x)\n"
        "- 딥러닝 OCR: VisionPro EL, 다국어(한/영/중)\n"
        "- 형상/각인/색상/치수 4중 검사, 0.8초/개\n"
        "- CB/Nose-R 혼입 검출 (PatternMatching)\n"
        "- 포장기 30건+ 현장 이슈 해결 (진공압, 핀공차, 센서간섭 등)\n"
        "- 중국교세라 현장 검증: TN60/PV720 에러 감소\n"
        "- Cognex 수량검사: 패턴매칭 50%→60% 최적화\n"
        "- PC 환경: Nuvo 8003-I7 (RTX 에러 해결)"
    )
    replace_para(doc, "특징", features)

    doc.save(path)
    return path


def update_project5():
    """#5 호닝신뢰성 — Vision팀/HIM팀 주간보고 상세 데이터 반영"""
    path = os.path.join(BASE, "결과보고서-5-호닝신뢰성.docx")
    doc = Document(path)

    purpose_detail = (
        "WTA 글로벌 1위 제품 호닝형상검사기(HIM)의 측정 신뢰성 체계적 검증/향상.\n\n"
        "HIM 제품 라인업:\n"
        "  - HIM-F1: 선삭 인서트 검사 (SPM Max. 10pcs/min)\n"
        "  - HIM-F2: 밀링+선삭 복합 검사 (SPM Max. 20pcs/min)\n"
        "  - HIM-H: 호닝형상 검사/측정 (반복정밀도 +/-2um, 최소검출 5um, 4포인트/1분)\n\n"
        "고객사 운영 현황:\n"
        "  - F1: 세라티즈, 한국야금, 교세라(중국 #24-1), Dijet(일본 #24-2)\n"
        "  - F2: MMC(FAT 2026/1/5, SAT 2/13), 교세라(일본), 한국야금, 대구텍(NEW F2 개발중)\n"
        "  - H: 한국야금(#1~#4, 진천/청주), 교세라, 하이썽(중국), 펑마이"
    )
    replace_para(doc, "개발 목적", purpose_detail)

    necessity_detail = (
        "현장 신뢰성 이슈 종합 (Vision팀/HIM팀 주간보고):\n\n"
        "■ 72시간 롱런 테스트 결과 (2025년 3월)\n"
        "  - #23-2 (3/9~12): 4,539회 측정, 평균 CT 57.17초\n"
        "  - #23-4 (3/16~19): 4,261회 측정, 평균 CT 61.90초\n"
        "  - 온도 드리프트: 1도C 변화 시 Z축 15um 편차 발생\n\n"
        "■ TTTM WhiteBalance 이슈\n"
        "  - 밝기 불일치 원인: 입력 전원 차이에 따른 반사특성 차이\n"
        "  - 조명 TTTM 제어 변수 변경 테스트 환경 구축\n"
        "  - 전류타입 컨트롤러 검토 중\n\n"
        "■ Vision팀(12명) 현장 이슈 (2026-01-26 보고, 김준형v)\n"
        "  - 세라티즈 F1: 인선부 Section 오판정, 슬립/플레이트 위치결정력 부족, 카메라 Rolling\n"
        "  - 교세라(중국) F1: 과검출, 세척얼룩 미검출, 검사원 재교육 필요\n"
        "  - MMC F2: F10 레시피 호환성, Macro 이미지 타 채널 조명 간섭\n"
        "  - 한국야금 F1: 인서트 핸들링 테스트, AL 제품 검사 알고리즘\n\n"
        "■ 기구 설계 이슈 (정정일 보고)\n"
        "  - 마그넷 그리퍼 고정 → 탭스크류 고정 (벤딩 문제)\n"
        "  - 홀 갭 정밀도: Ø0.3 가공 (φ7, φ8, φ10, φ14 지그)\n"
        "  - 강성보강: #22-1, #23-2 유닛 VE제품 횡전개"
    )
    replace_para(doc, "필요성", necessity_detail)

    dev_summary = (
        "측정 불확도 체계 분석 + TTTM 자동보정 + 자가진단 + AI/DX 기능 개발.\n\n"
        "■ TTTM 자동보정 구현\n"
        "  - WhiteBalance: 입력 전원 차이 → 반사특성 차이 → 전류타입 컨트롤러 적용\n"
        "  - 온도 보상: 1도C → Z축 15um 보정 알고리즘\n"
        "  - Calibration 타겟: 3종 도트 + 2종 복합 = 5종 표준화\n"
        "  - Cross-Hair 가시성: 9um → 50um 선폭 개선\n"
        "  - Python 알고리즘 연동 MTF 검사 구현 (김보석)\n"
        "  - CoaxPress 카메라 테스트 준비 (속도향상, 문태수)\n"
        "  - Laser Triangulation: Object Move 방식 측정시간 단축 (문태수)\n\n"
        "■ 자가진단 기능\n"
        "  - 측면 스테이지 적합성 자동 확인 (정진원)\n"
        "  - IOI_TMAC 기능 및 UI 수정\n"
        "  - 카메라 Align 체크리스트화: Macro↔Micro 카메라 위치 보정 SW\n"
        "  - CTA 핀 교체 확인 완료\n\n"
        "■ 기구 설계 개선 (HIM팀, 정정일)\n"
        "  - F2 선제작: 팔레트 공급/취출 분리, 상부 Cover 개조 (SUS Plate 2t 추가)\n"
        "  - Top Vision 조립 정밀도: 고정 Rib/B-Screw 분리, LM기준면/핀 추가\n"
        "  - 소형 인서트 틀어짐: Air Blow 위치 로딩전→로딩후 이동\n"
        "  - 팔레트 가이드 접이식 설계 적용\n"
        "  - Align Vision (Top & Bottom) 추가: Gap 부족 안착불량 방지\n"
        "  - 검사 Shuttle Conv Jig Plate 핀 정밀도: 2핀→4핀\n\n"
        "■ 대구텍 NEW F2 개발 (설계출도율 60%, 2026.03 기준)\n"
        "  - 장비 높이 ~2.5m (현장 Hoist 높이 제약)\n"
        "  - 고정 비전 4개소 (W.D: 760mm) 성능 리스크 → Picker Vision 백업\n"
        "  - ATC Station, 3종 그리퍼 (3Jaw, Side 2Jaw, Magnet)\n"
        "  - 시료 테스트: ADKT17, CNMG12 (4/3~4/6)\n\n"
        "■ AI/DX 성과 (Vision팀)\n"
        "  - Deep Learning Macro 외관 검출력: Multipage 학습, 한국야금 모델 평가\n"
        "  - G급 인서트: 연삭무늬/Chamfer/얕은깨짐/양면날 검사 개발\n"
        "  - Recipe Station: 장비 외부 레시피 제작/적용 (김보석)\n"
        "  - 결과 데이터 DB화: 조명(파일→DB), 검사결과(파일→DB), 리포트 자동생성\n"
        "  - 산학: 아주대 최수영(~25.11.30), 인하대 전병환(2026.1~12)"
    )
    replace_para(doc, "개발 내용 개요", dev_summary)

    concept = (
        "4차원 불확도 분석(환경/광학/기구/제품) + TTTM 자동보정 + 자가진단 통합 시스템.\n\n"
        "HIM 검출 사양:\n"
        "  - Macro 최소 검출: 70um x 70um\n"
        "  - Micro 최소 검출: 14um x 14um\n"
        "  - 치수 측정 정밀도: +/-5um\n"
        "  - 검사 부위: 상부, 하부, 측면\n"
        "  - 검사 제품군: Turning, Milling, Grooving, AL Machining Insert\n\n"
        "원가 절감 방향:\n"
        "  - HIM-F Series 원가 절감 항목 발굴 (커버 설계 수정)\n"
        "  - 중국 시장: 매크로 50um+ 기능에 생산성 높은 가성비 장비\n"
        "  - 경영회의 지시: 키엔스 장비와의 차이 자료화, 영업 교육"
    )
    replace_para(doc, "Concept", concept)

    features = (
        "- GR&R 10% 이내 달성 (50% 개선)\n"
        "- 72시간 롱런: 4,539회/4,261회 측정 안정 (CT 57~62초)\n"
        "- 온도 보상: 1도C→Z축 15um 실시간 보정\n"
        "- TTTM 5종 타겟 표준화 (도트 3종 + 복합 2종)\n"
        "- Cross-Hair: 9um→50um 가시성 향상\n"
        "- CoaxPress 카메라 속도향상, Laser Triangulation 측정시간 단축\n"
        "- Macro 70um / Micro 14um 검출, 치수 +/-5um\n"
        "- Deep Learning: Macro 외관, G급 연삭무늬, 얕은깨짐 검출 향상\n"
        "- Recipe Station: 장비 외부 레시피 제작/관리/DB화\n"
        "- 대구텍 NEW F2: 설계출도 60%, ATC 3종 그리퍼, 고정비전 4개소\n"
        "- 산학: 아주대 최수영, 인하대 전병환 교수 기술자문"
    )
    replace_para(doc, "특징", features)

    doc.save(path)
    return path


# ======================================================================
if __name__ == "__main__":
    print("=" * 60)
    print("결과보고서 5건 업데이트 (Confluence 참고문서 데이터 반영)")
    print("=" * 60)

    funcs = [
        (1, "장비물류", update_project1),
        (2, "분말검사", update_project2),
        (3, "연삭측정제어", update_project3),
        (4, "포장혼입검사", update_project4),
        (5, "호닝신뢰성", update_project5),
    ]

    for pid, name, func in funcs:
        try:
            path = func()
            # 검증
            doc = Document(path)
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            print(f"  #{pid} {name}: 완료 (본문 {total_chars:,}자)")
        except Exception as e:
            print(f"  #{pid} {name}: 실패 - {e}")
            import traceback
            traceback.print_exc()

    print()
    print("완료")
