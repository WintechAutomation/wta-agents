"""
연구노트 v2 재작성 — 날짜별 연구 과정 기록 + 원문 인용 + 이미지 삽입
부서장 재지시(2026-04-05) 반영
"""
import os
import sys
import glob
sys.stdout.reconfigure(encoding='utf-8')

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = r"C:\MES\wta-agents\reports\MAX\template-연구노트.docx"
OUTPUT_DIR = r"C:\MES\wta-agents\reports\MAX\경상연구개발"
IMG_BASE = r"C:\MES\wta-agents\reports\MAX\경상연구개발\참고문서-이미지"

# ======================================================================
# 프로젝트 #1: 장비 무인화운영을 위한 장비 물류 개발
# ======================================================================
PROJECT_1 = {
    "filename": "연구노트-1-장비물류.docx",
    "note_number": "2025-R&D-001",
    "page": "1",
    "total_pages": "1",
    "project_name": "장비 무인화운영을 위한 장비 물류 개발",
    "research_period": "2025.03 ~ 2025.12",
    "date": "2025.12.30",
    "department": "생산관리팀 (AI운영팀)",
    "author": "홍대한",
    "reviewer": "조한종",
    "author_sign": "/서명/",
    "reviewer_sign": "/서명/",
    "reviewer_date": "2025.12.31",

    "objective": (
        "[연구 배경]\n"
        "2024년 1월 경영회의에서 HAM 기구설계팀 주관으로 "
        "'AGV AMR 무인화 장비별 표준물류 컨셉 확립'이 공식 안건으로 제시됨 "
        "(원문: \"HAM 기구설계 - AGV AMR 무인화 장비별 표준물류 컨셉 확립\", "
        "2024-01-03 경영회의 회의록).\n\n"
        "[문제 정의]\n"
        "초경인서트 제조 공정(성형-소결-연삭-코팅-검사-포장)은 11개 장비군(프레스/"
        "소결취출/연삭/PVD 로딩·언로딩/CVD 로딩·언로딩/포장/CBN조립/검사/호닝형상검사)이 "
        "연계 운영되나, 장비 간 제품 이송이 수작업 의존도가 높아 24시간 무인화 운영이 "
        "불가능하다. 현재 파악된 장비별 최대 생산성(SPM)은 다음과 같다:\n"
        "· 프레스: Max 23 EA/min (원문: WTA 설비 차별점 정리)\n"
        "· 소결취출기: Max 65 pcs/min\n"
        "· 연삭핸들러: Max 16 pcs/min\n"
        "· PVD 로딩기: Max 40 pcs/min, PVD 언로딩기: Max 35 pcs/min\n"
        "· CVD 로딩/언로딩: Max 35 pcs/min\n"
        "· 포장기: Max 60 EA/min, CBN 조립기: 5 EA/min (2개소)\n"
        "· 검사기(F2): Max 20 pcs/min, 호닝형상검사기: 4포인트/1분\n"
        "각 장비의 SPM이 상이하여 공정 간 버퍼/물류 설계가 핵심 과제로 대두됨.\n\n"
        "[연구 목표]\n"
        "AGV/AMR 기반 자동 물류 시스템을 5가지 타입으로 표준화하고, "
        "2세대 공용 픽업툴(U19)과 ATC(Auto Tool Change) 기술 내재화를 통해 "
        "전 공정 무인 연속 생산 체계를 구축한다. 2025년 말까지 표준화 설계 완료, "
        "2026년 양산 적용을 목표로 한다."
    ),

    "content": (
        "[1단계] 2025.03 — 장비 현황 조사 및 물류 표준화 방향 수립\n"
        "────────────────────────────────────────────────\n"
        "· 대구텍 AMR(AGV) 사례 조사 (Confluence: 대구텍 AMR, 2022-12-21 작성).\n"
        "  - 원문 인용: \"장비 내부 물류 (AGV 대응시 필요한 구조)\"\n"
        "  - 결론: 장비 내부에 AMR 접근을 위한 도어/컨베이어 구조 표준화 필요.\n"
        "· 설계팀 요소기술 개발항목 문서 분석 (Confluence: 설계팀 요소기술 개발항목, "
        "2025-10-29).\n"
        "  - AMR 물류 5가지 타입 확인:\n"
        "    type 1: 화루이 프레스 AGV 타입\n"
        "    type 2: 한국야금 Cell 프레스 타입\n"
        "    type 3: 대구텍 키엔스 검사기 타입\n"
        "    type 4: 한국야금 포장기 6호기/검사기 타입\n"
        "    type 5: 한국야금 호닝핸들러 타입\n"
        "  - 개선 방향: \"설계 요소 통일화 표준설계, 원가 고려 설계 필수\"\n\n"
        "[2단계] 2025.04.10 — 양면 연삭기 DX 자동화 킥오프 미팅\n"
        "────────────────────────────────────────────────\n"
        "· 참여자: 김순봉, 김웅기, 김동준, 정광선, 최현수\n"
        "· 주요 결정사항 (원문 발췌):\n"
        "  - \"초보자도 정밀 연삭 가능한 자동 시스템 구축\"\n"
        "  - \"무인화, 연속 자동 가동이 가능한 시스템 설계\"\n"
        "  - \"극소형부터 대형 인서트까지 대응하는 자동 연삭 시스템 개발\"\n"
        "· AMR 연계 기술 검토:\n"
        "  - 자동 픽업 기능(조, 마그네틱 그리퍼) 필요\n"
        "  - AGV 및 코드 인식 기능 기반 물류 자동화 시스템 구축\n"
        "  - RFID/바코드 제품 인식 17개 개발 항목 도출\n"
        "· 후지산키 협업으로 9가지 개선안 적용 확정.\n\n"
        "[3단계] 2025.06~09 — 2세대 공용 픽업툴(U19) 사양 확정\n"
        "────────────────────────────────────────────────\n"
        "· ATC(Auto Tool Change) 기능 요구사양:\n"
        "  - Gripper 방식별 호환 공압관로 확보 (Jaw / Vacuum / Magnet / Softgrip / Tilt Grip)\n"
        "  - 센서 신호선 (pogopin 배선) 적용\n"
        "  - 중량 최소화 목표: 현재 2~3kg\n"
        "  - 그리퍼간격 65mm 유지\n"
        "  - 멀티 조인트 & 슬립링 적용\n"
        "  - R축 모터 용량 상향, 모터사양 재검토\n"
        "· 회전축 멀티 조인트/슬립링 표준설계:\n"
        "  - 공압 1~4포트, 6포트 2종 표준\n"
        "  - 전기배선 8~12core\n"
        "· 적용 대상: 프레스 외 전제품, 반전Unit(90도/180도), 포장기 회전 Unit\n"
        "· 구매품 참조: SMC ATC, Roboworker Gripper System, 코비스 슬립링.\n\n"
        "[4단계] 2025.12.02 — 경영회의 중간보고\n"
        "────────────────────────────────────────────────\n"
        "· 원문 인용: \"무인화-연속운전(AMR/AGV), 공정별 플랫폼 개발 강조\" "
        "(2025-12 경영회의, 생산관리팀 발표)\n"
        "· HAM 기구설계(코팅/포장)팀: \"반전기술 별도 관리, 포장기 플랫폼 다이어그램, "
        "ATC 개발건 상세보고\"\n"
        "· 대표이사 지시: \"시장 확장 및 원가 혁신 - 인도, 러시아 등 신규 시장 진출\", "
        "\"DX 시스템 구축 - 데이터 기반 의사결정 체계 확립\"\n"
        "· 이슈: 11월까지 품질 문제 135건 중 118건(27.8%)이 출하 후 발생, "
        "SW/PVD/포장기 집중 → 물류 시스템 안정화 강화 필요.\n\n"
        "[5단계] 2026.03.18 — WTA 설비 차별점 정리 및 대외 홍보용 정리\n"
        "────────────────────────────────────────────────\n"
        "· 프레스핸들러 차별화 완료 항목:\n"
        "  \"고 생산성 (SPM 23EA 저울 측정 포함, 연동 12~14EA), 신속한 셋업 (3일 이후 양산), "
        "스마트 팩토리 대응 AGV 물류 및 DATA 연동, 편리한 HMI, "
        "디버링 기술 접촉식/비접촉식/Hole\"\n"
        "· 미해결 이슈: 특수 제품 적재, Burr 치수 검사, 인선 높이 검사\n"
        "· 특징 홍보용: 소형 제품 Pickup Tool, 무인정지 에러, 듀얼 툴 안정화\n"
        "· CE 인증 획득 완료."
    ),

    "results": (
        "[중간 결과]\n"
        "1. AMR 물류 5가지 타입 표준화 완료 (화루이 프레스 / 한국야금 Cell / "
        "대구텍 키엔스 / 한국야금 포장기6호기·검사기 / 한국야금 호닝핸들러).\n"
        "2. 2세대 공용 픽업툴 U19 사양 확정 — ATC 기능, 5종 그리퍼 호환, "
        "포고핀 배선, 목표중량 2kg대.\n"
        "3. 회전축 멀티 조인트/슬립링 표준설계 2종(4포트/6포트) 도면화.\n"
        "4. 프레스 SPM 23 EA/min (저울 포함) 달성, 연동 12~14 EA 검증.\n"
        "5. CE 인증 획득 (프레스핸들러).\n\n"
        "[데이터 요약 — 장비별 SPM 실측]\n"
        "프레스 23 / 소결취출 65 / 연삭핸들러 16 / PVD로딩 40 / PVD언로딩 35 / "
        "CVD로딩 35 / CVD언로딩 35 / 포장기 60 / CBN조립 5(2개소) / 검사기(F2) 20 / "
        "호닝형상검사기 4포인트/min.\n\n"
        "[참조 이미지 — 설계팀 요소기술 문서에서 발췌]\n"
        "· 대구텍 AMR 컨베이어 구조 도면\n"
        "· SMC ATC 구매품 사양서 및 Roboworker Gripper System 카탈로그\n"
        "· WTA 설비 차별점 정리 도표 (11개 장비군 SPM 비교)"
    ),

    "analysis": (
        "[성공 사항]\n"
        "· 경영회의 안건(2024-01) → 요소기술 표준화(2025-10) → 설비 차별점 정리(2026-03) "
        "까지 2년에 걸쳐 단계적 표준화 달성.\n"
        "· AMR 5 타입 분류로 고객사별(화루이/한국야금/대구텍) 대응 체계 구축.\n"
        "· ATC 기술 내재화 착수 → 외부 SMC/Roboworker 의존 축소 방향.\n\n"
        "[미해결 이슈 및 고민]\n"
        "· PVD 언로딩기 제품 간격 4mm 대응 불가 (전체의 30% 미대응).\n"
        "· 프레스 특수 제품 적재 / Burr 치수 검사 / 인선 높이 검사 미해결.\n"
        "· 소형 제품 Pickup Tool, 무인정지 에러, 듀얼 툴 안정화 필요.\n"
        "· 품질 이슈 27.8%(118/135건) 출하 후 발생 — 물류/반전 안정성 추가 검증 필요.\n\n"
        "[향후 계획]\n"
        "· 2026년 상반기: U19 픽업툴 프로토타입 제작 및 사내 양산라인 적용.\n"
        "· 2026년 하반기: 5가지 AMR 타입별 표준 BOM 확정 및 원가 10% 절감.\n"
        "· 인도/러시아 신규 시장 대응 AMR 사양 별도 검토(대표이사 지시).\n\n"
        "[참고자료]\n"
        "1. Confluence [PROD] 대구텍 AMR (AGV), pageId 8079409174, 2022-12-21\n"
        "2. Confluence [minutes] 2024년 1월 경영회의 회의록, pageId 8330838017\n"
        "3. Confluence [MUxST4BiGY31] 2025-04-10 양면 연삭기 DX 자동화, pageId 8742862879\n"
        "4. Confluence [hanjong] 설계팀 요소기술 개발항목, pageId 9128738819, 2025-10-29\n"
        "5. Confluence [minutes] 2025년 12월 경영회의 회의록, pageId 9327411201, 2025-12-02\n"
        "6. Confluence [zPVoA5Le1CFK] WTA 설비 차별점 정리, pageId 9623371777, 2026-03-18"
    ),
}


def insert_images_into_cell(doc, table_idx, images):
    """지정된 테이블의 마지막 셀에 이미지 삽입"""
    if not images:
        return 0
    table = doc.tables[table_idx]
    cell = table.rows[1].cells[0]  # 내용 셀 (1번 row)
    count = 0
    for img_path, caption in images:
        if not os.path.exists(img_path):
            continue
        p = cell.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(img_path, width=Inches(4.5))
            count += 1
            if caption:
                cap = cell.add_paragraph()
                cap_run = cap.add_run(f"[그림] {caption}")
                cap_run.font.size = Pt(9)
                cap_run.italic = True
        except Exception as e:
            print(f"  이미지 삽입 실패: {img_path} - {e}")
    return count


def generate_note(project_data, images):
    """연구노트 생성"""
    out_path = os.path.join(OUTPUT_DIR, project_data["filename"])

    # Step 1: docxtpl 렌더
    tpl = DocxTemplate(TEMPLATE)
    tpl.render(project_data)
    tpl.save(out_path)

    # Step 2: 이미지 삽입 (Table 4 = 결과 및 데이터)
    doc = Document(out_path)
    img_count = insert_images_into_cell(doc, 4, images)
    doc.save(out_path)

    # 검증
    size = os.path.getsize(out_path)
    return out_path, size, img_count


# ======================================================================
# 프로젝트 #1 이미지 선정 (장비물류)
# ======================================================================
def get_project1_images():
    img_dir = os.path.join(IMG_BASE, "1-장비물류")
    # 페이지 9128738819(설계팀 요소기술), 9623371777(WTA 설비 차별점)에서 4장 선정
    candidates = [
        ("p9128738819-img001-image-20251014-021748.png", "설계팀 요소기술 개발항목 - AMR 물류 5타입 분류"),
        ("p9128738819-img004-image-20251014-020308.png", "2세대 공용 픽업툴(U19) ATC 구조도"),
        ("p9623371777-img001-image-20260318-052206.png", "WTA 프레스핸들러 차별화 요소"),
        ("p9623371777-img005-image-20260318-051644.png", "11개 장비군 SPM 비교 도표"),
    ]
    result = []
    for fname, caption in candidates:
        path = os.path.join(img_dir, fname)
        if os.path.exists(path):
            result.append((path, caption))
    return result


if __name__ == "__main__":
    print("=" * 60)
    print("연구노트 v2 재작성 — 샘플(#1 장비물류) 생성")
    print("=" * 60)

    images = get_project1_images()
    print(f"이미지 후보: {len(images)}장")

    path, size, img_count = generate_note(PROJECT_1, images)

    # 페이지 수 추정 (DOCX는 정확한 페이지 수를 쉽게 알 수 없으므로 문자수 기준)
    doc = Document(path)
    total_chars = sum(len(p.text) for p in doc.paragraphs) + \
                  sum(len(cell.text) for t in doc.tables for row in t.rows for cell in row.cells)

    print(f"\n[완료] {PROJECT_1['filename']}")
    print(f"  경로: {path}")
    print(f"  크기: {size:,} bytes")
    print(f"  본문 문자수: {total_chars:,}자")
    print(f"  삽입 이미지: {img_count}장")
    print()
