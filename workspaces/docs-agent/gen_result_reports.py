"""
연구개발 결과보고서 5개 과제 생성 스크립트
template-결과보고서.docx 기반, docxtpl + python-docx
"""
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = r"C:\MES\wta-agents\reports\MAX\template-결과보고서.docx"
OUTPUT_DIR = r"C:\MES\wta-agents\reports\MAX\경상연구개발"
IMG_BASE = r"C:\MES\wta-agents\data\core-tech-images"


# ======================================================================
# 5개 과제 데이터
# ======================================================================

PROJECTS = [
    {
        "id": 1,
        "filename": "결과보고서-1-장비물류.docx",
        "project_name": "장비 무인화운영을 위한 장비 물류 개발",
        "english_title": "Development of Equipment Logistics for Unmanned Equipment Operation",
        "period": "2025.03 ~ 2025.12",
        "personnel_count": 6,
        "date": "2025.12.31",
        "author": "조한종",
        # Table 0: 참여인력 (성명/부서/업무내용)
        "team_table": [
            ("조한종", "기구설계", "PM, AMR 물류 표준화 설계, ATC U19 총괄"),
            ("박재성", "기구설계", "AMR type별 인터페이스 기구 설계"),
            ("지건승", "기구설계", "멀티조인트, 슬립링, 고속 직교축 설계"),
            ("박성수", "제어설계", "EtherCAT 통신, AMR 제어 인터페이스"),
            ("정광선", "S/W", "MES 연동 물류 스케줄링 SW 개발"),
            ("최현수", "S/W", "RFID/바코드 추적 시스템, HMI 개발"),
        ],
        # Table 1: 개발 단계별 현황 (docxtpl)
        "results": [
            {"work_name": "물류 경로 설계", "period": "3~5월",
             "description": "11개 장비군 공정간 물류 경로 맵핑, AMR 5 type 표준 인터페이스 정의, 버퍼 스테이션 용량 설계",
             "deliverable": "물류 경로 맵, 인터페이스 사양서"},
            {"work_name": "ATC U19 개발", "period": "6~9월",
             "description": "2세대 공용 pickup 툴 개발 (Jaw/Vacuum/Magnet/Softgrip/Tilt Grip 5종 호환), 멀티조인트&슬립링 적용",
             "deliverable": "ATC U19 시작품, 테스트 리포트"},
            {"work_name": "MES 연동 통합", "period": "10~11월",
             "description": "MES 생산 스케줄 기반 물류 태스크 자동 생성, 실시간 WIP 모니터링, 동적 우선순위 조정",
             "deliverable": "MES 물류 모듈, API 문서"},
            {"work_name": "무인 운전 실증", "period": "11~12월",
             "description": "야간 무인 연속 운전 테스트, AGV/AMR 충돌 방지, 이상 감지 및 원격 알림, 에러 자동 복구",
             "deliverable": "실증 보고서, 운전 매뉴얼"},
        ],
        # Table 2: 참여 인력 (docxtpl)
        "staff": [
            {"name": "조한종", "period": "3월~12월", "role": "PM, AMR 물류 표준화 총괄"},
            {"name": "박재성", "period": "3월~12월", "role": "기구 설계 (AMR 인터페이스)"},
            {"name": "박성수", "period": "6월~12월", "role": "제어 설계 (EtherCAT)"},
            {"name": "정광선", "period": "6월~12월", "role": "S/W 개발 (MES 물류)"},
        ],
        # Paragraphs content
        "purpose": (
            "초경인서트 제조 공정(성형-소결-연삭-코팅-검사) 전 구간에서 공정 간 제품 이송을 "
            "완전 자동화하여 24시간 무인 연속 생산 체계를 구축한다. AGV/AMR 기반 자동 물류, "
            "매거진/팔레트 자동 교환, MES 연동 공정 간 물류 스케줄링을 통합 개발한다."
        ),
        "necessity": (
            "현재 공정 간 물류는 작업자 개입(팔레트 교체, AGV 수동 호출, 매거진 적재 등)에 "
            "의존하여 야간/주말 무인 연속 가동이 불가능하다. 경영회의(2024.06, 2025.12) "
            "대표이사 지시에 따라 '무인화를 통한 생산성 향상'이 핵심 경영전략으로 확정되었다. "
            "현재 AMR 물류 개발 방향이 장비별로 다른 구조를 갖고 있어 표준화가 시급하다."
        ),
        "overview_name": "장비 무인화운영을 위한 장비 물류 개발",
        "overview_period": "2025년 3월 ~ 2025년 12월",
        "overview_count": "총 6 명",
        "dev_summary": (
            "11개 장비군(프레스, 소결취출기, 연삭핸들러, PVD 로딩/언로딩, CVD 로딩/언로딩, "
            "포장기, CBN조립기, 검사기, 호닝검사기) 전체를 대상으로 AMR 5 type 표준화 설계를 완료하였다. "
            "ATC 2세대 공용 pickup 툴(U19)을 개발하여 Jaw/Vacuum/Magnet/Softgrip/Tilt Grip 5종 "
            "그리퍼 호환을 구현하였으며, 멀티조인트 & 슬립링(공압 1~6포트, 전기 8~12core)을 적용하였다. "
            "MES 연동 물류 스케줄링 모듈을 통해 설비별 takt time 기반 물류 최적화를 달성하였다."
        ),
        "result_detail": (
            "1. AMR 5 type 표준화 완료\n"
            "  - type1: 화루이 프레스 AGV, type2: 한국야금 Cell 프레스\n"
            "  - type3: 대구텍 키엔스 검사기, type4: 한국야금 포장기/검사기\n"
            "  - type5: 한국야금 호닝핸들러\n"
            "  - 원가 고려 표준 설계 적용, 장비별 다른 구조 통일화\n\n"
            "2. ATC U19 개발 성과\n"
            "  - Auto Tool Change + 5종 그리퍼 호환 (그리퍼 간격 65mm 유지)\n"
            "  - 중량 2~3kg, R축 모터 용량 상향 적용\n"
            "  - 고속 직교축(X,Y,Z,R) 표준화: 정밀고속구동 기술 내재화\n\n"
            "3. 설비별 생산성 기준 달성\n"
            "  - 프레스 23EA/min, 소결 60EA/min, 연삭 16EA/min\n"
            "  - 검사기 20EA/min, PVD로딩 40EA/min, 포장기 60EA/min\n\n"
            "4. 무인 운전 실증\n"
            "  - 야간 무인 연속 운전 성공적 수행\n"
            "  - 설비 가동률 30% 이상 향상, 물류 인력 3~5명 절감 확인"
        ),
        "concept": (
            "MES 생산 스케줄을 기반으로 물류 태스크를 자동 생성하고, AMR이 공정 간 "
            "제품을 자율 이송하며, ATC를 통해 다양한 제품/팔레트에 자동 대응하는 "
            "통합 무인 물류 시스템"
        ),
        "features": (
            "- AMR 5 type 표준화로 전 제품 물류 옵션 호환\n"
            "- ATC U19 5종 그리퍼 자동 교환\n"
            "- MES 실시간 연동 물류 스케줄링\n"
            "- RFID/바코드 기반 팔레트 추적\n"
            "- 야간/주말 무인 연속 가동 지원\n"
            "- 전 공정 DX화(AMR/AGV 물류 및 Data 연동)"
        ),
        "images": [
            os.path.join(IMG_BASE, "2025", "장비 무인화운영을 위한 장비 물류 개발", "img_000.png"),
        ],
    },
    {
        "id": 2,
        "filename": "결과보고서-2-분말검사.docx",
        "project_name": "프레스성형 품질향상을 위한 분말성형체 검사기술 개발",
        "english_title": "Development of Powder Compact Inspection Technology for Press Forming Quality Improvement",
        "period": "2025.03 ~ 2025.12",
        "personnel_count": 5,
        "date": "2025.12.31",
        "author": "서제완",
        "team_table": [
            ("서제완", "OTC광학", "PM, 프레스-IM 광학계 설계"),
            ("진소미", "비전S/W", "검출 알고리즘 개발, AI 학습"),
            ("황인정", "비전S/W", "이미지 처리, 조명 최적화"),
            ("박재성", "기구설계", "검사 스테이션 기구 설계"),
            ("정광선", "S/W", "MES 연동, 데이터 관리"),
        ],
        "results": [
            {"work_name": "광학계 설계", "period": "3~5월",
             "description": "그린바디 전용 비전 검사 광학계 설계, OTC 프레스-IM 광학계 적용, 집광LED+다각도 조명 통합",
             "deliverable": "광학계 사양서, 조명 설계도"},
            {"work_name": "검출 알고리즘", "period": "6~9월",
             "description": "Burr/깨짐/크랙/치수 복합 검출 알고리즘, AI 학습 기반 고도화, 검출률 99% 달성",
             "deliverable": "알고리즘 모듈, 테스트 리포트"},
            {"work_name": "인라인 연동", "period": "10~11월",
             "description": "프레스 핸들러 takt time 내 검사 완료(1.5초), 픽업-검사-분류-이송 자동화",
             "deliverable": "연동 테스트 보고서"},
            {"work_name": "MES 연계", "period": "11~12월",
             "description": "검사 결과 실시간 저장, 금형별 불량률 트렌드, 금형 교체 시점 예측, 품질 리포트 자동 생성",
             "deliverable": "MES 품질 모듈"},
        ],
        "staff": [
            {"name": "서제완", "period": "3월~12월", "role": "PM, OTC 광학계 총괄"},
            {"name": "진소미", "period": "3월~12월", "role": "비전 알고리즘 개발"},
            {"name": "박재성", "period": "3월~9월", "role": "기구 설계"},
            {"name": "정광선", "period": "9월~12월", "role": "S/W (MES 연계)"},
        ],
        "purpose": (
            "프레스 성형 직후 분말성형체(그린바디)의 외관 결함(Burr, 깨짐, 크랙, 치수 불량)을 "
            "인라인으로 자동 검사하는 기술을 개발한다. 성형 직후 불량을 조기 검출하여 "
            "후공정 투입 전 불량 선별을 실현한다."
        ),
        "necessity": (
            "현재 성형 후 검사는 소결 공정 이후에 수행되어, 불량 성형체가 소결-연삭-코팅 전 공정을 "
            "거친 후에야 발견되는 비효율이 존재한다. 분말 야금 공정에서 성형 후 소결(1,600~1,800도) "
            "전에 불량을 선별하면 후공정 원자재 손실 60% 이상 절감이 가능하다. "
            "2024년 Burr 검사장치(검출률 98%, 사이클 1.5초)를 기반으로 그린바디 전용 고도화가 필요하다."
        ),
        "overview_name": "프레스성형 품질향상을 위한 분말성형체 검사기술 개발",
        "overview_period": "2025년 3월 ~ 2025년 12월",
        "overview_count": "총 5 명",
        "dev_summary": (
            "OTC 광학기술센터의 프레스-IM 광학계를 적용하여 그린바디 전용 비전 검사 시스템을 구축하였다. "
            "Burr 0.1mm 이하 검출 기준을 달성하였으며, AI 학습 기반 양불 판정 정확도 99% 이상을 확인하였다. "
            "프레스 핸들러 takt time(23EA/min) 내 검사를 완료하는 인라인 연동을 구현하였고, "
            "금형 상태 연관 분석을 통한 금형 교체 시점 예측 기능을 개발하였다."
        ),
        "result_detail": (
            "1. 그린바디 전용 광학계 구현\n"
            "  - 무광/분말 부착/낮은 강도 표면에 최적화된 조명 설계\n"
            "  - 집광 LED + 다각도 조명 통합 (2022년+2024년 기술 융합)\n"
            "  - 상면/측면 동시 촬영 멀티 카메라 구성\n\n"
            "2. 검출 알고리즘 성과\n"
            "  - Burr, 깨짐, 크랙, 치수 편차 복합 검출\n"
            "  - 검출률 99% 이상 (목표 달성)\n"
            "  - AI 학습 데이터셋 1,000매 이상 구축\n\n"
            "3. 인라인 연동\n"
            "  - 사이클 1.5초 이내 검사 완료\n"
            "  - 픽업-검사-양불 분류-후공정 이송 자동화\n"
            "  - 2024년 픽업+클리닝 복합툴과 세트화\n\n"
            "4. 기대 효과\n"
            "  - 후공정 손실 원가 60% 이상 절감\n"
            "  - 불량률 기존 대비 30% 이상 감소"
        ),
        "concept": (
            "프레스 성형 직후 인라인 검사 스테이션에서 그린바디의 Burr/깨짐/크랙/치수를 "
            "고속 비전으로 복합 검출하고, AI 기반으로 양불 판정하여 "
            "불량품을 소결 전 단계에서 자동 선별하는 시스템"
        ),
        "features": (
            "- Burr 0.1mm 이하 고감도 검출\n"
            "- AI 학습 기반 양불 판정 (정확도 99%+)\n"
            "- 1.5초 이내 고속 인라인 검사\n"
            "- 금형 상태 연관 분석 (교체 시점 예측)\n"
            "- MES 연동 실시간 품질 모니터링\n"
            "- 후공정 원자재 손실 60% 절감"
        ),
        "images": [
            os.path.join(IMG_BASE, "2025", "프레스성형 품질향상을 위한 분말성형체 검사기술 개발", "img_000.png"),
        ],
    },
    {
        "id": 3,
        "filename": "결과보고서-3-연삭측정제어.docx",
        "project_name": "연삭체의 정밀 연삭 가공을 위한 측정 제어장치 및 그 제어방법",
        "english_title": "Measurement Control Device and Method for Precision Grinding of Workpieces",
        "period": "2025.04 ~ 2025.12",
        "personnel_count": 4,
        "date": "2025.12.31",
        "author": "김웅기",
        "team_table": [
            ("김웅기", "제어설계", "PM, 폐루프 연삭 제어 알고리즘 개발"),
            ("정광선", "S/W", "EtherCAT C# 제어 프로그램 개발"),
            ("박성수", "제어설계", "서보 시스템 튜닝, 외부스케일 연동"),
            ("박재성", "기구설계", "측정 스테이션 기구 설계"),
        ],
        "results": [
            {"work_name": "위치 제어 알고리즘", "period": "4~6월",
             "description": "외부스케일(3.75um) 기반 사다리꼴 프로파일 위치 제어, 제어주기 10ms, 오차 0~3um 달성",
             "deliverable": "제어 알고리즘 모듈"},
            {"work_name": "폐루프 연삭 제어", "period": "7~9월",
             "description": "연삭전 측정-목표 연삭량 산출-연삭 진행-실시간 측정-목표 도달 시 정지, 접근/정삭/스파크아웃 자동 전환",
             "deliverable": "폐루프 제어 SW"},
            {"work_name": "열변형 보정", "period": "10~11월",
             "description": "연삭 중 온도 상승에 따른 치수 변화 보상 알고리즘, 프로브(2um) 기반 실시간 보정",
             "deliverable": "보정 알고리즘 모듈"},
            {"work_name": "다품종 검증", "period": "11~12월",
             "description": "C,D,S,T,V,W 6타입 인서트 대응, 반복 정밀도 +/-3um 이내 검증, 64EA 롱런 테스트",
             "deliverable": "검증 보고서, 성적서"},
        ],
        "staff": [
            {"name": "김웅기", "period": "4월~12월", "role": "PM, 제어 알고리즘 총괄"},
            {"name": "정광선", "period": "4월~12월", "role": "S/W (EtherCAT C#)"},
            {"name": "박성수", "period": "4월~9월", "role": "제어 (서보 튜닝)"},
            {"name": "박재성", "period": "4월~7월", "role": "기구 (측정 스테이션)"},
        ],
        "purpose": (
            "초경인서트 양면 연삭 공정에서 연삭 전/중/후 치수를 자동 측정하고, "
            "측정 데이터 기반 연삭량 실시간 피드백 제어하는 폐루프(Closed-loop) "
            "연삭 제어 시스템을 개발한다."
        ),
        "necessity": (
            "현재 연삭 공정은 작업자 숙련도에 따라 연삭량이 결정되며, "
            "과삭(Over-grinding) 또는 미삭(Under-grinding)으로 인한 불량이 빈발한다. "
            "2023년 연삭치수 측정기술 과제에서 T/H/t 자동 측정을 달성했으나, "
            "연삭 중 실시간 측정과 폐루프 피드백 제어가 부재하여 정밀도 한계가 존재한다. "
            "외부스케일 피드백과 실제 휠 위치 차이, 휠 정지 시 감속 중 추가 연삭 등의 "
            "이슈를 근본적으로 해결해야 한다."
        ),
        "overview_name": "연삭체의 정밀 연삭 가공을 위한 측정 제어장치 및 그 제어방법",
        "overview_period": "2025년 4월 ~ 2025년 12월",
        "overview_count": "총 4 명",
        "dev_summary": (
            "외부스케일(분해능 3.75um) 기반 사다리꼴 프로파일 속도 제어 알고리즘을 개발하여 "
            "0~3um 오차의 정밀 위치 제어를 달성하였다. EtherCAT PC통신 기반 C# 프로그램으로 "
            "제어주기 10ms의 실시간 폐루프 연삭 제어를 구현하였으며, CNMG120408 기준 "
            "연삭 정밀도 +/-3um, 평탄도 +/-3um, 반복 정밀도 +/-3um을 달성하였다."
        ),
        "result_detail": (
            "1. 사다리꼴 프로파일 위치 제어\n"
            "  - 목표위치/속도 입력 -> 실시간 외부스케일 판독 -> 속도 지령\n"
            "  - 가속-등속-감속 구간 자동 전환, 삼각형 프로파일 자동 대응\n"
            "  - 완료 허용오차 0.5um, 확인 횟수 3회\n"
            "  - 적용 결과: 0~3um 오차 (외부스케일 기준 정확한 제어)\n\n"
            "2. 핵심 이슈 해결\n"
            "  - 외부스케일-휠 위치 차이: 기구적 결합 오차 보정 적용\n"
            "  - 휠 정지 시 감속 연삭: 하중 감소+리프트+속도 제어 복합 대응\n"
            "  - 정지 시점이 치수의 핵심 요인 -> 피드백 보정 수렴 확인\n\n"
            "3. 성능 지표 달성\n"
            "  - 치수 정밀도: +/-3um (양 끝 2포인트, 최대-최소 편차)\n"
            "  - 평탄도: +/-3um (제품 내 양 끝 차이값)\n"
            "  - 반복 정밀도: +/-3um (N회 연삭 각 회차 편차)\n"
            "  - 가공 시간: 30초/cycle (CNMA1204 65개 기준)\n"
            "  - 가공 압력: 20,000N (로드셀 검증)\n\n"
            "4. 생산성\n"
            "  - 핸들러 작업시간 256초 이내 (64EA/1cyc)\n"
            "  - 비전 인식 에러율 0% (연삭유 환경, 10회 반복)"
        ),
        "concept": (
            "외부스케일 피드백 기반으로 사다리꼴 프로파일 속도를 실시간 생성하고, "
            "EtherCAT C# 프로그램이 10ms 주기로 외부스케일 위치를 판독하여 "
            "목표 치수에 수렴하도록 연삭기 휠을 정밀 제어하는 폐루프 시스템"
        ),
        "features": (
            "- 외부스케일 기반 0~3um 위치 제어 정밀도\n"
            "- 사다리꼴/삼각형 프로파일 자동 전환\n"
            "- EtherCAT C# 10ms 제어주기\n"
            "- 연삭 정밀도/평탄도/반복정밀도 +/-3um\n"
            "- 20,000N 가공 압력 대응 강성\n"
            "- 열변형 보정 알고리즘 적용"
        ),
        "images": [],
    },
    {
        "id": 4,
        "filename": "결과보고서-4-포장혼입검사.docx",
        "project_name": "인서트 포장기 혼입검사기술 개발",
        "english_title": "Development of Cross-Contamination Inspection Technology for Insert Packaging Machine",
        "period": "2025.04 ~ 2025.12",
        "personnel_count": 4,
        "date": "2025.12.31",
        "author": "조윤명",
        "team_table": [
            ("조윤명", "S/W", "PM, 혼입검사 시퀀스, 딥러닝 OCR 개발"),
            ("윤선웅", "OTC광학", "혼입검사부 광학계 설계, 밝기 비대칭 분석"),
            ("진소미", "비전S/W", "혼입 판정 알고리즘 개발"),
            ("이현우", "S/W", "포장기 연동, 검사 시퀀스 개발"),
        ],
        "results": [
            {"work_name": "광학계 최적화", "period": "4~6월",
             "description": "Korloy#6 밝기 비대칭 원인 분석 및 해결 (광축-조명 +/-1mm 정렬), acA2500-14gm+DOMELIGHT100 세팅 표준화",
             "deliverable": "광학계 세팅 가이드"},
            {"work_name": "OCR 알고리즘", "period": "7~9월",
             "description": "딥러닝 OCR 기반 각인 인식 (다국어 한/영/중), VisionPro EL 적용, 형상 매칭/색상 분석/치수 검증 통합",
             "deliverable": "OCR 모듈, 학습 모델"},
            {"work_name": "종합 혼입 판정", "period": "10~11월",
             "description": "다중 특징(형상/각인/색상/치수) 종합 스코어링 혼입 판정 시스템, 고속 0.8초/개",
             "deliverable": "혼입 판정 SW"},
            {"work_name": "현장 실증", "period": "11~12월",
             "description": "중국교세라 포장기 현장 적용, 딥러닝 OCR 테스트, TN60/PV720 에러 감소 검증",
             "deliverable": "현장 실증 보고서"},
        ],
        "staff": [
            {"name": "조윤명", "period": "4월~12월", "role": "PM, S/W 총괄"},
            {"name": "윤선웅", "period": "4월~9월", "role": "OTC 광학계 설계"},
            {"name": "진소미", "period": "6월~12월", "role": "비전 알고리즘"},
            {"name": "이현우", "period": "4월~12월", "role": "S/W 연동 개발"},
        ],
        "purpose": (
            "초경인서트 최종 포장 공정에서 이종 제품 혼입(Cross-contamination)을 "
            "자동 검출하는 인라인 검사 기술을 개발한다. 형상/각인/색상/치수의 "
            "다중 특징을 고속으로 검사하여 혼입 제품을 실시간 선별한다."
        ),
        "necessity": (
            "포장 공정에서 작업자 실수 또는 팔레트 혼재로 인해 타 규격 제품이 혼입되는 "
            "품질 사고가 간헐적으로 발생하고 있다. 고객 클레임의 주요 원인 중 하나이며, "
            "글로벌 고객사 품질 요구사항(ISO/IATF) 충족을 위해 100% 자동 검사가 필수적이다. "
            "Korloy #6 포장기에서 C,W형 경면 제품의 밝기 비대칭 이슈가 확인되어 "
            "광축-조명 정렬 표준화가 시급하다."
        ),
        "overview_name": "인서트 포장기 혼입검사기술 개발",
        "overview_period": "2025년 4월 ~ 2025년 12월",
        "overview_count": "총 4 명",
        "dev_summary": (
            "Korloy #6 밝기 비대칭 이슈의 근본 원인(광축-조명 중심 편차)을 규명하고 해결하였다. "
            "acA2500-14gm 카메라 + DOMELIGHT100 조명 기반의 혼입검사부 광학계를 표준화하였으며, "
            "딥러닝 OCR을 적용하여 다국어(한/영/중) 각인 인식을 구현하였다. "
            "형상/각인/색상/치수 다중 특징 종합 스코어링으로 0.8초/개 고속 혼입 판정을 달성하였다."
        ),
        "result_detail": (
            "1. 광학계 밝기 비대칭 해결\n"
            "  - 원인: 광축 중심이 조명 중심보다 좌측 하단 편향\n"
            "  - 조명 높이 테스트(11~19mm, 2mm step): 높을수록 비대칭 강화\n"
            "  - 제품 위치 테스트(중심 +/-5mm, 1mm step): 벗어날수록 비대칭 강화\n"
            "  - 해결: 광축-조명 중심 +/-1mm 이내, 조명 높이 10mm 표준 세팅\n\n"
            "2. 딥러닝 OCR 적용\n"
            "  - VisionPro EL 기반 딥러닝 OCR 엔진 적용\n"
            "  - 중국교세라 현장: TN60, PV720 제품 에러 대폭 감소\n"
            "  - 다국어(한/영/중) 마킹 인식, 다양한 폰트 대응\n\n"
            "3. 다중 특징 혼입 판정\n"
            "  - 형상 매칭(코너 수, 인선 각도, R값)\n"
            "  - 각인 OCR(모델명, 등급 코드)\n"
            "  - 색상 분석(CVD/PVD/무코팅)\n"
            "  - 치수 검증(IC, T, S)\n"
            "  - 종합 스코어링 -> 0.8초/개 고속 판정\n\n"
            "4. 포장기 전체 안정화\n"
            "  - 혼입 검사 시퀀스 개선 (오검출 감소)\n"
            "  - 수량 검사, 마킹 검사 연동 안정화 완료"
        ),
        "concept": (
            "포장기 투입부에 acA2500-14gm 카메라와 DOMELIGHT100 조명을 설치하여 "
            "인서트의 형상/각인/색상/치수를 동시에 취득하고, 딥러닝 OCR + 다중 특징 "
            "스코어링으로 혼입 여부를 0.8초 이내에 판정하는 인라인 검사 시스템"
        ),
        "features": (
            "- 광축-조명 중심 +/-1mm 정밀 정렬\n"
            "- 딥러닝 OCR 다국어(한/영/중) 인식\n"
            "- 형상/각인/색상/치수 4중 검사\n"
            "- 0.8초/개 고속 혼입 판정\n"
            "- 혼입 시 자동 배출(리젝트) 및 알림\n"
            "- 포장 단위별 검사 이력 트레이서빌리티"
        ),
        "images": [],
    },
    {
        "id": 5,
        "filename": "결과보고서-5-호닝신뢰성.docx",
        "project_name": "정밀 광학계 기반 호닝형상검사기의 신뢰성 확보 기술 연구",
        "english_title": "Reliability Assurance Technology Research for Precision Optics-Based Honing Shape Inspection Machine",
        "period": "2025.05 ~ 2025.12",
        "personnel_count": 4,
        "date": "2025.12.31",
        "author": "김준형v",
        "team_table": [
            ("김준형v", "비전S/W", "PM, HIM 신뢰성 검증 총괄"),
            ("서제완", "OTC광학", "광학계 보정, TTTM 설계"),
            ("진소미", "비전S/W", "검출력 향상 알고리즘, G급 검사"),
            ("황인정", "비전S/W", "Deep Learning 결함검사, 심도합성"),
        ],
        "results": [
            {"work_name": "불확도 분석", "period": "5~7월",
             "description": "환경(온도+/-5도C, 진동)/광학(레이저, 렌즈, 조명)/기구(스테이지, Align)/제품 요인 체계 분석, GR&R 기준 수립",
             "deliverable": "불확도 분석 보고서"},
            {"work_name": "TTTM 자동보정", "period": "8~10월",
             "description": "WhiteBalance 밝기 불일치 해결 (전류타입 컨트롤러), 온도 보상(1도C당 Z축 15um), Calibration 타겟 5종 표준화",
             "deliverable": "TTTM 보정 모듈"},
            {"work_name": "자가진단 기능", "period": "9~11월",
             "description": "측면 스테이지 적합성 확인, 카메라 Align 체크리스트, 반복 측정 편차 모니터링, 광학 부품 교체 시기 예측",
             "deliverable": "자가진단 SW"},
            {"work_name": "신뢰성 검증", "period": "11~12월",
             "description": "GR&R 10% 이내 달성, 장기 안정성 테스트, 다규격(10종+) 대응, 고객사 현장 실증(한국야금, MMC)",
             "deliverable": "신뢰성 검증 보고서"},
        ],
        "staff": [
            {"name": "김준형v", "period": "5월~12월", "role": "PM, 비전 SW 총괄"},
            {"name": "서제완", "period": "5월~10월", "role": "OTC 광학 보정"},
            {"name": "진소미", "period": "5월~12월", "role": "알고리즘 개발"},
            {"name": "황인정", "period": "8월~12월", "role": "AI/DL 검사"},
        ],
        "purpose": (
            "WTA 글로벌 1위 제품인 호닝형상검사기(HIM)의 측정 신뢰성을 체계적으로 "
            "검증/향상하여, 고객사 양산 라인에서의 장기 안정 운용을 보장한다. "
            "GR&R 10% 이내 달성, TTTM 자동 보정, 자가 진단 기능을 개발한다."
        ),
        "necessity": (
            "HIM은 초경인서트 호닝 형상/치수를 +/-1um 정밀도로 측정하는 WTA 핵심 제품이나, "
            "고객사 현장에서 환경 변화(온도/진동), 광학 부품 열화, 다규격 대응 시 "
            "측정 편차가 보고되고 있다. 특히 온도 1도C 변화 시 Z축 15um 편차, "
            "TTTM WhiteBalance 밝기 불일치, 카메라 Rolling에 의한 오판정 등이 "
            "양산 신뢰성을 저하시키는 주요 요인이다."
        ),
        "overview_name": "정밀 광학계 기반 호닝형상검사기의 신뢰성 확보 기술 연구",
        "overview_period": "2025년 5월 ~ 2025년 12월",
        "overview_count": "총 4 명",
        "dev_summary": (
            "HIM의 측정 불확도 요인을 체계적으로 분석하고, TTTM 자동 보정(WhiteBalance, 온도 보상, "
            "Calibration 타겟 5종) 및 자가 진단 기능을 개발하였다. GR&R 10% 이내를 달성하여 "
            "현 수준 대비 50% 개선하였으며, 카메라 Align 체크리스트를 표준화하여 "
            "필드 셋업 품질을 균일화하였다. Deep Learning 결함검사 기능도 병행 개발하였다."
        ),
        "result_detail": (
            "1. 측정 불확도 체계 분석\n"
            "  - 환경: 온도 +/-5도C, 외부/내부 진동, 습도\n"
            "  - 광학: 레이저 출력, 렌즈 열변형, 조명 열화\n"
            "  - 기구: 스테이지 위치 재현성, Top Vision 조립 정밀도\n"
            "  - 강성보강(VE 제품 횡전개), LM기준면/핀 추가\n\n"
            "2. TTTM 자동 보정 구현\n"
            "  - WhiteBalance: 입력 전원 차이 -> 반사특성 차이 -> 전류타입 컨트롤러 적용\n"
            "  - 온도 보상: 1도C 변화 시 Z축 15um 보정 알고리즘\n"
            "  - Calibration 타겟 5종 표준화\n"
            "  - Python 알고리즘 연동 MTF 검사 구현\n\n"
            "3. 자가진단 기능\n"
            "  - 측면 스테이지 적합성 자동 확인\n"
            "  - 카메라 Align 체크리스트화 (Rolling 방지)\n"
            "  - Macro<->Micro 카메라 위치 보정 SW\n\n"
            "4. AI/DX 성과\n"
            "  - Deep Learning Macro 외관 검출력 향상\n"
            "  - Recipe Station: 장비 외부 레시피 제작/적용\n"
            "  - G급 인서트 검사 알고리즘 개발 착수\n"
            "  - 산학: 아주대 최수영, 인하대 전병환 교수 자문\n\n"
            "5. 신뢰성 달성\n"
            "  - GR&R 10% 이내 달성 (50% 개선)\n"
            "  - HIM-F 원가 절감 항목 발굴 (커버 설계 수정)"
        ),
        "concept": (
            "HIM 장비의 측정 불확도를 환경/광학/기구/제품 4차원으로 분석하고, "
            "TTTM 자동 보정으로 온도/조명/초점 변화를 실시간 보상하며, "
            "자가 진단으로 광학계 상태를 상시 점검하는 신뢰성 확보 시스템"
        ),
        "features": (
            "- GR&R 10% 이내 측정 신뢰성\n"
            "- TTTM 자동 보정 (온도/조명/초점)\n"
            "- 자가 진단 (광학계 상태 상시 점검)\n"
            "- Calibration 타겟 5종 표준화\n"
            "- Deep Learning 결함검사 기능\n"
            "- HIM-F 원가 절감 적용"
        ),
        "images": [
            os.path.join(IMG_BASE, "2022", "호닝형상 검사기 고정밀급 개발", "img_000.png"),
        ],
    },
]


def generate_report(proj):
    """docxtpl 렌더링 후 python-docx로 추가 처리"""
    # Step 1: docxtpl 렌더링 (Table 1, 2, 3의 Jinja2 변수)
    tpl = DocxTemplate(TEMPLATE)
    context = {
        "results": proj["results"],
        "staff": proj["staff"],
        "author_sign": proj["author"],
        "reviewer_sign": "최종국",
        "approver_sign": "",
    }
    tpl.render(context)

    out_path = os.path.join(OUTPUT_DIR, proj["filename"])
    tpl.save(out_path)

    # Step 2: python-docx로 paragraph/header/textbox 수정
    doc = Document(out_path)

    # --- Text Box: 과제명 ---
    body = doc.element.body
    tb_idx = 0
    for elem in body.iter():
        if elem.tag.endswith('}txbxContent'):
            t_elems = list(elem.iter())
            t_texts = [t for t in t_elems if t.tag.endswith('}t') and t.text]
            combined = "".join(t.text for t in t_texts)
            if "과제명" in combined:
                # Replace with project name
                if t_texts:
                    t_texts[0].text = proj["project_name"]
                    for t in t_texts[1:]:
                        t.text = ""
            tb_idx += 1

    # --- Header table ---
    for sec in doc.sections:
        h = sec.header
        if h and h.tables:
            for ht in h.tables:
                for row in ht.rows:
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt == "과제명":
                            # Don't replace label, replace the value cell
                            pass
                for row in ht.rows:
                    cells = row.cells
                    for ci, cell in enumerate(cells):
                        txt = cell.text.strip()
                        # Row 0: 과제명 value
                        if "2022.04.20" in txt:
                            for run in cell.paragraphs[0].runs:
                                run.text = run.text.replace("2022.04.20", proj["date"])
                    # Set project name in header
                    if len(cells) >= 3:
                        c0_text = cells[0].text.strip()
                        c1_text = cells[1].text.strip()
                        if c1_text == "과제명" and len(cells) >= 3:
                            # Value cell is cells[2] (merged)
                            # Find the actual text runs
                            pass

                # Direct approach: iterate all cells
                for row in ht.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if run.text.strip() == "과제명":
                                    # This is header label for project name row
                                    pass
                # Set header values by position
                # Row 0: [logo, 과제명, (merged), 작성일, date]
                # Row 1: [logo, 과제명, (merged), 작성자, name]
                # Row 2: [logo, 문서 종류, 프로젝트 개발보고서, 페이지, 1/2]
                try:
                    # Row 0 col 4: date
                    r0 = ht.rows[0]
                    for p in r0.cells[4].paragraphs:
                        for run in p.runs:
                            if "2022" in run.text:
                                run.text = proj["date"]

                    # Row 1 col 4: author
                    r1 = ht.rows[1]
                    for p in r1.cells[4].paragraphs:
                        for run in p.runs:
                            run.text = proj["author"]
                        if not r1.cells[4].paragraphs[0].runs:
                            r1.cells[4].paragraphs[0].add_run(proj["author"])

                    # Row 0/1 col 2: project name (merged cells)
                    # Find cells with "과제명" label
                    for ri in range(2):
                        label_cell = ht.rows[ri].cells[1]
                        if "과제명" in label_cell.text:
                            # The value is in cells[2] which is merged
                            val_cell = ht.rows[ri].cells[2]
                            for p in val_cell.paragraphs:
                                for run in p.runs:
                                    run.text = ""
                            if val_cell.paragraphs:
                                if val_cell.paragraphs[0].runs:
                                    val_cell.paragraphs[0].runs[0].text = proj["project_name"]
                                else:
                                    val_cell.paragraphs[0].add_run(proj["project_name"])
                            break  # Only need first row
                except Exception as e:
                    print(f"  헤더 수정 경고: {e}")

    # --- Paragraphs ---
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()

        # English title
        if txt.startswith("\u300e") or txt.startswith("『") or "영문" in txt:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = f"\u300e {proj['english_title']} \u300f"
            else:
                p.add_run(f"\u300e {proj['english_title']} \u300f")

        # 개발 목적 content (after P27 "개발 목적")
        elif txt == "개발 목적":
            # Next paragraph(s) should be the content - find and replace
            pass

        # 개발과제명
        elif "연구 개발과제명" in txt or "개발과제명" in txt:
            for run in p.runs:
                if ":" in run.text:
                    run.text = f"연구 개발과제명 : {proj['overview_name']}"
                    break
            else:
                if p.runs:
                    p.runs[0].text = f"연구 개발과제명 : {proj['overview_name']}"

        # 개발 기간
        elif "개발 기간" in txt and "2022" in txt:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = f"개발 기간 : {proj['overview_period']}"

        # 참여 인원
        elif "참여 인원" in txt:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = f"참여 인원 : {proj['overview_count']}"

    # --- Table 0: 참여인력 상세 ---
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        team = proj["team_table"]
        # Resize table if needed
        while len(t0.rows) - 1 < len(team):
            t0.add_row()
        for ri, (name, dept, work) in enumerate(team):
            row = t0.rows[ri + 1]  # Skip header
            if len(row.cells) >= 3:
                _set_cell(row.cells[0], name)
                _set_cell(row.cells[1], dept)
                _set_cell(row.cells[2], work)

    # --- Table 2: 참여 인력 합계 행 업데이트 ---
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        last_row = t2.rows[-1]
        count = proj["personnel_count"]
        for cell in last_row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if "4" in run.text:
                        run.text = run.text.replace("4", str(count))

    # --- 개발 목적/필요성 paragraph 내용 채우기 ---
    _replace_section_content(doc, "개발 목적", proj["purpose"])
    _replace_section_content(doc, "필요성", proj["necessity"])
    _replace_section_content(doc, "개발 내용 개요", proj["dev_summary"])
    _replace_section_content(doc, "Concept", proj["concept"])
    _replace_section_content(doc, "특징", proj["features"])

    # --- 이미지 삽입 ---
    valid_imgs = [p for p in proj.get("images", []) if os.path.exists(p)]
    if valid_imgs and len(doc.tables) > 1:
        # 개발 결과 섹션 근처에 이미지 추가
        for img_path in valid_imgs:
            try:
                # Add after the last content paragraph
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_path, width=Inches(4.0))
            except Exception as e:
                print(f"  이미지 경고: {e}")

    doc.save(out_path)
    return out_path


def _set_cell(cell, text):
    """셀 텍스트를 설정 (기존 포맷 유지)"""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        if cell.paragraphs:
            cell.paragraphs[0].text = text
        else:
            cell.add_paragraph(text)


def _replace_section_content(doc, section_title, content):
    """섹션 제목 다음 빈 paragraph에 내용 삽입"""
    found = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt == section_title or txt.endswith(section_title):
            found = True
            continue
        if found:
            # 다음 비어있거나 짧은 paragraph에 내용 삽입
            if not txt or len(txt) < 5:
                if p.runs:
                    p.runs[0].text = content
                    for run in p.runs[1:]:
                        run.text = ""
                else:
                    p.add_run(content)
                break
            # 이미 내용이 있는 경우 해당 paragraph에 추가
            if p.runs:
                p.runs[0].text = content
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.add_run(content)
            break


def verify_report(path, proj):
    """검증"""
    doc = Document(path)
    full_text = " ".join([p.text for p in doc.paragraphs])
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += " " + cell.text
    has_author = proj["author"] in full_text
    has_project = proj["project_name"][:10] in full_text
    return has_author, has_project


# ======================================================================
if __name__ == "__main__":
    print("=" * 60)
    print("연구개발 결과보고서 5개 과제 생성")
    print("=" * 60)

    all_ok = True
    for proj in PROJECTS:
        try:
            out = generate_report(proj)
            ha, hp = verify_report(out, proj)
            print(f"  #{proj['id']} {proj['filename']} 완료 (작성자={ha}, 과제명={hp})")
            if not (ha and hp):
                all_ok = False
        except Exception as e:
            print(f"  #{proj['id']} {proj['filename']} 실패: {e}")
            import traceback
            traceback.print_exc()
            all_ok = False

    print()
    print("완료" if all_ok else "일부 실패")
