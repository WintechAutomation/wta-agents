"""연구개발 계획서 5건 Word 문서 생성 스크립트
MD 원본 → template-계획서.docx 매핑 → 개별 docx 출력
"""
import copy
import re
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt

TEMPLATE = Path(r"C:\MES\wta-agents\reports\MAX\template-계획서.docx")
OUT_DIR = Path(r"C:\MES\wta-agents\reports\MAX")

# ── 과제별 데이터 정의 ──────────────────────────────────────

PROJECTS = [
    {
        "id": 1,
        "slug": "장비물류",
        "md": "연구개발-1-장비물류.md",
        "title": "장비 무인화운영을 위한 장비 물류 개발",
        "eng_title": "Development of Equipment Logistics for Unmanned Operation",
        "period": "2025년 3월 1일 ~ 2025년 12월 31일",
        "period_months": "10개월",
        "members": "총 5명 (기구설계 2, 제어설계 1, S/W 2)",
        "method": "AGV/AMR 기반 자동 물류 시스템 설계, MES 연동 스케줄링 개발, 현장 실증",
        "purposes": [
            "초경인서트 제조 공정(성형→소결→연삭→코팅→검사) 전 구간의 공정 간 제품 이송을 완전 자동화하여 장비 무인화 운영을 실현한다.",
            "AGV/AMR 기반 자동 물류, 매거진/팔레트 자동 교환, MES 연동 공정 간 물류 스케줄링을 통합 개발한다.",
            "야간/주말 24시간 무인 연속 생산 체계를 구축하여 설비 가동률 30% 이상 향상을 달성한다.",
        ],
        "necessities": [
            "현재 공정 간 물류는 작업자 개입(팔레트 교체, AGV 수동 호출, 매거진 적재)에 의존하여 야간/주말 무인 연속 가동이 불가능하다. 2022년 AGV 과제는 성형→소결 단일 구간만 대상이었으며, 전체 공정을 관통하는 통합 물류 시스템은 부재한 상태이다.",
            "MES 연동 스케줄링, AMR 다중 경로 관리, 공정별 버퍼 스테이션 자동 관리 등 시스템 수준의 물류 기술 개발이 필요하며, 이를 통해 인력 절감(3~5명), 리드타임 20% 단축, 운반 중 제품 손상/혼입 방지 등의 효과를 기대할 수 있다.",
        ],
        "strategies": [
            "1단계(3~5월): 전 공정 물류 경로 맵핑 및 버퍼 스테이션 용량 설계. 2022년 AGV 과제의 Fork형 AGV, 엘리베이터 유닛 기술을 기반으로 다구간 확장 설계를 수행한다. 설계팀 AMR 물류 4type 표준화 결과 및 대구텍 AMR 설계 사례를 반영한다.",
            "2단계(6~9월): 자동 팔레트/매거진 교환 시스템 개발 및 MES 연동 물류 스케줄링 구현. 2024년 리팔레팅 기술을 적용하여 이종 팔레트 간 자동 전환을 구현한다. 경영회의 무인화 지시(2024.1, 2025.12) 목표에 부합하는 로드맵을 준수한다.",
            "3단계(10~12월): 무인 운전 안전 시스템 통합, AGV/AMR 동선 최적화, 현장 실증 및 24시간 연속 가동 테스트를 수행한다.",
        ],
        "tasks": [
            {"name": "공정 간 물류 경로 설계", "period": "3개월", "start_date": "2025-03-01", "end_date": "2025-05-31", "resource": "기구설계 2명"},
            {"name": "팔레트/매거진 자동교환", "period": "4개월", "start_date": "2025-06-01", "end_date": "2025-09-30", "resource": "기구설계 2 + 제어 1"},
            {"name": "MES 연동 물류 스케줄링", "period": "4개월", "start_date": "2025-06-01", "end_date": "2025-09-30", "resource": "S/W 2명"},
            {"name": "무인 운전 통합 실증", "period": "3개월", "start_date": "2025-10-01", "end_date": "2025-12-31", "resource": "전원 (5명)"},
        ],
        "pm_name": "○○○", "pm_period": "10", "pm_role": "과제 총괄, 물류 설계",
        "sw_name": "○○○", "sw_period": "10", "sw_role": "MES 연동 S/W 개발",
        "budget_detail": "재료비 / 외주가공비 / 시험비", "budget_amount": "별도 책정",
        "constraints": "AGV/AMR 도입 시 기존 공장 레이아웃 변경 최소화 필요, 충돌 방지 및 비상정지 안전 인증 확보",
    },
    {
        "id": 2,
        "slug": "분말검사",
        "md": "연구개발-2-분말검사.md",
        "title": "프레스성형 품질향상을 위한 분말성형체 검사기술 개발",
        "eng_title": "Development of Green Body Inspection Technology for Press Forming Quality Improvement",
        "period": "2025년 3월 1일 ~ 2025년 12월 31일",
        "period_months": "10개월",
        "members": "총 5명 (기구설계 2, 비전 S/W 2, S/W 1)",
        "method": "비전 검사 시스템 설계, AI 기반 불량 검출 알고리즘 개발, 프레스 핸들러 인라인 연동",
        "purposes": [
            "프레스 성형 직후 분말성형체(그린바디)의 외관 결함(Burr, 깨짐, 크랙, 치수 불량)을 인라인으로 자동 검사하는 기술을 개발한다.",
            "성형 직후 단계에서 불량을 조기 검출하여 후공정 투입 전 불량 선별을 실현하고, 원자재 손실을 절감한다.",
            "AI 학습 기반 검사를 고도화하여 양불 판정 정확도 99% 이상을 달성하고, 금형 상태 연관 분석으로 금형 교체 시점을 예측한다.",
        ],
        "necessities": [
            "현재 성형 후 검사는 소결 공정 이후에 수행되어, 불량 성형체가 소결·연삭·코팅 전 공정을 거친 후에야 발견되는 비효율이 존재한다. 소결 전 불량 선별 시 후공정 손실 원가 60% 이상 절감이 가능하다.",
            "2024년 Burr 검사 장치는 소결 후 제품 대상이며, 분말성형체(그린바디)는 강도가 매우 낮고 표면 특성이 다르다. 비접촉 검사 방식 고도화, 그린바디 전용 조명/알고리즘, 미세 크랙 검출 기술이 추가로 필요하다.",
        ],
        "strategies": [
            "1단계(3~5월): 그린바디 전용 비전 검사 시스템 설계. OTC 프레스-IM 광학계 설계/셋업 완료 성과를 기반으로, 분말성형체 표면 특성(무광, 분말 부착, 낮은 강도)에 최적화된 조명 및 멀티 카메라 구성을 개발한다.",
            "2단계(6~9월): 불량 검출 알고리즘 개발. 2024년 Burr 검사 알고리즘(검출기준 0.1mm 이상) 기반에 그린바디 전용 파라미터를 튜닝하고, AI 학습 기반 양불 판정 정확도 99% 이상을 달성한다.",
            "3단계(10~12월): 프레스 핸들러 인라인 연동 및 MES 데이터 연계. takt time 1.5초 이내 검사 완료, 금형별/제품별 불량률 트렌드 분석 기능을 구현한다.",
        ],
        "tasks": [
            {"name": "그린바디 전용 비전시스템", "period": "3개월", "start_date": "2025-03-01", "end_date": "2025-05-31", "resource": "기구설계 2 + 비전 1"},
            {"name": "불량 검출 알고리즘 개발", "period": "4개월", "start_date": "2025-06-01", "end_date": "2025-09-30", "resource": "비전 S/W 2명"},
            {"name": "핸들러 인라인 연동", "period": "3개월", "start_date": "2025-08-01", "end_date": "2025-10-31", "resource": "기구설계 1 + S/W 1"},
            {"name": "MES 연계 및 실증", "period": "2개월", "start_date": "2025-11-01", "end_date": "2025-12-31", "resource": "전원 (5명)"},
        ],
        "pm_name": "○○○", "pm_period": "10", "pm_role": "과제 총괄, 비전 검사",
        "sw_name": "○○○", "sw_period": "10", "sw_role": "알고리즘/MES 연동",
        "budget_detail": "재료비 / 외주가공비 / 시험비", "budget_amount": "별도 책정",
        "constraints": "그린바디의 낮은 강도로 비접촉 검사 필수, 프레스 takt time 내 검사 완료 필요",
    },
    {
        "id": 3,
        "slug": "연삭측정제어",
        "md": "연구개발-3-연삭측정제어.md",
        "title": "연삭체의 정밀 연삭 가공을 위한 측정 제어장치 및 그 제어방법",
        "eng_title": "Measurement Control Device and Method for Precision Grinding",
        "period": "2025년 4월 1일 ~ 2025년 12월 31일",
        "period_months": "9개월",
        "members": "총 4명 (기구설계 1, 제어설계 1, S/W 2)",
        "method": "인프로세스 측정 기술 개발, 폐루프 연삭 제어 알고리즘 설계, 하드웨어 통합",
        "purposes": [
            "초경인서트 양면 연삭 공정에서 연삭 전·중·후 치수를 자동 측정하고, 측정 데이터를 기반으로 연삭량을 실시간 피드백 제어하는 측정 제어장치를 개발한다.",
            "연삭 과정에서 목표 치수에 도달하면 자동 정지하는 폐루프(Closed-loop) 연삭 제어 시스템을 구현하여, 과삭/미삭 불량을 90% 이상 감소시킨다.",
            "제품별 연삭 이력(측정값, 연삭량, 연삭 조건)을 자동 기록하고 SPC 기반 공정 능력 분석 체계를 구축한다.",
        ],
        "necessities": [
            "현재 연삭 공정은 작업자 숙련도에 따라 연삭량이 결정되며, 과삭(Over-grinding) 또는 미삭(Under-grinding)으로 인한 불량이 빈발한다. 자동 측정·제어로 작업자 숙련도 무관한 일정 품질 확보가 필요하다.",
            "2023년 과제는 연삭 전/후 측정→연삭량 계산→연삭기 전송의 오프라인 방식이었다. 본 과제는 연삭 중 실시간 측정(In-process measurement)과 폐루프 피드백 제어를 추가하여 목표 치수 대비 편차 ±1µm 이내를 달성하는 고정밀 제어 시스템으로 고도화한다.",
        ],
        "strategies": [
            "1단계(4~6월): 인프로세스 측정 기술 개발. 비접촉 레이저/에어 마이크로미터 기반 연삭 중 두께 실시간 측정(분해능 ±0.5µm 이하). 리니어스케일(3.75μm) 및 프로브(2μm) 적용 방안을 검증하고, 연삭 진동·절삭유 환경 대응 안정성을 확보한다.",
            "2단계(7~9월): 폐루프 연삭 제어 알고리즘 개발. 치수제어 폐루프 피드백 오차분석 결과를 반영하여 연삭 전 측정→목표 연삭량 산출→연삭 진행→실시간 측정→목표 도달 시 정지. 접근/정삭/스파크아웃 자동 전환, 열변형 보정 포함.",
            "3단계(10~12월): 측정 제어장치 하드웨어 통합 및 연삭기 CNC 인터페이스(EtherCAT/PROFINET) 연동. MES 품질 데이터 실시간 전송, 현장 실증 수행.",
        ],
        "tasks": [
            {"name": "인프로세스 측정 기술", "period": "3개월", "start_date": "2025-04-01", "end_date": "2025-06-30", "resource": "기구설계 1 + S/W 1"},
            {"name": "폐루프 제어 알고리즘", "period": "3개월", "start_date": "2025-07-01", "end_date": "2025-09-30", "resource": "제어설계 1 + S/W 1"},
            {"name": "H/W 통합 및 CNC 연동", "period": "3개월", "start_date": "2025-10-01", "end_date": "2025-12-31", "resource": "전원 (4명)"},
            {"name": "MES 연계 및 현장 실증", "period": "2개월", "start_date": "2025-11-01", "end_date": "2025-12-31", "resource": "S/W 2명"},
        ],
        "pm_name": "○○○", "pm_period": "9", "pm_role": "과제 총괄, 제어 설계",
        "sw_name": "○○○", "sw_period": "9", "sw_role": "알고리즘/MES 연동",
        "budget_detail": "재료비 / 외주가공비 / 시험비", "budget_amount": "별도 책정",
        "constraints": "연삭 진동·칩·절삭유 환경에서의 측정 안정성 확보 난이도 높음, 연삭기 CNC 통신 인터페이스 호환성",
    },
    {
        "id": 4,
        "slug": "포장혼입검사",
        "md": "연구개발-4-포장혼입검사.md",
        "title": "인서트 포장기 혼입검사기술 개발",
        "eng_title": "Development of Cross-contamination Inspection Technology for Insert Packaging",
        "period": "2025년 4월 1일 ~ 2025년 12월 31일",
        "period_months": "9개월",
        "members": "총 4명 (기구설계 1, 비전 S/W 2, S/W 1)",
        "method": "다중 특징 기반 제품 식별 시스템 설계, 고속 검사 광학계 개발, 포장기 인라인 연동",
        "purposes": [
            "초경인서트 최종 포장 공정에서 이종 제품 혼입(Cross-contamination)을 자동 검출하는 인라인 검사 기술을 개발한다.",
            "인서트의 형상·각인·색상·치수를 고속(0.8초/개 이내)으로 검사하여, 혼입 제품을 실시간 선별하는 시스템을 구현한다.",
            "포장 단위별 검사 이력을 기록하여 완전한 트레이서빌리티를 확보하고, 고객 클레임 원인인 혼입 불량을 근절한다.",
        ],
        "necessities": [
            "현재 포장 공정에서 작업자 실수 또는 팔레트 혼재로 인해 타 규격 제품이 혼입되는 품질 사고가 간헐적으로 발생하고 있다. 혼입 불량은 고객 클레임의 주요 원인 중 하나이며, 글로벌 고객사 품질 요구사항(ISO/IATF) 충족을 위해 100% 자동 검사가 필수적이다.",
            "기존 검사 과제들은 단일 불량 유형(Burr, 깨짐, 호닝 치수)을 대상으로 했으나, 혼입 검사는 제품 식별(Identification) 문제로 형상 유사 인서트 간 미세 차이(각인 문자, 코팅 색상, 코너 R값)를 고속으로 판별하는 별도 기술이 필요하다.",
        ],
        "strategies": [
            "1단계(4~6월): 다중 특징 기반 제품 식별 시스템 설계. 형상 매칭, 딥러닝 OCR 각인 인식, 코팅 색상 분석, 핵심 치수 검증의 4대 식별 모듈을 설계한다. Korloy#6 혼입 광학계 이슈 분석 및 Nose R 혼입 측정편차 데이터를 반영한다.",
            "2단계(7~9월): 고속 검사 광학계 개발. 2022년 집광 LED + 다각도 조명 통합, 텔레센트릭 렌즈 적용, 상면·측면 동시 촬영(0.8초/개 이내). 딥러닝 OCR 테스트 결과를 반영한 혼입 판정 알고리즘(다중 특징 스코어링 + AI 학습)을 개발한다.",
            "3단계(10~12월): 포장기 인라인 연동 및 현장 실증. 기존 공정 흐름 변경 최소화하여 포장기 투입부에 설치, 혼입 자동 배출(리젝트) 기능을 구현한다.",
        ],
        "tasks": [
            {"name": "제품 식별 시스템 설계", "period": "3개월", "start_date": "2025-04-01", "end_date": "2025-06-30", "resource": "비전 S/W 2 + 기구 1"},
            {"name": "고속 검사 광학계 개발", "period": "3개월", "start_date": "2025-07-01", "end_date": "2025-09-30", "resource": "기구설계 1 + 비전 1"},
            {"name": "혼입 판정 알고리즘", "period": "3개월", "start_date": "2025-07-01", "end_date": "2025-09-30", "resource": "비전 S/W 1 + S/W 1"},
            {"name": "포장기 연동 및 실증", "period": "3개월", "start_date": "2025-10-01", "end_date": "2025-12-31", "resource": "전원 (4명)"},
        ],
        "pm_name": "○○○", "pm_period": "9", "pm_role": "과제 총괄, 비전 설계",
        "sw_name": "○○○", "sw_period": "9", "sw_role": "알고리즘/인라인 연동",
        "budget_detail": "재료비 / 외주가공비 / 시험비", "budget_amount": "별도 책정",
        "constraints": "형상 유사 인서트 간 미세 차이 판별 난이도 높음, 포장 라인 takt time 내 검사 완료 필요",
    },
    {
        "id": 5,
        "slug": "호닝신뢰성",
        "md": "연구개발-5-호닝신뢰성.md",
        "title": "정밀 광학계 기반 호닝형상검사기의 신뢰성 확보 기술 연구",
        "eng_title": "Reliability Assurance Technology for Precision Optical Honing Inspection Machine",
        "period": "2025년 5월 1일 ~ 2025년 12월 31일",
        "period_months": "8개월",
        "members": "총 4명 (비전 S/W 2, 광학설계 1, S/W 1)",
        "method": "측정 불확도 체계 분석, 자동 보정/자가 진단 시스템 개발, 신뢰성 검증 프로토콜 수립",
        "purposes": [
            "2022년 개발한 호닝형상 검사기(HIM)의 측정 신뢰성을 체계적으로 검증·향상하여, 고객사 양산 라인에서의 장기 안정 운용을 보장한다.",
            "환경 변화(온도·진동·조명), 광학 부품 경년 열화, 다양한 제품 규격 대응에서 발생하는 측정 편차 요인을 분석하고 자동 보정·자가 진단 기능을 개발한다.",
            "GR&R(Gage Repeatability & Reproducibility) 10% 이내 달성 및 1,000시간 연속 운전 시 정밀도 유지를 보증한다.",
        ],
        "necessities": [
            "HIM은 WTA의 글로벌 1위 핵심 제품으로 ±1µm 정밀도의 호닝 형상·치수 측정 장비이나, 고객사 현장 적용 시 환경 변화, 광학 열화, 다규격 대응에서 측정 편차가 보고되고 있다. 글로벌 경쟁력 유지를 위해 장기 신뢰성 확보가 시급하다.",
            "2022년 HIM 개발은 측정 기능 구현에 초점이 맞춰져 있었으나, 양산 현장에서의 장기 신뢰성 확보는 별도 과제로 다뤄지지 않았다. 2024년 Burr 검사 장치 현장 적용에서 발견된 광학계 이슈(비네팅, 조명 균일도)가 HIM에도 유사하게 적용되며, ISO 17025 등 측정 장비 인증 요구사항 대응도 필요하다.",
        ],
        "strategies": [
            "1단계(5~7월): 측정 불확도 체계 분석. 환경·광학·기구·제품 요인별 측정 편차 정량화, GR&R 분석 체계 수립. HIM-F 원가절감 항목 발굴 결과와 2024년 Burr 검사 현장 경험의 광학계 이슈 사례를 반영한다.",
            "2단계(8~10월): 자동 보정(Auto-Calibration) 및 자가 진단 시스템 개발. 클리닝 장치(Air&Brush, 레이저/플라즈마) 비접촉 방식 검토 결과를 반영하고, 표준 시편 기반 주기 보정, 온도 보상 알고리즘, 조명 세기 자동 조정, 광축 자동 정렬 보정, 다채널 분할검사 기반 광학 부품 교체 시기 예측 기능을 구현한다.",
            "3단계(11~12월): 신뢰성 검증 프로토콜 실행. 환경 변화 조건별 테스트(온도 챔버, 가진기), 1,000시간 연속 운전, 다규격(10종+) 대응, 고객사 현장 실증(2개소 이상)을 수행한다.",
        ],
        "tasks": [
            {"name": "측정 불확도 체계 분석", "period": "3개월", "start_date": "2025-05-01", "end_date": "2025-07-31", "resource": "광학설계 1 + 비전 1"},
            {"name": "자동 보정 시스템 개발", "period": "3개월", "start_date": "2025-08-01", "end_date": "2025-10-31", "resource": "비전 S/W 2 + S/W 1"},
            {"name": "자가 진단 기능 개발", "period": "2개월", "start_date": "2025-09-01", "end_date": "2025-10-31", "resource": "S/W 1 + 비전 1"},
            {"name": "신뢰성 검증 및 현장 실증", "period": "2개월", "start_date": "2025-11-01", "end_date": "2025-12-31", "resource": "전원 (4명)"},
        ],
        "pm_name": "○○○", "pm_period": "8", "pm_role": "과제 총괄, 광학 설계",
        "sw_name": "○○○", "sw_period": "8", "sw_role": "보정/진단 S/W",
        "budget_detail": "재료비 / 외주가공비 / 시험비", "budget_amount": "별도 책정",
        "constraints": "고객사 현장 실증을 위한 장비 반출/설치 일정 조율 필요, 1,000시간 장기 테스트 장비 확보",
    },
]


def replace_paragraph_text(paragraph, old_text, new_text):
    """단락의 텍스트를 교체하되 서식(run style)은 유지"""
    if old_text in paragraph.text:
        # 첫 번째 run에 새 텍스트 설정, 나머지 run 비우기
        runs = paragraph.runs
        if runs:
            # 전체 텍스트 교체
            full = paragraph.text.replace(old_text, new_text)
            runs[0].text = full
            for r in runs[1:]:
                r.text = ""
        return True
    return False


def add_paragraph_after(doc, ref_paragraph, text, style_name=None):
    """특정 단락 뒤에 새 단락 추가"""
    from docx.oxml.ns import qn
    new_p = copy.deepcopy(ref_paragraph._element)
    # 텍스트 설정
    for r in new_p.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            t.text = ""
    # 첫 run에 텍스트 넣기
    runs = new_p.findall(qn('w:r'))
    if runs:
        ts = runs[0].findall(qn('w:t'))
        if ts:
            ts[0].text = text
        else:
            from lxml import etree
            t_elem = etree.SubElement(runs[0], qn('w:t'))
            t_elem.text = text
    ref_paragraph._element.addnext(new_p)
    return new_p


def generate_doc(proj):
    """하나의 과제에 대해 Word 문서 생성"""
    # docxtpl로 테이블 Jinja2 변수 렌더링
    tpl = DocxTemplate(str(TEMPLATE))

    context = {
        "author_title": "담당", "author_dept": "생산관리팀", "author_sign": "",
        "reviewer_title": "검토", "reviewer_dept": "생산관리팀", "reviewer_sign": "",
        "approver_title": "승인", "approver_dept": "대표이사", "approver_sign": "",
        "tasks": proj["tasks"],
        "pm_name": proj["pm_name"], "pm_period": proj["pm_period"], "pm_role": proj["pm_role"],
        "sw_name": proj["sw_name"], "sw_period": proj["sw_period"], "sw_role": proj["sw_role"],
        "budget_detail": proj["budget_detail"], "budget_amount": proj["budget_amount"],
    }

    tpl.render(context)

    # 중간 저장 후 python-docx로 단락 편집
    import tempfile, os
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = tmp.name
    tmp.close()
    tpl.save(tmp_path)

    doc = Document(tmp_path)

    # ── 단락 텍스트 교체 ──
    for p in doc.paragraphs:
        text = p.text.strip()

        # 표지 영문 표기명
        if "필요시 영문 표기명" in text:
            replace_paragraph_text(p, "필요시 영문 표기명", proj["eng_title"])

        # 개발 목적 1, 2, 3
        if text == "개발목적 1":
            replace_paragraph_text(p, "개발목적 1", proj["purposes"][0])
        elif text == "개발목적 2":
            replace_paragraph_text(p, "개발목적 2", proj["purposes"][1])
        elif text.startswith("개발목적 3"):
            replace_paragraph_text(p, text, proj["purposes"][2])

        # 필요성
        if text == "개발 필요성 1":
            replace_paragraph_text(p, "개발 필요성 1", proj["necessities"][0])
        elif text.startswith("개발 필요성 2"):
            replace_paragraph_text(p, text, proj["necessities"][1])

        # 개발 개요
        if "개발 과제명 : 개발과제명" in text:
            replace_paragraph_text(p, "개발과제명", proj["title"])
        if "20**년 **월 **일 ~ 20**년 **월 **일" in text:
            replace_paragraph_text(p, "20**년 **월 **일 ~ 20**년 **월 **일", proj["period"])
        if "참여 인원 : 총  명" in text or "참여 인원 : 총" in text:
            replace_paragraph_text(p, p.text, f"참여 인원 : {proj['members']}")
        if "개발 방법 :" in text and len(text) < 15:
            replace_paragraph_text(p, p.text, f"개발 방법 : {proj['method']}")

        # 추진전략
        if text == "추진전략 1":
            replace_paragraph_text(p, "추진전략 1", proj["strategies"][0])

    # 추진전략 2, 3 추가 (추진전략 1 뒤에)
    for i, p in enumerate(doc.paragraphs):
        if proj["strategies"][0] in p.text:
            # 역순으로 추가해야 순서 유지
            for s_idx in range(len(proj["strategies"]) - 1, 0, -1):
                add_paragraph_after(doc, p, proj["strategies"][s_idx])
            break

    # 제약사항 추가 - '개발체계 및 업무 분장' 앞 빈 단락에 삽입
    for i, p in enumerate(doc.paragraphs):
        if "개발체계 및 업무 분장" in p.text:
            if i > 0:
                prev_p = doc.paragraphs[i-1]
                if not prev_p.text.strip():
                    if prev_p.runs:
                        prev_p.runs[0].text = f"[제약사항] {proj['constraints']}"
                    else:
                        # runs가 없으면 새 run 추가
                        run = prev_p.add_run(f"[제약사항] {proj['constraints']}")
            break

    # 저장
    out_path = OUT_DIR / f"연구개발계획서-{proj['id']}-{proj['slug']}.docx"
    doc.save(str(out_path))
    os.unlink(tmp_path)
    print(f"  생성 완료: {out_path.name}")
    return out_path


# ── 메인 실행 ──
if __name__ == "__main__":
    print("=" * 60)
    print("연구개발 계획서 5건 Word 문서 생성")
    print("=" * 60)

    results = []
    for proj in PROJECTS:
        try:
            path = generate_doc(proj)
            results.append((proj["id"], proj["slug"], True, str(path)))
        except Exception as e:
            print(f"  [오류] {proj['slug']}: {e}")
            import traceback
            traceback.print_exc()
            results.append((proj["id"], proj["slug"], False, str(e)))

    print()
    print("=" * 60)
    print("결과 요약:")
    for pid, slug, ok, info in results:
        status = "성공" if ok else "실패"
        print(f"  #{pid} {slug}: {status} - {info}")
    print("=" * 60)
