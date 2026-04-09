# -*- coding: utf-8 -*-
"""경상연구개발 문서 보완 스크립트 - 참고문서 기반으로 v2 docx 내용 보강"""
import sys, io, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt

V2_DIR = 'C:/MES/wta-agents/reports/MAX/경상연구개발/v2'

def get_all_text(doc):
    """문서의 모든 텍스트를 추출"""
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return '\n'.join(texts)

def find_paragraph_by_text(doc, search_text, start_from=0):
    """특정 텍스트를 포함하는 paragraph 인덱스 반환"""
    for i in range(start_from, len(doc.paragraphs)):
        if search_text in doc.paragraphs[i].text:
            return i
    return -1

def append_text_to_paragraph(para, additional_text):
    """paragraph에 텍스트 추가"""
    if para.text.strip() and not para.text.strip().endswith('.') and not para.text.strip().endswith('다') and not para.text.strip().endswith(')'):
        para.text = para.text.rstrip() + '\n' + additional_text
    else:
        para.text = para.text.rstrip() + '\n\n' + additional_text

def add_text_after_paragraph(doc, index, text, style='s0'):
    """특정 paragraph 뒤에 새 paragraph 추가"""
    new_para = doc.paragraphs[index]._element
    from docx.oxml.ns import qn
    from lxml import etree
    p_elem = etree.SubElement(new_para.getparent(), qn('w:p'))
    new_para.addnext(p_elem)
    # 간단히 기존 paragraph의 텍스트에 추가하는 방식 사용
    return None

def enhance_table_cell(table, row_idx, col_idx, additional_text):
    """테이블 셀에 텍스트 추가"""
    try:
        cell = table.rows[row_idx].cells[col_idx]
        if additional_text not in cell.text:
            cell.text = cell.text.rstrip() + '\n' + additional_text
    except (IndexError, AttributeError):
        pass

def process_1_장비물류():
    """1-장비물류 프로젝트 문서 보완"""
    changes = []

    # === 연구개발계획서 ===
    path = os.path.join(V2_DIR, '연구개발계획서-1-장비물류.docx')
    doc = Document(path)
    all_text = get_all_text(doc)

    # 필요성 섹션 보강 - 설비별 SPM 수치 추가
    idx = find_paragraph_by_text(doc, '필요성')
    if idx >= 0:
        next_idx = find_paragraph_by_text(doc, '경영회의', idx)
        if next_idx >= 0:
            existing = doc.paragraphs[next_idx].text
            if '소결취출기' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n설비별 현재 물류 대응 수준 비교:\n- 프레스핸들러: SPM Max. 23 EA/min (저울측정 포함, 연동 12~14EA), 셋업 3일 후 양산, CE 인증 획득\n- 소결취출기: SPM Max. 65 EA/min, 혼입검사(형상, C/B, Nose-R) 기능 내장\n- 연삭핸들러: SPM Max. 16 EA/min, 특화 비전기술, 다양한 제품 대응\n- PVD 로딩기: SPM Max. 40 EA/min, 다양한 ROD 길이(220~500mm) 대응\n- 포장기: SPM Max. 60 EA/min, 완벽한 혼입 검사, 레이저·잉크 마킹 검사\n- 검사기: Macro 최소 검출 70μm, Micro 최소 14μm, 치수 정밀도 ±5μm'
                changes.append('계획서-필요성: 설비별 SPM 수치 및 현재 물류 대응 수준 추가')

    # 추진전략 1단계 - AMR 5 type 상세 추가
    idx = find_paragraph_by_text(doc, '1단계(3~5월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if 'type 1' not in existing:
            doc.paragraphs[idx].text = existing + '\n\nAMR 물류 5 type 상세:\n- type 1: 화루이 프레스 AGV 타입 (Fork형 AGV 기반)\n- type 2: 한국야금 Cell 프레스 타입 (Cell 라인 연동)\n- type 3: 대구텍 키엔스 검사기 타입 (검사 장비 물류)\n- type 4: 한국야금 포장기/검사기 타입 (6호기 연동)\n- type 5: 한국야금 호닝핸들러 타입 (호닝 공정 물류)'
            changes.append('계획서-추진전략 1단계: AMR 5 type 구체적 분류 추가')

    # 추진전략 2단계 - ATC U19 상세
    idx = find_paragraph_by_text(doc, '2단계(6~9월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if 'Softgrip' not in existing and 'Tilt' not in existing:
            doc.paragraphs[idx].text = existing + '\n\nATC U19 공용 pickup 툴 요구사항:\n- 5종 Gripper 호환: Jaw, Vacuum, Magnet, Softgrip, Tilt Grip\n- pogopin 배선을 통한 센서 신호선 대응\n- 중량 최소화 목표(현재 2~3kg)\n- 그리퍼 간격 65mm 유지\n- 멀티 조인트 & 슬립링 적용 (공압 1~4포트, 6포트 2종 표준설계, 전기배선 8~12core)\n- R축 모터 용량 상향 검토'
            changes.append('계획서-추진전략 2단계: ATC U19 기술 요구사항 상세 추가')

    doc.save(path)

    # === 결과보고서 ===
    path = os.path.join(V2_DIR, '결과보고서-1-장비물류.docx')
    doc = Document(path)

    # 개발 내용 개요 - 설비별 차별점 상세 추가
    idx = find_paragraph_by_text(doc, '개발 내용 개요')
    if idx >= 0:
        next_idx = find_paragraph_by_text(doc, '11개 장비군', idx)
        if next_idx >= 0:
            existing = doc.paragraphs[next_idx].text
            if 'CBN' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n■ 설비별 차별점 상세 (WTA 기술 자산)\n- 프레스: 고속 SPM 23, 신속 셋업 3일→양산, 디버링 기술(접촉식/비접촉식/Hole), 듀얼 툴 안정화\n- PVD 언로딩기: SPM Max. 35, 제품 간격 4mm 대응 이슈 해결 과제 (전체의 30%)\n- CVD 로딩/언로딩기: SPM Max. 35, 수직 제품 로딩, 카본판 높이 자동 보정\n- CBN 조립기: 2개소 기준 5EA/Min, 정밀 부착(0.05mm), 용접재 정밀정량 도포\n- 호닝형상검사기: 반복정밀도 ±2μm, 최소 5μm 이상 검사, 4포인트 1분 측정'
                changes.append('결과보고서-개발내용: 설비별 차별점 상세 수치 추가 (CBN, PVD, CVD 등)')

    # 개발 결과 - 고객사별 대응 상세 추가
    idx = find_paragraph_by_text(doc, 'Concept')
    if idx >= 0:
        next_idx = idx + 1
        if next_idx < len(doc.paragraphs):
            existing = doc.paragraphs[next_idx].text
            if '대구텍' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n  - 대구텍: AMR(AGV) 장비 내부 물류 구조 대응, 키엔스 검사기 연동\n  - 한국야금: Cell 프레스형+포장기/검사기형+호닝핸들러형 3종 동시 표준화'
                changes.append('결과보고서-Concept: 대구텍, 한국야금 고객사별 대응 상세 추가')

    doc.save(path)

    # === 연구노트 ===
    path = os.path.join(V2_DIR, '연구노트-1-장비물류.docx')
    doc = Document(path)

    # TABLE-2 (1. 연구 목표) 셀 보강
    if len(doc.tables) > 2:
        cell = doc.tables[2].rows[1].cells[0]
        existing = cell.text
        if '설계팀 요소기술' not in existing:
            cell.text = existing + '\n\n■ 연구 배경 (경영회의 지시사항)\n- 2024년 1월 경영회의: HAM 기구설계팀에서 "AGV AMR 무인화 장비별 표준물류 컨셉 확립" 보고\n- 2025년 12월 경영회의: 대표이사 "DX 시스템 구축: 데이터 기반 의사결정 체계 확립" 지시\n- 설계팀 요소기술 개발항목에 "AMR 물류 표준화 설계" 정식 등재 (5 type 분류)'
            changes.append('연구노트-연구목표: 경영회의 지시사항 및 설계팀 요소기술 개발항목 추가')

    # TABLE-3 (2. 실험/개발 내용) 셀 보강
    if len(doc.tables) > 3:
        cell = doc.tables[3].rows[1].cells[0]
        existing = cell.text
        if 'Components' not in existing:
            cell.text = existing + '\n\n■ 2. 고속 직교축 Components 기술 확보\n· X,Y,Z,R 형태의 복합축 구조 정밀고속구동 기술 내재화\n· Roboworker components 사업 3가지 item:\n  - 3D프린트 팔레트 부자재\n  - 프레스 금형 분말피더 shoes\n  - Linear Axes(belt)\n· 검사기, PVD, 포장기에 공용 적용 설계\n\n■ 3. 회전축 멀티 조인트 & 슬립링 표준화\n· 공압 1~4포트, 6포트 2종 표준설계\n· 전기배선 8~12core\n· 적용 대상: ATC 내재화, 전제품 반전Unit(90도/180도), 포장기 회전Unit'
            changes.append('연구노트-실험내용: 고속 직교축 Components, 멀티조인트·슬립링 표준화 기술 상세 추가')

    # TABLE-4 (3. 결과 및 데이터) 셀 보강
    if len(doc.tables) > 4:
        cell = doc.tables[4].rows[1].cells[0]
        existing = cell.text
        if '엘리베이터' not in existing:
            cell.text = existing + '\n· 엘리베이터 교체시간: 속도 100% 기준 29.89초 (SPM 15, 시간 비율 7.5%)\n· 프레스핸들러 SPM Max. 23 EA/min 달성 (저울측정 포함)\n· AMR type별 표준 인터페이스 사양서 작성 완료'
            changes.append('연구노트-결과: 엘리베이터 교체시간, SPM 실측 데이터 추가')

    doc.save(path)
    return changes

def process_2_분말검사():
    """2-분말검사 프로젝트 문서 보완"""
    changes = []

    # === 연구개발계획서 ===
    path = os.path.join(V2_DIR, '연구개발계획서-2-분말검사.docx')
    doc = Document(path)

    # 개발 목적 보강 - 검사 사양 상세
    idx = find_paragraph_by_text(doc, '개발 목적')
    if idx >= 0:
        next_idx = find_paragraph_by_text(doc, '성형 직후', idx)
        if next_idx >= 0:
            existing = doc.paragraphs[next_idx].text
            if '소결 취출기' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n후공정 연동 검사 사양:\n- 소결 취출기: 치수 측정(내접원, 비대칭) Macro ±5μm, 깨짐 100μm 이상, 밴딩/뒤틀림 판정\n- 연삭 핸들러: 높이 측정 (하드웨어적 접촉 측정)\n- 호닝 핸들러: 호닝 유무 검사, 혼입검사\n- 공통 혼입 검사: CB 혼입, Nose-R 혼입(0.2mm 이상), 현재 카메라 FOV 45×35mm(1.3MP, 0.03mm/pixel)'
                changes.append('계획서-개발목적: 후공정 연동 검사 사양 상세 추가')

    # 추진전략 1단계 - 조명 테스트 상세
    idx = find_paragraph_by_text(doc, '1단계(3~5월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if 'HV Macro' not in existing:
            doc.paragraphs[idx].text = existing + '\n\nOTC 광학기술센터 연계 작업:\n- HV Macro 조명 제작 진행 (고시인성 prototype)\n- 동축 조명, 돔 조명(간접→직접 8ch), DF 조명(1열→2열) 개발\n- Macro 조명 시인성 개선: 모델명 WRRDDRL188-S24W 개발'
            changes.append('계획서-추진전략 1단계: OTC 광학센터 HV Macro 조명 개발 상세 추가')

    # 추진전략 2단계 - 고해상도 광학계 정보
    idx = find_paragraph_by_text(doc, '2단계(6~9월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '12MP' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n고해상도 Macro 광학계 검토:\n- 현행 Pixel 수 12MP → 신규 25MP로 업그레이드 검토\n- Dynamic Range: 69.03dB vs 56.32dB, SNR: 44.55dB vs 42.13dB\n- 분해능: 13.9μm → 7.8μm (25MP 적용 시)\n- 5MP 카메라 추가 검토(NoseR 0.1mm 대응)'
            changes.append('계획서-추진전략 2단계: 고해상도 Macro 광학계 사양 비교 추가')

    doc.save(path)

    # === 결과보고서 ===
    path = os.path.join(V2_DIR, '결과보고서-2-분말검사.docx')
    doc = Document(path)

    # 개발 내용 개요 보강
    idx = find_paragraph_by_text(doc, '개발 내용 개요')
    if idx >= 0:
        next_idx = find_paragraph_by_text(doc, '측면 광학계', idx)
        if next_idx >= 0:
            existing = doc.paragraphs[next_idx].text
            if 'Chromatic confocal' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n■ 추가 기술 검토\n- KAIST 김정원 교수 3D profilometry 기술 센싱: ToF 부적합 → Chromatic confocal 방법 제시\n- Feasibility test 진행 (4/18 1차)\n- 포장기 혼입검사 측정편차 이슈 병행 해결\n- 팔레트 광학계 카메라 단종 대응: Basler acA1300-30gm → acA1300-60gm 전환'
                changes.append('결과보고서-개발내용: KAIST 3D profilometry, 카메라 단종 대응 추가')

    # Concept 보강
    idx = find_paragraph_by_text(doc, 'Concept')
    if idx >= 0:
        next_idx = idx + 1
        if next_idx < len(doc.paragraphs):
            existing = doc.paragraphs[next_idx].text
            if '인서트 제조 공정' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n분말야금 공정 특성 반영:\n- 분말 야금(P/M): 금속 분말을 가열·결합하여 초경합금 인서트 제조\n- 성형 공정: 분말 충진 → 가압(상하 압력) → 탈형 → 배출\n- 핵심 관리 포인트: 압력·중량 관리, 복잡 형상 대응\n- 검사 시점: 성형 직후(그린바디 상태)에서 Burr 검출이 가장 효과적'
                changes.append('결과보고서-Concept: 분말야금 공정 특성 및 검사 시점 근거 추가')

    doc.save(path)

    # === 연구노트 ===
    path = os.path.join(V2_DIR, '연구노트-2-분말검사.docx')
    doc = Document(path)

    # TABLE-3 (2. 실험/개발 내용) 보강
    if len(doc.tables) > 3:
        cell = doc.tables[3].rows[1].cells[0]
        existing = cell.text
        if '인서트 제조 공정' not in existing:
            cell.text = existing + '\n\n■ 2. 인서트 제조 공정과 검사 연계\n· 분말야금(P/M) → 성형(가압) → 소결(1,600~1,800°C) → 연삭 → 호닝 → 코팅(CVD/PVD)\n· 성형체 Burr: 금형 마모 판단 핵심 지표, 인선부 0.1mm 초과 시 불량\n· 코팅 전 단계에서 검출해야 공정 손실 최소화\n\n■ 3. 조명 비교 테스트 결과\n· 동축조명: Burr 경사면 반사로 시인 불가\n· T-BLU(Telecentric Back Light Unit): 배경 밝기 부족\n· FlatDome: Both부(배경) 밝기 극대화, Burr 실루엣 최적\n· 라운드 바조명: 45~135도 15도 step 테스트, 105도 이상 최적'
            changes.append('연구노트-실험내용: 인서트 제조 공정 연계, 조명 비교 테스트 상세 결과 추가')

    # TABLE-5 (4. 고찰 및 분석) 보강
    if len(doc.tables) > 5:
        cell = doc.tables[5].rows[1].cells[0]
        existing = cell.text
        if '48V Strobe' not in existing:
            cell.text = existing + '\n· 48V Strobe Mode 1000회 점등으로 광량 보상 성공\n· 소성체 제품이 코팅 제품보다 Both부 대비가 더 뚜렷하여 검출 유리\n\n[향후 과제]\n· 고해상도 Macro 광학계(12MP→25MP) 적용 시 분해능 13.9μm→7.8μm 향상 기대\n· KAIST 3D profilometry(Chromatic confocal) 기술 접목 검토\n· 포장기 label 검사 광학계(CIS 방식) 연동 검토'
            changes.append('연구노트-고찰: 48V Strobe 결과, 향후 과제(고해상도, 3D profilometry) 추가')

    doc.save(path)
    return changes

def process_3_연삭측정제어():
    """3-연삭측정제어 프로젝트 문서 보완"""
    changes = []

    # === 연구개발계획서 ===
    path = os.path.join(V2_DIR, '연구개발계획서-3-연삭측정제어.docx')
    doc = Document(path)

    # 개발 목적 보강 - 기술적 이슈 상세
    idx = find_paragraph_by_text(doc, '위치 피드백과 실제 휠')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '기구적 결합 오차' in existing and '평행도' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n기술적 이슈 상세:\n1) 측정부와 실제 연삭 부위의 위치 차이로 기구적 결합 오차(이격) 발생 → 휠 평행도 불안정, 연삭 대상물 치수 편차 증가\n2) 모터 Enable 정지 상태에서도 공압하중 변화 시 외부스케일 위치 변동 발생\n3) 연삭 중 외부스케일 수치가 진동하면서 하강하는 현상 관측\n4) 위치 도달 후 휠 감속 시간 동안 추가 연삭 발생 (하중 감소+휠 상승 버퍼구간 포함)\n5) 외부스케일을 피드백으로 직접 위치제어 시 버퍼길이 변화로 진동·오류 발생 → 모터 엔코더 기반 제어 필요하나 실제 외부스케일 수치와 차이'
            changes.append('계획서-개발목적: 연삭 위치 피드백 기술적 이슈 5가지 상세 추가')

    # 추진전략 1단계 보강
    idx = find_paragraph_by_text(doc, '1단계(4~6월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '사다리꼴 속도 프로파일' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n외부스케일 기반 별도 위치 제어 알고리즘 상세:\n- 목표위치·속도 입력 → 실시간 서보 시스템 속도 제어 실행\n- 이동거리·가감속 계산 후 사다리꼴 속도 프로파일 추출\n- 제어주기(10ms)마다 외부스케일 위치 판독, 프로파일 대비 속도 지령 반복\n- 삼각형/사다리꼴 프로파일 자동 판단 (가속거리 대비 이동거리 비교)\n- EtherCAT PC통신 기반 C# 프로그램 제어\n- 알고리즘 적용 시 0~3μm 오차 관측 (외부스케일 기준 정밀 제어 달성)'
            changes.append('계획서-추진전략 1단계: 사다리꼴 속도 프로파일 알고리즘 상세 추가')

    # 추진전략 2단계 보강 - 연삭 조건별 영향 분석
    idx = find_paragraph_by_text(doc, '2단계(7~9월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '정지시점(목표위치)' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n연삭 조건별 치수 영향 분석:\n- 제어 가능 요소: 연삭하중, 정지시점(목표위치), 회전방향·가감속, 속도(rpm)\n- 핵심 발견: 정지시점(목표위치)이 치수 결과에 가장 큰 영향, 나머지 요소는 연삭시간·품질(광택, 날 깨짐, Burr)에 영향\n- 치수 제어 전략: 정지시점만 변경하는 것이 최고 신뢰성, 치수 피드백 보정 적용\n  (예: 목표 4.86mm → 결과 4.96mm → 보정치 -0.010 적용)\n- 편차 관리: 연삭 결과물 편차가 크면 피드백 보정 신뢰성 저하 → 편차 최소화가 선결 과제'
            changes.append('계획서-추진전략 2단계: 연삭 조건별 치수 영향 분석 및 제어 전략 추가')

    # 제약사항 보강
    idx = find_paragraph_by_text(doc, '[제약사항]')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '피드백 편차' not in existing:
            doc.paragraphs[idx].text = existing + ' (3) 외부스케일 피드백 편차가 랜덤하게 발생하여 서보 시스템의 위치제어와 비교하면 응답성이 떨어지는 한계 존재. (4) 연삭 편차 요인: 연마재와 연삭체 접촉 밸런스(제어불가 기구적 요소) + 연삭 속도/궤적분포 밸런스(제어가능 요소).'
            changes.append('계획서-제약사항: 피드백 편차 랜덤 발생, 연삭 편차 요인 분석 추가')

    doc.save(path)

    # === 결과보고서 ===
    path = os.path.join(V2_DIR, '결과보고서-3-연삭측정제어.docx')
    doc = Document(path)

    # 개발 내용 개요 보강 - 성능지표 상세
    idx = find_paragraph_by_text(doc, '개발 내용 개요')
    if idx >= 0:
        next_idx = find_paragraph_by_text(doc, '양면 연삭 핸들러', idx)
        if next_idx >= 0:
            existing = doc.paragraphs[next_idx].text
            if '성능지표' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n■ 성능지표 달성 현황\n- 연삭 치수 정밀도: ±3μm 이내 (CNMG120408, 두께 4.860mm 기준, KOLAS 검증)\n- 연삭 평탄도: 제품 내 양 끝 2포인트 편차 ±3μm 이내\n- 반복 가공 정밀도: 동일 조건 연속 가공 시 뱃치 간 편차 ±3μm 이내\n- 연삭 소요시간: CNMA1204 65개 기준 30초/cycle 이내\n- 연삭 가공압력: 최대 2톤 하중 가능 (로드셀 측정 검증)\n- 가공치수 측정 정밀도: 영점 위치 반복측정 정밀도 1μm\n- 핸들러 픽업 속도: 2000mm/s에서 저진동 구동 달성\n- 비전 인식 에러율: 64개 10회 반복 테스트 (연삭유 미제거 상태)\n- 핸들러 생산성: 64EA/cycle 로딩-언로딩 256초 이내'
                changes.append('결과보고서-개발내용: 9개 성능지표 달성 현황 상세 추가')

    # Concept 보강 - 장비 비교 데이터
    idx = find_paragraph_by_text(doc, 'Concept')
    if idx >= 0:
        next_idx = idx + 1
        if next_idx < len(doc.paragraphs):
            existing = doc.paragraphs[next_idx].text
            if '후지산키' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n장비 비교 분석 (3종):\n- 후지산키 350 오버홀: Ø350 휠, 4축, 측면 분사, 스퍼기어(자동화 불리), 1976년 주물(강성 낮음)\n- WTA 1호기(후지산키 협업): Ø430 휠, 5축, 핀기어(자동화 유리), 주물 1톤 이상 증량(강성 향상)\n  → 제품 내 편차 최대 4μm, 제품 간 편차 최대 6μm 달성\n- 국가과제 개발장비: Ø530 휠, 5축, 연삭유 직접분사(열화 방지), 주물 최대(강성 최고)\n  → 단, 헤드 중량이 모터부 대비 2배 이상, 무게 중심 쏠림 이슈 존재'
                changes.append('결과보고서-Concept: 후지산키/WTA/국가과제 3종 장비 비교 데이터 추가')

    doc.save(path)

    # === 연구노트 ===
    path = os.path.join(V2_DIR, '연구노트-3-연삭측정제어.docx')
    doc = Document(path)

    # TABLE-3 (2. 실험/개발 내용) 보강
    if len(doc.tables) > 3:
        cell = doc.tables[3].rows[1].cells[0]
        existing = cell.text
        if '프로브 측정기' not in existing:
            cell.text = existing + '\n\n■ 2. 측정 정밀도 개선\n· 리니어 스케일 정밀도: 3.75μm\n· 프로브 측정기 정밀도: 2μm\n· 프로브 측정기를 제품 가까운 거리에 배치하여 정밀도 극대화\n· 핀 진직도/캐리어 유격/지그 평탄도 불일치 → 연마 불균형 발생 → 보정 방안 개발\n\n■ 3. DX 자동화 시스템\n· 자동 측정: 비전 방식 선정 (접촉식 대비 비용·개발기간·물류 연동성 우수)\n· 초음파 세척 → 에어나이프 건조 → 비전 측정 공정 분리·자동화\n· 클리닝 시 인서트 특성(사이즈, 재질) 반영 수압·공압 조절\n· CPK 기반 품질 관리, 1로트별 샘플 치수 데이터 관리'
            changes.append('연구노트-실험내용: 측정 정밀도(프로브/스케일), DX 자동화 시스템 상세 추가')

    # TABLE-4 (3. 결과 및 데이터) 보강
    if len(doc.tables) > 4:
        cell = doc.tables[4].rows[1].cells[0]
        existing = cell.text
        if '드레싱' not in existing:
            cell.text = existing + '\n· 외부스케일 기반 위치제어 알고리즘: 0~3μm 오차 달성\n· 사다리꼴 프로파일 제어: 10ms 주기, EtherCAT C# 구현\n· 자동 드레싱 스테이션: 엘리베이터+피더 통합 설계\n· 드레싱 스톤+연삭 지그 자동 공급·회수 시스템 구축\n· 건식 vs 습식 연마 비교 테스트 수행 (열화 현상 차이 분석)\n· WTA 1호기 실측: 제품 내 최대 4μm, 제품 간 최대 6μm'
            changes.append('연구노트-결과: 알고리즘 성과, 드레싱 스테이션, 실측 데이터 추가')

    # TABLE-5 (4. 고찰 및 분석) 보강
    if len(doc.tables) > 5:
        cell = doc.tables[5].rows[1].cells[0]
        existing = cell.text
        if '주물' not in existing:
            cell.text = existing + '\n\n[기술적 고찰]\n· 주물 설비 특성: 진동 감쇠성(흡진성), 열 안정성, 높은 강성(동일 부피 판금 대비)\n· 내구성 핵심 요소: 피로수명, 내마모성, 내열성, 균열 저항성, 재료 조직 안정성\n· 강성 핵심 요소: 정적 강성(연삭 휠 동심도/평탄도), 동적 강성(진동 억제), 열 변형 민감도\n· 국가과제 장비 이슈: 헤드 무게 중심 쏠림 → 주축부 주물 부하·연삭 정밀도 저하 우려\n\n[향후 계획]\n· 소량 다품종 대응: RFID 기반 드레싱룸 지그 자동 식별\n· AGV 물류 연동: 팔레트 안정성 확보(진동 제거)\n· GPT 기반 교육 플랫폼 검토'
            changes.append('연구노트-고찰: 주물 설비 특성, 내구성·강성 분석, 향후 계획 추가')

    doc.save(path)
    return changes

def process_4_포장혼입검사():
    """4-포장혼입검사 프로젝트 문서 보완"""
    changes = []

    # === 연구개발계획서 ===
    path = os.path.join(V2_DIR, '연구개발계획서-4-포장혼입검사.docx')
    doc = Document(path)

    # 개발 목적 보강
    idx = find_paragraph_by_text(doc, '인서트의 형상·각인·색상·치수')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '케이스 내 수량' not in existing or '마킹 유무' not in existing:
            pass  # 이미 포함되어 있으면 스킵
        if '수량검사 Vision' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n현장 이슈 기반 검사 요구사항:\n- 수량검사 Vision 오검출 → 커버 조립부 에러 연계 발생 사례\n- 팔레트 검출 오류: 팔레트와 유사 특징의 오검출 → 다중 검출 시 무시 기능 필요\n- 마킹 검사: 잉크젯·레이저 마킹 후 문자 유무/품질 검사, 마킹색·촬상 각도에 따른 이미지 품질 저하 대응\n- 라벨 부착 검사: 라벨 말림, 2장 출력, 부착 실패 등 다양한 이슈 대응'
            changes.append('계획서-개발목적: 현장 이슈(수량검사 오검출, 팔레트 오류, 마킹·라벨) 추가')

    # 필요성 보강 - 포장기 현장 이슈 통계
    idx = find_paragraph_by_text(doc, '포장 단계에서의 혼입은 고객 클레임')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '포장기#4' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n현장 CS 이력 분석 (4개 고객사 포장기):\n- 쫜스 포장기 #2~#4(신·구공장): 혼입검사, 케이스 공급/반전, 그리퍼 동심도, 센서 오동작 등 30건 이상 이슈\n- 한국야금 포장기 #1: 케이블 손상, 스텝모터 무언정지\n- 교세라 포장기: 마킹 검출 카메라 불량, 이미지 전송 불량\n- 공통 이슈: Zebra 프린터 타임아웃, 온도 센서 표시 오류, PC 재부팅(RTX 에러)'
            changes.append('계획서-필요성: 포장기 CS 이력 분석 4개 고객사 이슈 통계 추가')

    # 추진전략 2단계 - 알고리즘 상세
    idx = find_paragraph_by_text(doc, '2단계(7~9월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if 'PatternMatching' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n향상된 위치검출 툴 개발 상세:\n- PatternMatching: 다각형 폐루프 Tracker로 인서트 영역 설정\n- CB 방향 Align: 학습 기반 방향 자동 감지 및 회전 보정\n- 앞/뒷면 구분: 표면 패턴 비교 기반\n- CB 혼입 판단: ChipBreaker 형태 분류\n- NoseR 혼입 모델기능: NoseR 값 비교 기반 혼입 판정\n- 효과: R형 인서트 측면 검사 가능, 치수 반복측정 코너 기준 정렬'
            changes.append('계획서-추진전략 2단계: 향상된 위치검출 툴(PatternMatching, CB Align 등) 상세 추가')

    doc.save(path)

    # === 결과보고서 ===
    path = os.path.join(V2_DIR, '결과보고서-4-포장혼입검사.docx')
    doc = Document(path)

    # 개발 내용 개요 보강
    idx = find_paragraph_by_text(doc, '개발 내용 개요')
    if idx >= 0:
        next_idx = find_paragraph_by_text(doc, '혼입분류 기능', idx)
        if next_idx >= 0:
            existing = doc.paragraphs[next_idx].text
            if '알고리즘팀' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n■ 알고리즘팀 검출력 향상 성과\n- G급 연삭무늬 검사: 연삭무늬 무력화 검사 알고리즘 개발\n- Chamfer 제품: 고배율/저배율 영역 분리 개선\n- 날방향 제품: 측면 영역 분리 개발 및 Batch Test 완료\n- 얕은 깨짐 검출력: 이미지 처리로 광량 향상 효과 테스트\n- 저배율 50μm 이상 검출: CO 이미지 검출력 개선\n- 양면날 제품: 측면 영역 분리 개선, 상부 검사 대응'
                changes.append('결과보고서-개발내용: 알고리즘팀 검출력 향상 성과 6건 추가')

    # Concept 보강
    idx = find_paragraph_by_text(doc, 'Concept')
    if idx >= 0:
        next_idx = idx + 1
        if next_idx < len(doc.paragraphs):
            existing = doc.paragraphs[next_idx].text
            if 'AI' not in existing or 'Deep-Learning' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\nAI/Deep Learning 적용:\n- 서버-검사기 네트워크 연결 완료\n- Deep-Learning 결함검사: Macro 외관 검출력 테스트\n- 정상이미지 과검출 테스트, Multipage 학습 조건 테스트\n- 한국야금 모델 성능 평가 수행\n- OCR: 문자인식 이미지 전후처리, 모델 서치 및 성능 테스트'
                changes.append('결과보고서-Concept: AI/Deep Learning 결함검사 적용 상세 추가')

    doc.save(path)

    # === 연구노트 ===
    path = os.path.join(V2_DIR, '연구노트-4-포장혼입검사.docx')
    doc = Document(path)

    # TABLE-3 (2. 실험/개발 내용) 보강
    if len(doc.tables) > 3:
        cell = doc.tables[3].rows[1].cells[0]
        existing = cell.text
        if '케이스 공급' not in existing:
            cell.text = existing + '\n\n■ 2. 포장기 기구 이슈 대응\n· 케이스 공급 기울어짐: FD1 0.3mm, FD2 0.5mm → 클램프 레벨링 기구 개선\n· 케이스 반전 틀어짐: 적재 오차 → 반전지그 여유 +1.2mm→4mm 확장\n· 그리퍼 동심도 개선: 툴 헤드 교체 시 맞춤핀 탈착 어려움 → 기구 개선\n· 그리퍼 캡 종류별 픽업 높이 차이 → 범용성 캡 제작 방안 모색\n· 커버 누름쇠 감지 이슈 → 센서 방식 변경 또는 커버 누름쇠 변경\n\n■ 3. 혼입검사 광학계 대응\n· 팔레트 광학계 카메라 단종: Basler acA1300-30gm → acA1300-60gm 전환\n· 포장기 label 검사: CIS(Contact Image Sensor) 방식 검토\n· 마킹검사 조명 개선: 마킹색·촬상 각도에 따른 이미지 품질 저하 대응'
            changes.append('연구노트-실험내용: 포장기 기구 이슈 대응 5건, 광학계 대응 3건 추가')

    # TABLE-5 (4. 고찰 및 분석) 보강
    if len(doc.tables) > 5:
        cell = doc.tables[5].rows[1].cells[0]
        existing = cell.text
        if '레시피 서버' not in existing:
            cell.text = existing + '\n\n[향후 계획]\n· 레시피 서버 구축: 검사 결과 데이터 DB화 (파일→DB 전환)\n· Recipe Station: 장비 외부에서 레시피 제작·적용 (이미지 수신→레시피 작성→송신→적용)\n· Cell 라인 Press Handler 외관검사 연동: 인선높이·버높이 측정 기능 통합\n· 향상된 검출 툴 정밀 티칭 및 초점 티칭 기능 고도화'
            changes.append('연구노트-고찰: 레시피 서버, Recipe Station, Cell 라인 연동 향후 계획 추가')

    doc.save(path)
    return changes

def process_5_호닝신뢰성():
    """5-호닝신뢰성 프로젝트 문서 보완"""
    changes = []

    # === 연구개발계획서 ===
    path = os.path.join(V2_DIR, '연구개발계획서-5-호닝신뢰성.docx')
    doc = Document(path)

    # 개발 목적 보강 - 72시간 롱런 데이터 상세
    idx = find_paragraph_by_text(doc, '72시간 롱런')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '평균 CT' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n72시간 롱런 테스트 상세 결과:\n- #23-2호기 (3/9~12): 4,539회 측정, 평균 CT 57.17초\n- #23-4호기 (3/16~19): 4,261회 측정, 평균 CT 61.90초\n- 온도 변화 1°C당 Z축 높이 15μm 변화 확인\n- 초기 가동 2.5~3시간(22°C 이하) 측정값 부정확 구간 존재\n- FFU 가동 시 진동에 의한 산포 증가 현상'
            changes.append('계획서-개발목적: 72시간 롱런 테스트 상세(CT, 온도 영향, FFU 진동) 추가')

    # 필요성 보강 - 기구 개선 상세
    idx = find_paragraph_by_text(doc, '기구 개선 검토')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if 'Pallet 공급/취출' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n기구 개선 주요 항목 예시:\n- Pallet 공급/취출 분리 방식 개조\n- 상부 Door 마모 문제: SUS Plate 2t 추가\n- NG Pallet: 틀 공급 타입 → 4장 개별 공급 & Sensing\n- 리턴 Conv QR리더기 추가, Air Blow 위치 이동 (소형 인서트 틀어짐 대책)\n- 석정반 외곽 막음 커버, FFU·헤파필터 노후화 교체\n- 마그넷 그리퍼 2종 테스트 적용 (소형 돌기, 대형 Type)'
            changes.append('계획서-필요성: 기구 개선 주요 항목 6건 구체적 예시 추가')

    # 추진전략 2단계 - TTTM 상세
    idx = find_paragraph_by_text(doc, '2단계(8~10월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if 'Calibration 타겟' not in existing:
            doc.paragraphs[idx].text = existing + '\n\nTTTM 기술 상세:\n- 광학계 Calibration 타겟: Dot 타겟 3종 + 복합 타겟 2종 설계\n- Cross Hair 시인성 향상: 선폭 9→50μm 확대\n- 광학계 align process 확립: 모듈 정렬 체계화\n- 각 고객사향(MMC, Korloy, 한국교세라) 셋업 후 정렬 상태 점검\n- 시료 측정 기능 추가: TTTM 장치 개조\n\nCoaxPress 카메라 테스트 준비 (속도향상 목적)\nLaser Triangulation: 광삼각법 측정방식 변경 (Object Move), 측정시간 단축'
            changes.append('계획서-추진전략 2단계: TTTM Calibration 타겟, CoaxPress, Laser Triangulation 추가')

    # 추진전략 3단계 - 현장 실증 상세
    idx = find_paragraph_by_text(doc, '3단계(11~12월)')
    if idx >= 0:
        existing = doc.paragraphs[idx].text
        if '알루미늄 인서트' not in existing:
            doc.paragraphs[idx].text = existing + '\n\n다규격 검증 대상:\n- CNMA, TNMA 등 표준 인서트\n- 양면날 밀링 인서트: AI 검사 가능성 검토 (MMC 관심사항)\n- 알루미늄 인서트: 별도 크리닝 장치 필요 (Air & Brush → 비접촉 방식 검토)\n- G급 인서트: 깨짐 외 결함 기준서 기반 적용 범위 정의\n- CBN 제품: 고배율 테스트 (다인정공 3차 시료)'
            changes.append('계획서-추진전략 3단계: 다규격 검증 대상(양면날, 알루미늄, G급, CBN) 상세 추가')

    doc.save(path)

    # === 결과보고서 ===
    path = os.path.join(V2_DIR, '결과보고서-5-호닝신뢰성.docx')
    doc = Document(path)

    # 개발 내용 개요 보강
    idx = find_paragraph_by_text(doc, '개발 내용 개요')
    if idx >= 0:
        next_idx = find_paragraph_by_text(doc, '기구 개선 26개', idx)
        if next_idx >= 0:
            existing = doc.paragraphs[next_idx].text
            if '#1~2호기' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n■ 호닝검사기 장비별 개선 현황\n- #1~2호기: V.E제품 횡전개 (강성보강, 내재화 Stage, Bolt 체결방식), 개조 부품 100% 입고\n- #23-3~4호기: 지그 고정방식 변경 (마그넷→Tap), 지그별(Ø7/8/10/14) Ø0.3 홀 가공\n- Top Vision부: 조립 정밀도 개선을 위한 재설계 (고정 Rib와 B/Screw 분리, LM기준면·위치 고정핀 추가)\n- 하이썽(중국) 출하 장비: #23-4호기 미비사항 횡전개 조치'
                changes.append('결과보고서-개발내용: 호닝검사기 장비별(#1-4) 개선 현황 상세 추가')

    # Concept 보강 - 비전팀 성과
    idx = find_paragraph_by_text(doc, 'Concept')
    if idx >= 0:
        next_idx = idx + 1
        if next_idx < len(doc.paragraphs):
            existing = doc.paragraphs[next_idx].text
            if '산학' not in existing:
                doc.paragraphs[next_idx].text = existing + '\n\n비전팀 핵심 기술 성과:\n- 이미지 심도합성 알고리즘: 측면 반사광 간섭 개선\n- 자가검사기능(전처리): 측면 스테이지 평탄도 측정 완료, IOI_TMAC 기능 및 UI 수정\n- 고배율 마스킹: 밀링/그루빙 제품 고배율 마스킹 기능 개발\n- Python 알고리즘 연동: MTF 검사 알고리즘 Python 구현\n\n산학 협력:\n- 아주대 최수영 교수: 산학자문 협약 (~25.11.30)\n- 인하대 전병환 교수: 기술자문 협약 (2026.1.1~12.31)\n- 인하대 산학공동프로젝트 진행'
                changes.append('결과보고서-Concept: 비전팀 핵심 기술 4건, 산학 협력 3건 추가')

    doc.save(path)

    # === 연구노트 ===
    path = os.path.join(V2_DIR, '연구노트-5-호닝신뢰성.docx')
    doc = Document(path)

    # TABLE-3 (2. 실험/개발 내용) 보강
    if len(doc.tables) > 3:
        cell = doc.tables[3].rows[1].cells[0]
        existing = cell.text
        if '72Hr 롱런' not in existing:
            cell.text = existing + '\n\n■ 2. 72Hr 롱런 안정성 테스트\n· #23-2 (3/9~12): 4,539회 측정, 평균 CT 57.17초, 온도변화 1°C당 Z축 15μm\n· #23-4 (3/16~19): 4,261회 측정, 평균 CT 61.90초\n· 초기 가동 22°C 이하 2.5~3시간 측정값 부정확 구간 확인\n· FFU 가동 시 진동 산포 증가 → 방진 대책 수립\n\n■ 3. 대구텍 NEW F2 검사기 (신규 개발 연계)\n· 설계 출도율: 60% 달성 (석정반, 반전기, Top Vision, 3-side Vision, 엘리베이터, 팔레트피더)\n· ATC Station: Gripper 3종(3Jaw, Side 2Jaw, Magnet), Side Chuck Jig\n· 장비 높이 2.5m, 현장 Hoist 높이 대비 상부 Frame 조립 방안 검토\n· 고정 비전 4개소(W.D: 760mm) 성능 Risk 대책 마련'
            changes.append('연구노트-실험내용: 72Hr 롱런 테스트 상세, 대구텍 NEW F2 신규 개발 연계 추가')

    # TABLE-5 (4. 고찰 및 분석) 보강
    if len(doc.tables) > 5:
        cell = doc.tables[5].rows[1].cells[0]
        existing = cell.text
        if '원가절감' not in existing:
            cell.text = existing + '\n\n[원가절감 성과]\n· HIM-F Series 원가절감 항목 발굴\n· 호닝형상검사기 커버 설계 수정으로 절감 방안 확인\n· Macro 조명 시인성 개선: 기구 부품 제작(삼호테크), PCB Inner/Outer Ring SMT\n\n[향후 과제]\n· CoaxPress 카메라 도입으로 측정 속도 향상\n· Laser Triangulation 광삼각법 적용(Object Move 방식, 측정시간 단축)\n· 알루미늄 인서트 크리닝: 전병환 교수 제안 Chemical 세정→초음파→Dry Hot Air 검토\n· 중국 교세라 마그넷 그리퍼 스크래치 대책: PEEK 재질 변경, 인선부 도피 타입 3종 테스트'
            changes.append('연구노트-고찰: 원가절감 성과, 향후 과제(CoaxPress, 크리닝, 그리퍼 대책) 추가')

    doc.save(path)
    return changes


# === 메인 실행 ===
if __name__ == '__main__':
    all_changes = {}

    print("=" * 60)
    print("경상연구개발 문서 보완 작업 시작")
    print("=" * 60)

    print("\n[1/5] 장비물류 프로젝트 보완 중...")
    all_changes['1-장비물류'] = process_1_장비물류()
    for c in all_changes['1-장비물류']:
        print(f"  ✓ {c}")

    print("\n[2/5] 분말검사 프로젝트 보완 중...")
    all_changes['2-분말검사'] = process_2_분말검사()
    for c in all_changes['2-분말검사']:
        print(f"  ✓ {c}")

    print("\n[3/5] 연삭측정제어 프로젝트 보완 중...")
    all_changes['3-연삭측정제어'] = process_3_연삭측정제어()
    for c in all_changes['3-연삭측정제어']:
        print(f"  ✓ {c}")

    print("\n[4/5] 포장혼입검사 프로젝트 보완 중...")
    all_changes['4-포장혼입검사'] = process_4_포장혼입검사()
    for c in all_changes['4-포장혼입검사']:
        print(f"  ✓ {c}")

    print("\n[5/5] 호닝신뢰성 프로젝트 보완 중...")
    all_changes['5-호닝신뢰성'] = process_5_호닝신뢰성()
    for c in all_changes['5-호닝신뢰성']:
        print(f"  ✓ {c}")

    print("\n" + "=" * 60)
    print("보완 완료 요약")
    print("=" * 60)
    total = 0
    for proj, clist in all_changes.items():
        print(f"\n[{proj}] {len(clist)}건")
        for c in clist:
            print(f"  - {c}")
        total += len(clist)
    print(f"\n총 {total}건 보완 완료")
