"""
연구노트 5개 과제 생성 스크립트
template-연구노트.docx 기반, docxtpl + python-docx + 이미지 삽입
"""
import os
import sys
from copy import deepcopy
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = r"C:\MES\wta-agents\reports\MAX\template-연구노트.docx"
OUTPUT_DIR = r"C:\MES\wta-agents\reports\MAX\경상연구개발"
IMG_BASE = r"C:\MES\wta-agents\data\core-tech-images"

# ======================================================================
# 프로젝트 데이터
# ======================================================================

PROJECTS = [
    {
        "id": 1,
        "filename": "연구노트-1-장비물류.docx",
        "note_number": "RND-2025-001",
        "page": "1",
        "total_pages": "4",
        "project_name": "장비 무인화운영을 위한 장비 물류 개발",
        "research_period": "2025.03 ~ 2025.12",
        "date": "2025.04.01",
        "department": "생산관리팀 (AI운영팀)",
        "author": "조한종",
        "reviewer": "최종국",
        "author_sign": "",
        "reviewer_sign": "",
        "reviewer_date": "2025.04.01",
        "objective": (
            "초경인서트 제조 공정(성형-소결-연삭-코팅-검사) 전 구간에서 공정 간 제품 이송을 "
            "완전 자동화하여 24시간 무인 연속 생산 체계를 구축한다.\n\n"
            "핵심 목표:\n"
            "1) AMR 5 type 표준화 설계 (화루이/한국야금 Cell/대구텍 키엔스/포장기/호닝핸들러)\n"
            "2) ATC 2세대 공용 pickup 툴 (U19) 개발 - 멀티조인트 & 슬립링 적용\n"
            "3) MES 연동 물류 스케줄링 - 실시간 WIP 모니터링 및 동적 우선순위 조정\n"
            "4) 설비별 takt time 기반 물류 최적화 (프레스 23EA/min, 소결 60EA/min, 검사 20EA/min)"
        ),
        "content": (
            "1. 사용 장비/자재\n"
            "  - AMR: 5 type 표준화 설계 (type1: 화루이 프레스 AGV, type2: 한국야금 Cell 프레스, "
            "type3: 대구텍 키엔스 검사기, type4: 한국야금 포장기/검사기, type5: 한국야금 호닝핸들러)\n"
            "  - ATC 2세대 U19: Jaw/Vacuum/Magnet/Softgrip/Tilt Grip 호환, 그리퍼 간격 65mm, 중량 2~3kg\n"
            "  - 멀티조인트 & 슬립링: 공압 1~4/6포트, 전기 8~12core 표준설계\n"
            "  - RFID/바코드 기반 팔레트 추적 시스템\n\n"
            "2. 실험 조건\n"
            "  - 11개 장비군(프레스, 소결취출기, 연삭핸들러, PVD 로딩/언로딩, CVD 로딩/언로딩, "
            "포장기, CBN조립기, 검사기, 호닝검사기) 각각에 대한 물류 인터페이스 설계\n"
            "  - 버퍼 스테이션: 투입/취출 이중 버퍼 구조\n"
            "  - AGV/AMR 동선 최적화: 교차 회피 알고리즘 적용\n\n"
            "3. 수행 절차\n"
            "  Phase 1 (3~5월): 공정간 물류 경로 맵핑 및 AMR type별 표준 인터페이스 설계\n"
            "  Phase 2 (6~9월): ATC U19 제작, 멀티조인트/슬립링 적용 시작, 단위 테스트\n"
            "  Phase 3 (10~12월): MES 연동 통합 테스트, 야간 무인 운전 실증\n\n"
            "[사진 부착: AMR 물류 시스템 개념도]"
        ),
        "results": (
            "■ 설계 검토 결과 (1차)\n\n"
            "1. AMR 표준화 설계 현황\n"
            "  - 5 type 구분 완료, type별 장비 내부 물류 AGV 대응 구조 정의\n"
            "  - 현재 장비별로 다른 구조로 개발 진행 중인 문제 확인 → 통일화 필요\n\n"
            "2. 설비별 생산성 기준 데이터\n"
            "  | 설비 | SPM | 비고 |\n"
            "  | 프레스 | 23 EA/min | 저울측정 포함 |\n"
            "  | 소결취출기 | 60 EA/min | 100~200만/월 |\n"
            "  | 연삭핸들러 | 16 EA/min | 20만/월 |\n"
            "  | 검사기(F2) | 20 EA/min | 대구텍 24EA/min |\n"
            "  | PVD로딩 | 40 EA/min | |\n"
            "  | 포장기 | 60 EA/min | |\n\n"
            "3. ATC U19 설계 착수\n"
            "  - Auto Tool Change 기능, 5종 그리퍼 호환\n"
            "  - R축 모터 용량 상향 검토 (멀티조인트 하중 증가)\n\n"
            "[데이터 표/그래프 부착 영역]"
        ),
        "analysis": (
            "1. AMR 물류 표준화의 핵심 과제\n"
            "  - 5 type의 장비-AMR 인터페이스를 통일화해야 원가 절감과 유지보수 효율 확보 가능\n"
            "  - 경영회의(2024.06) 대표이사 지시: '무인화를 통한 생산성 향상'이 핵심 경영전략\n"
            "  - 경영회의(2025.12) 방향: '각 공정별 AGV를 이용한 무인화 구현'\n\n"
            "2. 설계 표준화 필수 요소\n"
            "  - 고속 직교 축 표준화: X,Y,Z,R 형태 복합축 정밀고속구동 기술 내재화\n"
            "  - Roboworker Components 사업 벤치마킹 (3D프린트 팔레트, 분말피더, Linear Axes)\n\n"
            "3. 리스크 요인\n"
            "  - 소형 제품(1.5mm~) 핸들링 정밀도 확보\n"
            "  - 분진/연삭유 환경에서의 센서 내구성\n"
            "  - 다양한 팔레트 규격 호환 (이종 팔레트 간 리팔레팅 기술 필요)\n\n"
            "향후 계획: Phase 2 ATC U19 시작품 제작 및 단위 테스트 진행"
        ),
        "images": [
            os.path.join(IMG_BASE, "2025", "장비 무인화운영을 위한 장비 물류 개발", "img_000.png"),
        ],
    },
    {
        "id": 2,
        "filename": "연구노트-2-분말검사.docx",
        "note_number": "RND-2025-002",
        "page": "1",
        "total_pages": "4",
        "project_name": "프레스성형 품질향상을 위한 분말성형체 검사기술 개발",
        "research_period": "2025.03 ~ 2025.12",
        "date": "2025.04.01",
        "department": "생산관리팀 (AI운영팀)",
        "author": "서제완",
        "reviewer": "최종국",
        "author_sign": "",
        "reviewer_sign": "",
        "reviewer_date": "2025.04.01",
        "objective": (
            "프레스 성형 직후 분말성형체(그린바디)의 외관 결함(Burr, 깨짐, 크랙, 치수 불량)을 "
            "인라인으로 자동 검사하는 기술을 개발한다.\n\n"
            "핵심 목표:\n"
            "1) Burr 0.1mm 이하 검출 (2024년 Burr 검사기 기반 고도화)\n"
            "2) OTC 광학기술센터 프레스-IM 광학계 설계 적용\n"
            "3) 검출률 99% 이상, 사이클 1.5초 이내\n"
            "4) 금형 상태 연관 분석 → 금형 교체 시점 예측"
        ),
        "content": (
            "1. 사용 장비/자재\n"
            "  - 카메라: 고해상도 산업용 (상면/측면 멀티 카메라 구성)\n"
            "  - 조명: OTC 프레스-IM 광학계 - 집광 LED + 다각도 조명 통합\n"
            "  - 프레스 핸들러: 기존 SPM 23EA/min 핸들러 연동\n"
            "  - 참조 선행기술: 2024년 Burr 검사장치 (검출률 98%, 사이클 1.5초)\n\n"
            "2. 실험 조건\n"
            "  - 검사 대상: 분말성형체(그린바디) - 소결 전 상태, 강도 낮음, 무광 표면\n"
            "  - Burr 판정 기준: 0.1mm 이상 → 불량\n"
            "  - 그린바디 전용 조명 파라미터 최적화 필요\n"
            "  - AI 학습 데이터셋: 양품/불량 각 1,000매 이상 수집 목표\n\n"
            "3. 수행 절차\n"
            "  Phase 1 (3~5월): 그린바디 표면 특성 분석, 전용 조명 설계\n"
            "  Phase 2 (6~9월): 검출 알고리즘 개발 (Burr/깨짐/크랙/치수 복합 검출)\n"
            "  Phase 3 (10~12월): 프레스 핸들러 인라인 연동, 롱런 테스트\n\n"
            "[사진 부착: 분말성형 공정 및 검사 광학계]"
        ),
        "results": (
            "■ 사전 조사 및 광학계 설계 결과 (1차)\n\n"
            "1. 분말 야금 제조 공정 분석\n"
            "  - 분말 준비 → 성형(프레스 가압) → 소결(1,600~1,800도) → 연삭 → 코팅 → 검사\n"
            "  - 성형 직후 그린바디는 강도가 매우 낮아 비접촉 검사 필수\n\n"
            "2. 선행 과제 성과 활용\n"
            "  - 2024 Burr 검사장치: 고속 비전 1.5초, 검출률 98%, 핸들러 세트화 완료\n"
            "  - 2024 픽업+클리닝 복합툴: 검사 투입 전 성형체 표면 상태 확보\n\n"
            "3. OTC 광학계 사양\n"
            "  - 프레스-IM 전용 광학계: 저배율/고배율 이중 구성 검토\n"
            "  - 고객사별 제품 다양성: C형, W형, 특수 형상 대응 필요\n\n"
            "[표/그래프 부착 영역]"
        ),
        "analysis": (
            "1. 그린바디 검사의 기술적 도전\n"
            "  - 소결 전 상태(무광, 분말 부착, 낮은 강도)에서의 비접촉 검사 최적화 필수\n"
            "  - 2024 Burr 검사기는 '소결 후 제품' 대상 → 그린바디 전용 파라미터 재설계\n\n"
            "2. 조기 불량 검출의 경제적 효과\n"
            "  - 소결 전 불량 선별 → 후공정(소결/연삭/코팅) 원자재 손실 60% 이상 절감 기대\n"
            "  - 금형 상태 연관 분석으로 금형 교체 시점 예측 → 유지보수 효율화\n\n"
            "3. 남은 과제\n"
            "  - 미세 크랙 검출 알고리즘 (표면 질감과 크랙 구분)\n"
            "  - AI 학습 데이터 확보 (양품/불량 분류 기준 정립)\n"
            "  - 분진 환경에서의 광학계 내구성 확보\n\n"
            "향후 계획: Phase 2 알고리즘 개발 착수, 그린바디 시료 확보"
        ),
        "images": [
            os.path.join(IMG_BASE, "2025", "프레스성형 품질향상을 위한 분말성형체 검사기술 개발", "img_000.png"),
            os.path.join(IMG_BASE, "2024", "프레스 성형체의 burr 검사 장치 개발", "img_000.png"),
        ],
    },
    {
        "id": 3,
        "filename": "연구노트-3-연삭측정제어.docx",
        "note_number": "RND-2025-003",
        "page": "1",
        "total_pages": "4",
        "project_name": "연삭체의 정밀 연삭 가공을 위한 측정 제어장치 및 그 제어방법",
        "research_period": "2025.04 ~ 2025.12",
        "date": "2025.04.15",
        "department": "생산관리팀 (AI운영팀)",
        "author": "김웅기",
        "reviewer": "최종국",
        "author_sign": "",
        "reviewer_sign": "",
        "reviewer_date": "2025.04.15",
        "objective": (
            "초경인서트 양면 연삭 공정에서 연삭 전/중/후 치수를 자동 측정하고, "
            "측정 데이터 기반 연삭량 실시간 피드백 제어하는 폐루프(Closed-loop) 연삭 제어 시스템을 개발한다.\n\n"
            "핵심 목표:\n"
            "1) 외부스케일 기반 사다리꼴 프로파일 위치 제어 - 오차 0~3um\n"
            "2) EtherCAT C# 실시간 제어 프로그램 개발\n"
            "3) CNMG120408 기준 연삭 정밀도 +/-3um 이내\n"
            "4) 열변형 보정 알고리즘 (연삭 중 온도 상승 보상)"
        ),
        "content": (
            "1. 사용 장비/SW\n"
            "  - 연삭기: 양면연삭기 (비트리파이드 #400 휠)\n"
            "  - 외부스케일: 분해능 3.75um (리니어 스케일)\n"
            "  - 프로브: 접촉식 2um 분해능\n"
            "  - 제어: EtherCAT PC통신 기반 C# 프로그램\n"
            "  - 측정 대상: CNMG120408 초경인서트, ISO 기준 두께 4.860mm\n\n"
            "2. 실험 조건\n"
            "  - 외부스케일 피드백 위치 제어: 사다리꼴 프로파일 형태\n"
            "    * 제어주기 10ms, 완료허용오차 0.5um\n"
            "    * 가속구간 -> 등속구간 -> 감속구간 자동 전환\n"
            "  - 연삭 정지시점(목표위치) 정밀 제어\n"
            "  - 연삭 하중 20,000N (2톤) 강성 검증\n"
            "  - 연삭량 0.2mm 기준 테스트\n\n"
            "3. 수행 절차\n"
            "  Phase 1 (4~6월): 외부스케일 기반 위치 제어 알고리즘 구현 및 검증\n"
            "  Phase 2 (7~9월): 폐루프 연삭 제어 통합 (측정-연산-제어 루프)\n"
            "  Phase 3 (10~12월): 열변형 보정, 다품종 대응, 반복 정밀도 검증"
        ),
        "results": (
            "■ 위치 제어 알고리즘 설계 및 검증 결과\n\n"
            "1. 외부스케일 기반 위치 제어 이슈 분석\n"
            "  - 이슈 1: 외부스케일과 실제 휠 위치 차이 발생\n"
            "    * 측정부와 연삭 발생부 사이 기구적 결합 오차(이격)\n"
            "    * 모터 Enable 정지 상태에서도 공압하중 변화로 외부스케일 위치 변동\n"
            "  - 이슈 2: 휠 정지 시 감속 시간 동안 추가 연삭 진행\n"
            "    * 하중 감소 + 휠 리프트 + 감속 시간 복합 영향\n\n"
            "2. 대안: 사다리꼴 프로파일 속도 제어\n"
            "  - 목표위치/속도 입력 -> 실시간 외부스케일 판독 -> 속도 지령 반복\n"
            "  - 적용 결과: 0~3um 오차 관측 (비교적 정확한 제어 가능)\n\n"
            "3. 성능 지표 정의\n"
            "  | 항목 | 목표 | 측정방법 |\n"
            "  | 치수 정밀도 | +/-3um | 양 끝 2포인트 측정, 최대-최소 편차 |\n"
            "  | 평탄도 | +/-3um | 제품 내 양 끝 차이값 최대치 |\n"
            "  | 반복 정밀도 | +/-3um | N회 연삭 각 회차 최대/최소 편차 |\n"
            "  | 가공 시간 | 30초/cycle | 65개 CNMA1204 기준 |\n"
            "  | 가공 압력 | 20,000N | 로드셀 측정 |"
        ),
        "analysis": (
            "1. 핵심 발견\n"
            "  - 연삭 결과 치수의 핵심 요인은 '정지 시점(목표위치)' → 다른 조건(하중, RPM) 변경 없이 "
            "정지 시점만 변경하는 것이 치수 제어에 가장 신뢰성 높음\n"
            "  - 정지 시점과 결과 치수는 선형적 관계 → 피드백 보정으로 수렴 가능\n\n"
            "2. 편차 관리의 중요성\n"
            "  - 단일 결과물이 아닌 다수 제품 동시 연삭 시 편차가 존재\n"
            "  - 편차가 크면 피드백 보정의 신뢰성 저하\n"
            "  - 연삭 편차 원인: (1) 연마재-연삭체 접촉 밸런스 (제어불가), (2) 연삭 속도/궤적 밸런스 (제어가능)\n\n"
            "3. C# 제어 프로그램 특성\n"
            "  - EtherCAT PC통신 기반이므로 서보 내장 위치제어 대비 응답성 열세\n"
            "  - 보완: 제어주기 10ms, 확인 횟수 3회, 삼각형/사다리꼴 프로파일 자동 전환\n\n"
            "향후 계획: Phase 2 폐루프 통합, 실제 연삭 환경에서 검증 진행"
        ),
        "images": [],
    },
    {
        "id": 4,
        "filename": "연구노트-4-포장혼입검사.docx",
        "note_number": "RND-2025-004",
        "page": "1",
        "total_pages": "4",
        "project_name": "인서트 포장기 혼입검사기술 개발",
        "research_period": "2025.04 ~ 2025.12",
        "date": "2025.04.15",
        "department": "생산관리팀 (AI운영팀)",
        "author": "조윤명",
        "reviewer": "최종국",
        "author_sign": "",
        "reviewer_sign": "",
        "reviewer_date": "2025.04.15",
        "objective": (
            "초경인서트 최종 포장 공정에서 이종 제품 혼입(Cross-contamination)을 "
            "자동 검출하는 인라인 검사 기술을 개발한다.\n\n"
            "핵심 목표:\n"
            "1) Korloy #6 광학계 이슈 해결 - 광축과 조명 중심 정렬 +/-1mm 이내\n"
            "2) 딥러닝 OCR 기반 각인 인식 - 다국어(한/영/중) 대응\n"
            "3) acA2500-14gm 카메라 기반 고속 검사 (0.8초/개)\n"
            "4) 형상/각인/색상/치수 다중 특징 혼입 판정"
        ),
        "content": (
            "1. 사용 장비/자재\n"
            "  - 카메라: acA2500-14gm (Basler)\n"
            "  - 렌즈: M2514-MP2\n"
            "  - 조명: DOMELIGHT100\n"
            "  - 설계치 LWD: 169mm, FOV: 38.251 x 28.856mm (mag. 0.1490x)\n"
            "  - 실측치 LWD: 169mm, FOV 단변 약 31mm (mag. 약 0.138x) - 설계치와 차이 있음\n"
            "  - 딥러닝 OCR 엔진: VisionPro EL 기반\n\n"
            "2. 실험 조건\n"
            "  - Korloy #6 밝기 비대칭 이슈 재현 테스트:\n"
            "    * 조명 높이 11mm~19mm (2mm step) 조정\n"
            "    * 제품 위치 중앙 기준 상하 5mm (1mm step) 이동\n"
            "  - C형, W형 경면 제품에서 한쪽 비정상 밝기 현상 확인\n"
            "  - 중국교세라 현장 적용: 딥러닝 OCR 테스트, TN60/PV720 에러 감소 확인\n\n"
            "3. 수행 절차\n"
            "  Phase 1 (4~6월): 광학계 광축-조명 정렬 세팅 기준 정립, 밝기 비대칭 해결\n"
            "  Phase 2 (7~9월): 딥러닝 OCR 적용 및 다국어 마킹 인식 알고리즘\n"
            "  Phase 3 (10~12월): 다중 특징(형상/각인/색상/치수) 종합 혼입 판정 시스템\n\n"
            "[사진 부착: 혼입검사부 광학계 및 밝기 비대칭 테스트]"
        ),
        "results": (
            "■ 광축 정렬 및 OCR 테스트 결과\n\n"
            "1. Korloy #6 밝기 비대칭 분석\n"
            "  - 원인: 광축 중심이 조명 중심보다 좌측 하단에 위치\n"
            "  - 조명 높아질수록(6.2mm->14.2mm) 비대칭 현상 강화\n"
            "  - 제품 위치가 중심에서 벗어날수록 비대칭 강화\n"
            "  - 해결: 광축-조명 중심 상하좌우 +/-1mm 이내 세팅, 조명 높이 10mm 권장\n\n"
            "2. 딥러닝 OCR 현장 적용 (중국교세라)\n"
            "  - 코드 수정 개발 완료, 현장 1~2일 적용 테스트\n"
            "  - TN60, PV720 등 제품 에러 상당 감소\n"
            "  - 혼입 검사 시퀀스 개선 (오검출 감소)\n\n"
            "3. 포장기 이슈 현황 (실측 데이터)\n"
            "  - 수량 검사 Vision 검출 오류, 팔레트 검출 오류 등 개선 완료\n"
            "  - 경면 제품 & 랜드부 경사 큰 제품은 측정 가능 제품군 제한 필요\n\n"
            "[데이터 표/이미지 부착 영역]"
        ),
        "analysis": (
            "1. 광학계 세팅 표준화 필요\n"
            "  - 혼입검사부 초기 세팅 시 광축-조명 중심 정렬이 핵심\n"
            "  - 경면 제품(C형, W형)의 재발 가능성 → 측정 가능 제품군 정의 필수\n\n"
            "2. OCR 기술 발전 방향\n"
            "  - 딥러닝 기반 OCR이 기존 룰 기반 대비 에러율 크게 감소\n"
            "  - 다국어 대응 (한/영/중) 및 다양한 폰트(RomanS, SogoeUI, HY견고딕) 처리 필요\n"
            "  - 마킹 품질 자체의 편차(빛 방향/각도에 따른 식별성) 해결 필요\n\n"
            "3. 포장기 전체 시스템 안정성\n"
            "  - 포장기 셋업 기간 장기화 문제(경쟁사 5일 이내 vs WTA 초과)\n"
            "  - 복합 마킹(레이저+잉크, 상면/측면) 대응 필요\n"
            "  - 혼입검사 외 수량검사/마킹검사/케이스적재 연동 안정화\n\n"
            "향후 계획: Phase 2 딥러닝 OCR 확대 적용 및 다국어 인식 성능 향상"
        ),
        "images": [],
    },
    {
        "id": 5,
        "filename": "연구노트-5-호닝신뢰성.docx",
        "note_number": "RND-2025-005",
        "page": "1",
        "total_pages": "4",
        "project_name": "정밀 광학계 기반 호닝형상검사기의 신뢰성 확보 기술 연구",
        "research_period": "2025.05 ~ 2025.12",
        "date": "2025.05.01",
        "department": "생산관리팀 (AI운영팀)",
        "author": "김준형v",
        "reviewer": "최종국",
        "author_sign": "",
        "reviewer_sign": "",
        "reviewer_date": "2025.05.01",
        "objective": (
            "WTA 글로벌 1위 제품인 호닝형상검사기(HIM)의 측정 신뢰성을 체계적으로 검증/향상하여 "
            "고객사 양산 라인 장기 안정 운용을 보장한다.\n\n"
            "핵심 목표:\n"
            "1) GR&R 10% 이내 달성 (현 수준 대비 50% 개선)\n"
            "2) TTTM(Tool/Tip/Target Management) 자동 보정 체계 구축\n"
            "3) 온도 변화 대응 - 1도C당 Z축 15um 편차 보정\n"
            "4) Calibration 타겟 5종 표준화"
        ),
        "content": (
            "1. 사용 장비/SW\n"
            "  - HIM-F1: 선삭 인서트 검사 (SPM Max.10pcs/min)\n"
            "  - HIM-F2: 밀링+선삭 복합 검사 (SPM Max.20pcs/min)\n"
            "  - HIM-H: 호닝형상 검사/측정기\n"
            "  - 광학계: 저배율+고배율 이중 구성, 텔레센트릭 렌즈\n"
            "  - Macro 검출: 70um x 70um, Micro 검출: 14um x 14um\n"
            "  - 치수 측정 정밀도: +/-5um\n"
            "  - Vision팀: 김준형v, 진소미, 황인정, 정진원 외 8명\n\n"
            "2. 실험 조건\n"
            "  - 환경 변수: 온도 변화(+/-5도C), 진동(외부/내부), 습도\n"
            "  - 광학 변수: 레이저 출력 편차, 렌즈 열변형, 조명 열화\n"
            "  - TTTM WhiteBalance: 밝기 불일치 원인 분석\n"
            "    * 조명 제어 변수, 온도 차이, 입력 전원 차이에 따른 반사특성\n"
            "  - F2 선제작 장비: Pallet 공급/취출 분리, 상부 Cover 개조\n\n"
            "3. 수행 절차\n"
            "  Phase 1 (5~7월): 측정 불확도 체계 분석, GR&R 기준 수립\n"
            "  Phase 2 (8~10월): TTTM 자동보정, 온도 보상 알고리즘 개발\n"
            "  Phase 3 (11~12월): 장기 안정성 테스트, 고객사 현장 실증"
        ),
        "results": (
            "■ HIM 현황 조사 및 신뢰성 이슈 분석\n\n"
            "1. HIM 제품 라인업 현황\n"
            "  - F1: 한국야금, 교세라(중국), Dijet(일본), 다인정공\n"
            "  - F2: MMC(일본), 교세라(일본), 한국야금, 대구텍(NEW F2 개발중)\n"
            "  - H: 한국야금(4대), 교세라, 하이썽(중국), 펑마이\n\n"
            "2. 주요 신뢰성 이슈\n"
            "  - TTTM WhiteBalance: 입력 전원 차이에 따른 반사특성 차이 → 전류타입 컨트롤러 검토\n"
            "  - 카메라 Align: Rolling으로 인한 검출 결함 오판정\n"
            "  - 세라티즈 F1: 인선부 Section 오판정, 플레이트 위치결정력 부족\n"
            "  - 소형 인서트 틀어짐: Air Blow 위치 이동으로 대응\n\n"
            "3. 기구 설계 개선\n"
            "  - 강성보강(VE 제품 횡전개), 내재화 Stage\n"
            "  - Top Vision 부 조립 정밀도 개선: 고정 Rib/B-Screw 분리, LM기준면/핀 추가\n"
            "  - HIM-F 원가 절감 항목 발굴: 커버 설계 수정으로 절감 방안\n\n"
            "[데이터 표/이미지 부착 영역]"
        ),
        "analysis": (
            "1. 신뢰성 핵심 요인\n"
            "  - 온도 변화 보상이 가장 시급: 1도C 변화 시 Z축 15um 편차 발생\n"
            "  - TTTM WhiteBalance 불일치는 전류타입 컨트롤러로 근본 해결 가능\n"
            "  - 카메라 Align 체크리스트화 → 필드 셋업 품질 표준화\n\n"
            "2. 고객사 대응 전략\n"
            "  - 대구텍 NEW F2: 20pcs/min 이상 요구, ATC 적용, 신규 Macro 광학계\n"
            "  - 경영회의 지시: '키엔스 장비와의 차이를 자료화하여 영업 교육'\n"
            "  - 중국 시장: 매크로 기능 강화로 가성비 높은 장비 전략\n\n"
            "3. AI/DX 방향\n"
            "  - Deep Learning 결함검사: Macro 외관 검출력 향상, Multipage 학습\n"
            "  - Recipe Station: 장비 외부에서 레시피 제작/적용\n"
            "  - 산학 협력: 아주대 최수영 교수(산학자문), 인하대 전병환 교수(기술자문)\n\n"
            "향후 계획: Phase 2 TTTM 자동보정 구현, 온도 보상 알고리즘 적용 테스트"
        ),
        "images": [
            os.path.join(IMG_BASE, "2022", "호닝형상 검사기 고정밀급 개발", "img_000.png"),
        ],
    },
]


def generate_note(proj):
    """docxtpl로 렌더링 후 이미지를 삽입하여 최종 연구노트 생성"""
    # Step 1: docxtpl 렌더링
    tpl = DocxTemplate(TEMPLATE)
    context = {
        "note_number": proj["note_number"],
        "page": proj["page"],
        "total_pages": proj["total_pages"],
        "project_name": proj["project_name"],
        "research_period": proj["research_period"],
        "date": proj["date"],
        "department": proj["department"],
        "author": proj["author"],
        "reviewer": proj["reviewer"],
        "author_sign": proj["author_sign"],
        "reviewer_sign": proj["reviewer_sign"],
        "reviewer_date": proj["reviewer_date"],
        "objective": proj["objective"],
        "content": proj["content"],
        "results": proj["results"],
        "analysis": proj["analysis"],
    }
    tpl.render(context)

    out_path = os.path.join(OUTPUT_DIR, proj["filename"])
    tpl.save(out_path)

    # Step 2: 이미지 삽입 (python-docx)
    if proj["images"]:
        doc = Document(out_path)
        # 결과 섹션(Table 4) 끝에 이미지 삽입
        valid_imgs = [p for p in proj["images"] if os.path.exists(p)]
        if valid_imgs:
            # 결과 테이블(Table 4, index 4)의 데이터 셀에 이미지 추가
            if len(doc.tables) > 4:
                result_cell = doc.tables[4].rows[1].cells[0]
                for img_path in valid_imgs:
                    p = result_cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    try:
                        run.add_picture(img_path, width=Inches(4.5))
                    except Exception as e:
                        run.text = f"[이미지 삽입 실패: {os.path.basename(img_path)}]"
                        print(f"  이미지 경고: {e}")

        doc.save(out_path)

    return out_path


def verify_note(path, proj):
    """간단 검증: 작성자, 과제명 포함 여부"""
    doc = Document(path)
    full_text = " ".join([p.text for p in doc.paragraphs])
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += " " + cell.text

    has_author = proj["author"] in full_text
    has_project = proj["project_name"][:10] in full_text
    return has_author, has_project


# ======================================================================
# Main
# ======================================================================
if __name__ == "__main__":
    print("=" * 60)
    print("연구노트 5개 과제 생성")
    print("=" * 60)

    results = []
    for proj in PROJECTS:
        try:
            out = generate_note(proj)
            ha, hp = verify_note(out, proj)
            results.append((proj["id"], proj["filename"], True, ha, hp))
            print(f"  #{proj['id']} {proj['filename']} 완료")
        except Exception as e:
            results.append((proj["id"], proj["filename"], False, False, False))
            print(f"  #{proj['id']} {proj['filename']} 실패: {e}")

    print()
    print("검증:")
    for pid, fname, ok, ha, hp in results:
        status = "OK" if ok else "FAIL"
        print(f"  #{pid} {fname}: {status}, 작성자={ha}, 과제명={hp}")

    print()
    print("완료")
