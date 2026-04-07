"""연구개발계획서 #1 장비물류 — 최종 보완 스크립트
참고문서(Confluence/Jira) 데이터 기반 구체화
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.shared import Pt
from lxml import etree
import copy

DOC_PATH = Path(r"C:\MES\wta-agents\reports\MAX\경상연구개발\연구개발계획서-1-장비물류.docx")


def set_cell_text(cell, text):
    """셀의 텍스트를 교체 (서식 유지)"""
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
            return
    # runs가 없으면 첫 paragraph에 추가
    if cell.paragraphs:
        cell.paragraphs[0].text = text


def replace_paragraph_text_preserving_style(para, new_text):
    """단락 텍스트 교체 (첫 run 서식 유지)"""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def add_paragraph_after(ref_para, text):
    """단락 뒤에 같은 스타일의 새 단락 추가"""
    from docx.oxml.ns import qn
    new_p = copy.deepcopy(ref_para._element)
    # 모든 run 텍스트 비우기
    for r in new_p.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            t.text = ""
    # 첫 run에 텍스트
    runs = new_p.findall(qn('w:r'))
    if runs:
        ts = runs[0].findall(qn('w:t'))
        if ts:
            ts[0].text = text
            ts[0].set(qn('xml:space'), 'preserve')
    ref_para._element.addnext(new_p)


def main():
    doc = Document(str(DOC_PATH))

    # ─── 1. 개발 목적 보강 (단락 28~30) ───
    purposes = [
        "초경인서트 제조 전 공정(성형→소결→연삭→코팅→검사→포장)에서 공정 간 제품 이송을 완전 자동화하여 장비 무인화 운영을 실현한다. 현재 11개 장비군(프레스핸들러, 소결취출기, 연삭핸들러, PVD/CVD 로딩·언로딩, 포장기, CBN조립기, 검사기, 호닝형상검사기)이 개별적으로 AGV/AMR 물류 대응 기능을 갖추고 있으나, 공정 간 통합 물류 시스템은 부재한 상태이다.",
        "AGV/AMR 기반 자동 물류, 매거진/팔레트 자동 교환, MES 연동 공정 간 물류 스케줄링을 통합 개발한다. 설계팀이 정립한 AMR 물류 5가지 표준 타입(화루이 프레스 AGV형, 한국야금 Cell 프레스형, 대구텍 키엔스 검사기형, 한국야금 포장기/검사기형, 한국야금 호닝핸들러형)을 기반으로 장비별 물류 표준화를 추진한다.",
        "야간/주말 24시간 무인 연속 생산 체계를 구축한다. 2024년 1월 및 2025년 12월 경영회의에서 대표이사가 '무인화-연속운전에 필요한 인식·측정·판단 구성과 공정별 물류 플랫폼 개발'을 핵심 전략으로 지시하였으며, 본 과제는 이에 직접 부합하는 실행 과제이다.",
    ]
    purpose_idx = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "전 구간의 공정 간 제품 이송을 완전 자동화" in t:
            replace_paragraph_text_preserving_style(p, purposes[0])
            purpose_idx = 1
        elif purpose_idx == 1 and "AGV/AMR 기반 자동 물류" in t:
            replace_paragraph_text_preserving_style(p, purposes[1])
            purpose_idx = 2
        elif purpose_idx == 2 and "야간/주말 24시간" in t:
            replace_paragraph_text_preserving_style(p, purposes[2])
            purpose_idx = 0

    # ─── 2. 필요성 보강 (단락 33~34) ───
    necessities = [
        "현재 공정 간 물류는 작업자 개입(팔레트 교체, AGV 수동 호출, 매거진 적재)에 의존하여 야간/주말 무인 연속 가동이 불가능하다. 2022년 AGV 과제는 성형→소결 단일 구간만 대상이었으며, 전체 공정을 관통하는 통합 물류 시스템은 부재한 상태이다. 또한 장비별로 서로 다른 구조의 AMR 물류가 개발되고 있어(현재 5개 type), 설계 요소 통일화와 원가 고려 설계가 시급하다.",
        "경영회의(2024.1, 2025.12)에서 무인화·연속운전 플랫폼 개발이 시장 경쟁력의 핵심으로 강조되었다. 경쟁사 실패 사례를 교훈으로 리스크 관리하며 물류 플랫폼을 추진해야 한다. 구체적으로 MES 연동 스케줄링, AMR 다중 경로 관리, 공정별 버퍼 스테이션 자동 관리, RFID/바코드 기반 제품 식별 등 시스템 수준의 물류 기술이 필요하며, 이를 통해 물류 인력 3~5명 절감, 리드타임 20% 단축, 운반 중 제품 손상/혼입 방지를 달성한다.",
    ]
    nec_idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if "현재 공정 간 물류는 작업자 개입" in t:
            replace_paragraph_text_preserving_style(p, necessities[0])
            nec_idx = 1
        elif nec_idx == 1 and "MES 연동 스케줄링" in t:
            replace_paragraph_text_preserving_style(p, necessities[1])
            nec_idx = 0

    # ─── 3. 추진전략 보강 ───
    strategies = [
        "1단계(3~5월): 전 공정 물류 경로 맵핑 및 버퍼 스테이션 용량 설계. 2022년 AGV 과제의 Fork형 AGV, 엘리베이터 유닛 기술을 기반으로 다구간 확장 설계를 수행한다. 설계팀이 정립한 AMR 물류 5가지 표준 타입(type 1~5)의 설계 요소를 통일화하고, 대구텍 AMR 설계 사례(키엔스 검사기 대응 구조)를 반영한다. 각 장비군별 물류 인터페이스(프레스 SPM 23EA/min, 소결취출 60EA/min, 검사기 20EA/min 등)의 takt time을 고려한 버퍼 용량을 산출한다.",
        "2단계(6~9월): 자동 팔레트/매거진 교환 시스템 개발 및 MES 연동 물류 스케줄링 구현. 2024년 리팔레팅 기술을 적용하여 이종 팔레트 간 자동 전환을 구현한다. 공용 pickup 툴 ATC 내재화(2세대 U19, Jaw/Vacuum/Magnet/Softgrip/Tilt Grip 호환), 회전축 멀티조인트&슬립링 표준설계를 물류 시스템에 통합한다. 경영회의 무인화 지시(2024.1, 2025.12) 목표에 부합하는 로드맵을 준수한다.",
        "3단계(10~12월): 무인 운전 안전 시스템 통합, AGV/AMR 동선 최적화(다수 AGV 운영, 접근 제한 시스템), 현장 실증 및 24시간 연속 가동 테스트를 수행한다. RFID/바코드 기반 제품 식별, 엘리베이터 내 인식 기술을 적용하여 제품 식별→공급→이송→가공→회수 전 과정 자동화를 검증한다. 진동 제거 및 팔레트 내 인서트 안정성을 확보한다.",
    ]
    strat_idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if "1단계(3~5월)" in t:
            replace_paragraph_text_preserving_style(p, strategies[0])
            strat_idx = 1
        elif strat_idx >= 1 and "2단계(6~9월)" in t:
            replace_paragraph_text_preserving_style(p, strategies[1])
            strat_idx = 2
        elif strat_idx >= 2 and "3단계(10~12월)" in t:
            replace_paragraph_text_preserving_style(p, strategies[2])
            strat_idx = 0

    # ─── 4. 제약사항 보강 ───
    for p in doc.paragraphs:
        if "[제약사항]" in p.text:
            replace_paragraph_text_preserving_style(p,
                "[제약사항] (1) AGV/AMR 도입 시 기존 공장 레이아웃 변경 최소화 필요 — 컨베이어 벨트 물류 이송, 팔레트 공급/취출 등 기구 설계 변경 수반. "
                "(2) 충돌 방지 및 비상정지 안전 인증(CE) 확보 필수 — 다수 AGV 동시 운영 시 접근 제한 시스템 구축. "
                "(3) 장비별 takt time 차이(프레스 23EA/min ~ 검사기 20EA/min)에 따른 물류 동기화 난이도. "
                "(4) 연삭유, 분진 등 열악한 환경에서의 비전 검출·센서 내구성 확보. "
                "(5) 소형 제품(⌀1.5~1.8mm) 핸들링 정밀도 및 안정성."
            )
            break

    # ─── 5. 개발 방법 보강 ───
    for p in doc.paragraphs:
        if "개발 방법 :" in p.text:
            replace_paragraph_text_preserving_style(p,
                "개발 방법 : AGV/AMR 기반 자동 물류 시스템 설계(5 type 표준화), 공용 ATC 픽업 툴 내재화, "
                "MES 연동 스케줄링 개발, RFID/바코드 제품 식별 시스템 구축, 현장 실증(24시간 무인 가동 테스트)"
            )
            break

    # ─── 6. WBS 테이블 (Table 1) 보강 ───
    table1 = doc.tables[1]
    wbs_data = [
        ("공정별 물류경로 설계 및 AMR 표준화", "3개월", "2025-03-01", "2025-05-31", "기구설계 2명 (박재성, 지건승)"),
        ("팔레트/매거진 자동교환 및 ATC 통합", "4개월", "2025-06-01", "2025-09-30", "기구설계 2 + 제어 1 (박성수)"),
        ("MES 연동 물류 스케줄링 / HMI", "4개월", "2025-06-01", "2025-09-30", "S/W 2명 (정광선, 최현수)"),
        ("무인운전 통합실증 및 안전시스템", "3개월", "2025-10-01", "2025-12-31", "전원 (5명)"),
    ]
    for ri, (name, period, start, end, resource) in enumerate(wbs_data):
        row = table1.rows[ri + 1]
        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], period)
        set_cell_text(row.cells[2], start)
        set_cell_text(row.cells[3], end)
        set_cell_text(row.cells[4], resource)

    # ─── 7. 참여인력 테이블 (Table 2) 보강 ───
    table2 = doc.tables[2]
    # P.M 행 (row 1)
    set_cell_text(table2.rows[1].cells[1], "조한종")
    set_cell_text(table2.rows[1].cells[2], "10")
    set_cell_text(table2.rows[1].cells[3], "과제 총괄, AMR 물류 표준화")
    # S/W 행 (row 2)
    set_cell_text(table2.rows[2].cells[1], "정광선")
    set_cell_text(table2.rows[2].cells[2], "10")
    set_cell_text(table2.rows[2].cells[3], "핸들러 S/W, MES 물류 연동")
    # 계 행 (row 3)
    set_cell_text(table2.rows[3].cells[1], "( 5 ) 명")
    set_cell_text(table2.rows[3].cells[2], "( 5 ) 명")
    set_cell_text(table2.rows[3].cells[3], "기구2 + 제어1 + S/W2")

    # ─── 8. 관리자 테이블 (Table 0) 보강 ───
    table0 = doc.tables[0]
    set_cell_text(table0.rows[1].cells[0], "조한종")
    set_cell_text(table0.rows[1].cells[1], "생산관리팀")
    set_cell_text(table0.rows[1].cells[2], "과제 총괄")
    set_cell_text(table0.rows[2].cells[0], "박재성")
    set_cell_text(table0.rows[2].cells[1], "기구설계팀")
    set_cell_text(table0.rows[2].cells[2], "기구설계 검토")
    set_cell_text(table0.rows[3].cells[0], "")
    set_cell_text(table0.rows[3].cells[1], "대표이사")
    set_cell_text(table0.rows[3].cells[2], "")

    # ─── 9. 헤더 머리글 표 — 작성일 업데이트 ───
    for section in doc.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "2020.00.00" in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "2020.00.00" in run.text:
                                    run.text = run.text.replace("2020.00.00", "2025.03.01")

    # ─── 10. 참여인원 수 일치 ───
    for p in doc.paragraphs:
        if "참여 인원 : 총 5명" in p.text:
            replace_paragraph_text_preserving_style(p,
                "참여 인원 : 총 5명 (기구설계 2, 제어설계 1, S/W 2) — P.M 조한종, S/W 정광선·최현수, 기구 박재성·지건승, 제어 박성수")
            break

    # ─── 저장 ───
    doc.save(str(DOC_PATH))
    print(f"저장 완료: {DOC_PATH}")

    # ─── 검증 ───
    doc2 = Document(str(DOC_PATH))
    checks = {
        "목적 보강": any("11개 장비군" in p.text for p in doc2.paragraphs),
        "필요성 보강": any("5개 type" in p.text for p in doc2.paragraphs),
        "전략 보강": any("5가지 표준 타입" in p.text for p in doc2.paragraphs),
        "제약사항 보강": any("takt time 차이" in p.text for p in doc2.paragraphs),
        "PM 조한종": doc2.tables[2].rows[1].cells[1].text.strip() == "조한종",
        "SW 정광선": doc2.tables[2].rows[2].cells[1].text.strip() == "정광선",
        "WBS 보강": "AMR 표준화" in doc2.tables[1].rows[1].cells[0].text,
        "헤더 날짜": any("2025.03.01" in cell.text
                       for s in doc2.sections for t in s.header.tables
                       for r in t.rows for cell in r.cells),
    }
    print("\n검증 결과:")
    all_ok = True
    for name, ok in checks.items():
        status = "OK" if ok else "FAIL"
        print(f"  {name}: {status}")
        if not ok:
            all_ok = False

    if all_ok:
        print("\n모든 항목 정상 반영 완료")
    else:
        print("\n일부 항목 확인 필요")


if __name__ == "__main__":
    main()
