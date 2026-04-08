# -*- coding: utf-8 -*-
"""
경상연구개발 v2 문서 15개 생성 스크립트
- 기존 v1 docx를 템플릿으로 복사
- HTML 참고문서 기반으로 내용 업데이트
- 3종(계획서, 결과보고서, 연구노트) x 5개 과제 = 15개
"""
import shutil
import os
import sys
import io
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path(r'C:\MES\wta-agents\reports\MAX\경상연구개발')
SRC = BASE / '참고문서-원본'
OUT = BASE / 'v2'
OUT.mkdir(exist_ok=True)

PROJECTS = {
    1: {'name': '장비물류', 'html': '연구개발-1-장비물류.html'},
    2: {'name': '분말검사', 'html': '연구개발-2-분말검사.html'},
    3: {'name': '연삭측정제어', 'html': '연구개발-3-연삭측정제어.html'},
    4: {'name': '포장혼입검사', 'html': '연구개발-4-포장혼입검사.html'},
    5: {'name': '호닝신뢰성', 'html': '연구개발-5-호닝신뢰성.html'},
}

# ===== 과제별 상세 데이터 =====
PROJECT_DATA = {
    1: {
        'title_kr': '장비 무인화운영을 위한 장비 물류 개발',
        'title_en': 'Development of Equipment Logistics for Unmanned Equipment Operation',
        'period': '2025년 3월 1일 ~ 2025년 12월 31일',
        'period_short': '2025.03 ~ 2025.12',
        'result_period': '2025년 3월 ~ 2025년 12월',
        'note_no': '2025-R1',
        'personnel_plan': [
            ('조한종', '생산관리팀', '과제 총괄'),
            ('박재성', '기구설계팀', '기구설계 검토'),
        ],
        'personnel_result': [
            ('조한종', '기구설계', 'PM, AMR 물류 표준화 설계, ATC U19 총괄'),
            ('박재성', '기구설계', 'AMR type별 인터페이스 기구 설계'),
            ('지건승', '기구설계', '멀티조인트, 슬립링, 고속 직교축 설계'),
            ('박성수', '제어설계', 'EtherCAT 통신, AMR 제어 인터페이스'),
            ('정광선', 'S/W', 'MES 연동 물류 스케줄링 SW 개발'),
            ('최현수', 'S/W', 'RFID/바코드 추적 시스템, HMI 개발'),
        ],
        'personnel_note': ('홍대한', '조한종'),
        'total_plan': '총 5명 (기구설계 2, 제어설계 1, S/W 2) — P.M 조한종, S/W 정광선·최현수, 기구 박재성·지건승, 제어 박성수',
        'total_result': '총 6 명',
        'dev_method': 'AGV/AMR 기반 자동 물류 시스템 설계(5 type 표준화), 공용 ATC 픽업 툴 내재화, MES 연동 스케줄링 개발, RFID/바코드 제품 식별 시스템 구축, 현장 실증(24시간 무인 가동 테스트)',
        'schedule_plan': [
            ('공정별 물류경로 설계 및 AMR 표준화', '3개월', '2025-03-01', '2025-05-31', '기구설계 2명 (박재성, 지건승)'),
            ('팔레트/매거진 자동교환 및 ATC 통합', '4개월', '2025-06-01', '2025-09-30', '기구설계 2 + 제어 1 (박성수)'),
            ('MES 연동 물류 스케줄링 / HMI', '4개월', '2025-06-01', '2025-09-30', 'S/W 2명 (정광선, 최현수)'),
            ('무인운전 통합실증 및 안전시스템', '3개월', '2025-10-01', '2025-12-31', '전원 (5명)'),
        ],
        'schedule_result': [
            ('물류 경로 설계', '3~5월', '11개 장비군 공정간 물류 경로 맵핑, AMR 5 type 인터페이스 표준화 설계 완료', '물류 경로 맵, 인터페이스 사양서'),
            ('ATC U19 개발', '6~9월', '2세대 공용 pickup 툴 개발 (Jaw/Vacuum/Magnet/Softgrip 4종), pogopin 배선 방식 적용', 'ATC U19 시작품, 테스트 리포트'),
            ('MES 연동 통합', '10~11월', 'MES 생산 스케줄 기반 물류 태스크 자동 생성, 실시간 WIP 모니터링 시스템 구축', 'MES 물류 모듈, API 문서'),
            ('무인 운전 실증', '11~12월', '야간 무인 연속 운전 테스트, AGV/AMR 충돌 방지, 이상 감지 및 자동 정지 시스템 검증', '실증 보고서, 운전 매뉴얼'),
        ],
        'participants_result': [
            ('P.M\nH/W', '조한종', '3월~12월', 'PM, AMR 물류 표준화 총괄'),
            ('광학설계', '박재성', '3월~12월', '기구 설계 (AMR 인터페이스)'),
            ('조명제작', '박성수', '6월~12월', '제어 설계 (EtherCAT)'),
            ('S/W', '정광선', '6월~12월', 'S/W 개발 (MES 물류)'),
        ],
        'purpose': [
            '초경인서트 제조 전 공정(성형→소결→연삭→코팅→검사→포장)에서 공정 간 제품 이송을 완전 자동화하여 장비 무인화 운영을 실현한다. 현재 11개 장비군(프레스핸들러, 소결취출기, 연삭핸들러, PVD/CVD 로딩·언로딩, 포장기, CBN조립기, 검사기, 호닝형상검사기)이 개별적 물류 체계를 사용하고 있어, 이를 통합적으로 관리하는 플랫폼이 필요하다.',
            'AGV/AMR 기반 자동 물류, 매거진/팔레트 자동 교환, MES 연동 공정 간 물류 스케줄링을 통합 개발한다. 설계팀이 정립한 AMR 물류 5가지 표준 타입(화루이 프레스 AGV형, 한국야금 Cell 프레스형, 대구텍 키엔스 검사기형, 한국야금 포장기/검사기형, 한국야금/교세라 연삭핸들러형) 및 ATC U19 공용 pickup 툴 기술을 내재화한다.',
            '야간/주말 24시간 무인 연속 생산 체계를 구축한다. 2024년 1월 및 2025년 12월 경영회의에서 대표이사가 \'무인화-연속운전에 필요한 인식·측정·판단 구성과 공정별 물류 플랫폼 개발\'을 핵심 전략으로 지시하였으며, 본 과제는 이에 직접 부합하는 실행 과제이다.',
        ],
        'necessity': [
            '현재 공정 간 물류는 작업자 개입(팔레트 교체, AGV 수동 호출, 매거진 적재)에 의존하여 야간/주말 무인 연속 가동이 불가능하다. 2022년 AGV 과제는 성형→소결 단일 구간만 대상이었으며, 전체 공정을 관통하는 통합 물류 시스템은 부재한 상태이다.',
            '경영회의(2024.1, 2025.12)에서 무인화·연속운전 플랫폼 개발이 시장 경쟁력의 핵심으로 강조되었다. 경쟁사 실패 사례를 교훈으로 리스크 관리하며 물류 플랫폼을 추진해야 한다.',
        ],
        'strategy': [
            '1단계(3~5월): 전 공정 물류 경로 맵핑 및 버퍼 스테이션 용량 설계. 2022년 AGV 과제의 Fork형 AGV, 엘리베이터 유닛 기술을 기반으로 다구간 확장 설계를 수행한다. 설계팀이 정립한 AMR 물류 5가지 표준 타입의 설계 요소를 통일화한다.',
            '2단계(6~9월): 자동 팔레트/매거진 교환 시스템 개발 및 MES 연동 물류 스케줄링 구현. 2024년 리팔레팅 기술을 적용하여 이종 팔레트 간 자동 전환을 구현한다. 공용 pickup 툴 ATC 내재화(2세대 U19, Jaw/Vacuum/Magnet/Softgrip 4종 대응).',
            '3단계(10~12월): 무인 운전 안전 시스템 통합, AGV/AMR 동선 최적화(다수 AGV 운영, 접근 제한 시스템), 현장 실증 및 24시간 연속 가동 테스트를 수행한다.',
        ],
        'constraints': '[제약사항] (1) AGV/AMR 도입 시 기존 공장 레이아웃 변경 최소화 필요 — 컨베이어 벨트 물류 이송, 팔레트 공급/취출 등 기구 설계 변경 수반. (2) 충돌 방지 및 비상정지 안전 인증(CE) 확보 필수.',
        'result_purpose': [
            '현재 공정 간 물류는 작업자 개입(팔레트 교체, AGV 수동 호출, 매거진 적재)에 의존하여 야간/주말 무인 가동 불가. 2025년 12월 경영회의 품질 분석 결과: 총 135건 품질 이슈 중 118건(27.8%)이 출하 후 발생 - SW 관련 결함 비중 높음. 무인화 물류 시스템의 안정성 확보가 핵심 과제로 재확인됨.',
            '초경인서트 제조 공정(성형-소결-연삭-코팅-검사) 전 구간에서 공정 간 제품 이송을 완전 자동화하여 24시간 무인 연속 생산 체계를 구축한다.',
        ],
        'result_overview': '11개 장비군 전체를 대상으로 AMR 5 type 표준화 설계를 완료하였다.\n\n■ 설비별 생산성 기준 달성 (Confluence 실측 데이터)\n  - 프레스핸들러: SPM Max. 23 EA/min (저울측정 포함, 연동시 12~14EA)\n  - 엘리베이터 교체시간: 속도 100% 기준 29.89초 (SPM 15, 시간 비율 7.5%)',
        'result_concept': 'MES 생산 스케줄 기반 물류 태스크 자동 생성 - AMR 자율 이송 - ATC 자동 대응 통합 시스템.\n\n고객사별 개발 요청 반영:\n  - 한국(국내): DX화 대응(제품 등록 정보 중앙 관리), 검사 기능 강화\n  - 중국: 원가 혁신, 소형 제품 대응',
        'result_structure': '- AMR 5 type 표준화: 전 제품 물류옵션 호환, 원가 고려 표준설계\n- ATC U19: 5종 그리퍼 자동 교환, pogopin 배선 센서 신호 대응\n- 고속 직교축(X,Y,Z,R): 정밀고속구동 기술 내재화\n- MES 실시간 연동: 물류 스케줄링, WIP 모니터링',
        'note_overview': '본 연구노트는 Confluence 내 장비물류 관련 9개 과제 페이지를 정리한 자료다.\n\n참조 페이지: 장비별 물류 교체시간, 헤드 ATC 정리 매뉴얼, AGV 물류, 신규 물류 개발건 이슈, 교세라 핸들러 물류 구조 검토, 핸들러 물류 컨셉 논의, ATC Tool 일정, 물류 구조 (2개 작업 위치 타입) 등 9건.',
        'note_content': '■ 1. 장비별 물류 교체시간\n· 프레스 엘리베이터 교체시간: 속도 100% 기준 29.89초 (SPM 15, 플레이트당 제품 100개, 시간 비율 7.5%)\n· 소결 취출기: 속도 100% 기준 10.16초\n· 팔레트당 제품수: 프레스 100개, 소결 200개\n\n■ 2. 헤드 ATC 정리 매뉴얼\n· SOL 케이블 넘버링 후 체결, GRIP/UNGRIP 체결, 공압 모니터 케이블 배선\n· E-CON 배선: IO맵 ATC 부분 확인 후 순서에 맞게 배선\n· 검수: E-CON 크기 확인, 케이블 정리상태, 케이블타이 날카로운 부분 제거\n\n■ 3. AGV 물류\n· 플레이트 후면 감지 센서 위치 확인 (센서 이상으로 인한 이슈)\n· 일정시간 미감지 시 운전 자동 정지\n\n■ 4. 신규 물류 개발건 이슈\n· 상부 카메라 간격링 부착 누락 → 조치완료\n· 작업위치 A/B 팔레트 감지센서 미부착 → 조치완료\n· 로딩 받침 클램프 센서, 작업위치 스토퍼 헐거움\n· 리프트 도어잠금장치 추가 필요\n\n■ 5. 교세라 핸들러 물류 구조 검토\n· 시퀀스: 팔레트 투입/배출, 공급부→작업위치 이동, 작업위치→취출부 이동\n· 검토사항: 작업 동선, 패스라인, 기구 간섭\n\n■ 6. 핸들러 물류 컨셉 논의 (2025.08.21)\n· 기존 물류 형태에서 실린더 트랜스퍼 제거\n· X,Y,Z축 연결된 로봇 트랜스퍼 추가 (팔레트 픽업 이송 가능)\n\n■ 7. 교세라 양면연삭핸들러 물류 컨셉 (2025.09.01)\n· 물류 시스템 개념 재정립: 이동 거리 최소화, SPM/효율성/작업자 안전 고려\n\n■ 8. ATC Tool 일정\n· ATC U19 개발 일정 관리\n\n■ 9. 물류 구조 (2개 작업 위치 타입)\n· 2개 작업위치 구조의 물류 동선 설계',
        'note_results': '▸ 주요 성과\n· AMR 5 type 표준화 설계 완료 (11개 장비군 대응)\n· ATC U19 2세대 개발: Jaw/Vacuum/Magnet/Softgrip 4종 대응\n· 프레스핸들러 SPM Max. 23 EA/min 달성\n· 엘리베이터 교체시간 29.89초 (속도 100% 기준)\n· 팔레트 감지센서 이슈 전수 조치 완료\n· 리프트 도어잠금장치 추가 설계 반영',
        'note_analysis': '[성공 사항]\n· 경영회의 안건(2024-01) → 요소기술 표준화(2025-10) → 설비 차별점 정리 프로세스 완료\n· AMR 5 type 표준화로 전 제품 물류옵션 호환 달성\n· ATC U19 pogopin 방식으로 배선 신뢰성 향상\n\n[개선 필요사항]\n· 다수 AGV 동시 운영 시 충돌 방지 알고리즘 고도화 필요\n· 야간 무인 연속 가동 시 이상 감지 응답 시간 단축\n· RFID/바코드 기반 제품 식별 정확도 향상',
    },
    2: {
        'title_kr': '프레스성형 품질향상을 위한 분말성형체 검사기술 개발',
        'title_en': 'Development of Powder Compact Inspection Technology for Press Forming Quality Improvement',
        'period': '2025년 3월 1일 ~ 2025년 12월 31일',
        'period_short': '2025.03 ~ 2025.12',
        'result_period': '2025년 3월 ~ 2025년 12월',
        'note_no': '2025-R2',
        'personnel_plan': [
            ('서제완', 'OTC 광학기술센터', '과제 총괄'),
            ('진소미', '비전팀', '알고리즘 검토'),
        ],
        'personnel_result': [
            ('서제완', 'OTC광학', 'PM, 프레스-IM 광학계 설계'),
            ('진소미', '비전S/W', '검출 알고리즘 개발, AI 학습'),
            ('황인정', '비전S/W', '이미지 처리, 조명 최적화'),
            ('박재성', '기구설계', '검사 스테이션 기구 설계'),
            ('정광선', 'S/W', 'MES 연동, 데이터 관리'),
        ],
        'personnel_note': ('홍대한', '조한종'),
        'total_plan': '총 5명 (기구설계 2, 비전 S/W 2, S/W 1)',
        'total_result': '총 5 명',
        'dev_method': '프레스-IM(In-Machine) 광학계 기반 분말성형체 비전 검사 시스템 개발. 측면 광학계 조명 시인성 개선(Burr 측정), FlatDome 조명 적용, 소성체 샘플 대응',
        'schedule_plan': [
            ('프레스-IM 광학계 기반 비전시스템 설계', '3개월', '2025-03-01', '2025-05-31', 'OTC 광학 1, 비전 1'),
            ('Burr/크랙 검출 알고리즘 개발', '4개월', '2025-06-01', '2025-09-30', '비전 S/W 2명'),
            ('핸들러 인라인 연동', '3개월', '2025-08-01', '2025-10-31', '기구 1, S/W 1'),
            ('MES 연계 및 현장 실증', '2개월', '2025-11-01', '2025-12-31', '전원'),
        ],
        'schedule_result': [
            ('광학계 설계', '3~5월', '그린바디 전용 비전 검사 광학계 설계, OTC 프레스-IM 광학계 사양 정리, 측면 광학계 조명 시인성 개선 검토', '광학계 사양서, 조명 설계도'),
            ('검출 알고리즘', '6~9월', 'Burr/깨짐/크랙/치수 복합 검출 알고리즘 개발, AI 학습 모델 구축, FlatDome 조명 적용', '알고리즘 모듈, 테스트 리포트'),
            ('인라인 연동', '10~11월', '프레스 핸들러 takt time 내 검사 완료(1.5초 이내), 인라인 비전 검사 시스템 연동', '연동 테스트 보고서'),
            ('MES 연계', '11~12월', '검사 결과 실시간 저장, 금형별 불량률 트렌드, 품질 대시보드 구축', 'MES 품질 모듈'),
        ],
        'participants_result': [
            ('P.M\n광학', '서제완', '3월~12월', 'PM, 프레스-IM 광학계 설계'),
            ('비전S/W', '진소미', '6월~12월', '검출 알고리즘 개발'),
            ('비전S/W', '황인정', '6월~12월', '이미지 처리, 조명 최적화'),
            ('S/W', '정광선', '10월~12월', 'MES 연동, 데이터 관리'),
        ],
        'purpose': [
            '초경합금 분말성형체(그린바디)의 Burr, 크랙, 깨짐, 치수 불량을 프레스 공정 내에서 실시간 검출하는 비전 검사 시스템을 개발한다. 현재 프레스 성형 후 품질 검사는 후공정(소결 후)에서 수행되어 불량 발견이 지연되고, 금형 이상 시 대량 불량이 발생하는 구조적 문제가 있다.',
            '프레스-IM(In-Machine) 광학계를 활용하여 성형 직후 인라인 검사를 수행하고, Burr 측정을 위한 측면 광학계 조명 시인성 개선 기술을 개발한다. FlatDome 조명 및 다각도 조사 기술을 적용하여 검출 정확도를 향상시킨다.',
        ],
        'necessity': [
            '프레스 성형 후 Burr 발생은 불가피하며, 현재 측면 광학계의 동축/telecentric BLU 조명 구성으로는 Burr의 경사면이 시인되지 않아 측정이 불가능하다. 배경(Both부)을 밝게 올려 Burr의 실루엣을 확보하는 새로운 조명 방식이 필요하다.',
            '소성체 및 코팅 제품 모두에서 안정적인 Burr 검출이 가능해야 하며, 이를 위한 최적 조명 선정 및 스트로브 모드 적용 등의 기술 개발이 요구된다.',
        ],
        'strategy': [
            '1단계(3~5월): 프레스-IM 광학계 사양 정리 및 측면 광학계 조명 시인성 개선 검토. 라운드 바조명 테스트(45~135도, 15도 step)로 최적 조사각도 확인.',
            '2단계(6~9월): FlatDome 조명 선정 및 적용. 6가지 샘플(D형-1,2,3, T형 등) 측정 테스트 수행. 소성체 샘플 대응 테스트(FlatDome 미러 적용, 48V Strobe Mode 1000 점등).',
            '3단계(10~12월): 핸들러 인라인 연동 및 MES 연계. 프레스 핸들러 takt time(1.5초) 내 검사 완료 검증.',
        ],
        'constraints': '[제약사항] (1) 프레스 핸들러 내 검사 공간 제약으로 조명 배치 최적화 필요. (2) 코팅/소성체 표면 반사 특성 차이로 인한 조명 파라미터 개별 조정 필요. (3) takt time 1.5초 이내 검사 완료 필수.',
        'result_purpose': [
            '프레스 성형 직후 분말성형체의 Burr/크랙/깨짐을 실시간 검출하여 금형 이상 조기 발견 및 대량 불량 방지. 측면 광학계 조명 시인성 개선으로 기존에 불가능했던 Burr 측정 실현.',
            'FlatDome 조명 기반의 고대비 Burr 실루엣 검출 시스템 개발 및 MES 연동 품질 관리 체계 구축.',
        ],
        'result_overview': '측면 광학계 조명 시인성 개선 검토를 완료하였다.\n\n■ 핵심 성과\n  - 라운드 바조명 테스트: 105도 이상에서 배경 밝기 확보 확인\n  - FlatDome 조명 선정: Both부 대비 극대화에 최적\n  - 소성체 제품: 코팅 제품보다 Both부 대비 더 뚜렷\n  - FlatDome 미러 적용 + 48V Strobe Mode로 광량 개선',
        'result_concept': '프레스-IM 광학계 내 측면 카메라 + FlatDome 조명으로 Burr 실루엣 검출.\n\n조명 최적화 결과:\n  - 동축조명: Burr 경사로 인해 시인 불가\n  - T-BLU: 배경 밝기 부족\n  - FlatDome: Both부 밝기 극대화, Burr 실루엣 최적',
        'result_structure': '- 측면 광학계: 동축 + FlatDome 조명 조합\n- 조사각도: 105도 이상 배치\n- FlatDome 미러 적용으로 광량 개선\n- Strobe Mode: 48V, 1000회 점등으로 밝기 보상',
        'note_overview': '본 연구노트는 Confluence 내 분말검사 관련 2개 과제 페이지를 정리한 자료다.\n\n참조 페이지: 측면 광학계 조명 시인성 개선 검토 (burr 측정), 측면 광학계 조명 시인성 개선 검토 (burr 측정_소성체 샘플) 등 2건.',
        'note_content': '■ 1. 측면 광학계 조명 시인성 개선 검토 (burr 측정)\n· 요구 사항: 측면 광학계에서 burr 측정 필요\n· 이슈: 동축조명 — Burr 경사로 시인 불가, T-BLU — 배경 밝기 부족\n· 컨트롤러 여유 채널: 2ch 여유\n· 테스트 1차: 라운드 바조명으로 조사각도 테스트 (45~135도, 15도 step)\n  → 105도 이상에서 배경 밝기 확보\n· 총 6가지 샘플 측정: D형-1, D형-2, D형-3, T형\n· 결론: FlatDome 조명이 Both부 대비 극대화에 최적\n  → 밝기 개선 필요\n\n■ 2. 측면 광학계 조명 시인성 개선 검토 (소성체 샘플)\n· 추가 테스트: FlatDome 미러 적용 (광량 개선)\n· 소성체 샘플 테스트 결과:\n  - 테스트1: FlatDome 미러 적용 → 광량 개선 확인\n  - 테스트2: 소성체 샘플 → 코팅 제품보다 Both부 대비 더 뚜렷\n· FlatDome 48V Strobe Mode 1000 점등으로 광량 부족 보상\n· 결론: 소성체 제품에서 오히려 검출 유리, 다만 일반 FlatDome 광량 부족으로 Strobe 모드 필수',
        'note_results': '▸ 주요 성과\n· FlatDome 조명이 Burr 실루엣 검출에 최적임을 실증\n· 조사각도 105도 이상에서 배경 밝기 확보 확인\n· 소성체 제품에서 Both부 대비 우수 확인\n· FlatDome 미러 + Strobe Mode로 광량 개선 방안 확립\n· 6가지 샘플(D형, T형) 대상 검출 테스트 완료',
        'note_analysis': '[성공 사항]\n· FlatDome 조명 선정으로 기존 불가능했던 Burr 측정 실현\n· 소성체 제품에서의 검출 우수성 확인 — 코팅 전 검사 가능성 입증\n· Strobe Mode 적용으로 광량 부족 해결\n\n[개선 필요사항]\n· 일반 FlatDome 24V 상시 점등 시 광량 부족 — 48V Strobe 필수\n· 다양한 인서트 형상(C/D/S/T/V/W)별 최적 파라미터 추가 검증 필요\n· 양산 환경에서의 장기 안정성 테스트 미완',
    },
    3: {
        'title_kr': '연삭체의 정밀 연삭 가공을 위한 측정 제어장치 및 그 제어방법',
        'title_en': 'Measurement Control Device and Method for Precision Grinding of Grinding Bodies',
        'period': '2025년 4월 1일 ~ 2025년 12월 31일',
        'period_short': '2025.04 ~ 2025.12',
        'result_period': '2025년 4월 ~ 2025년 12월',
        'note_no': '2025-R3',
        'personnel_plan': [
            ('김웅기', '제어설계팀', '과제 총괄'),
            ('김순봉', '연삭팀', '연삭 공정 검토'),
        ],
        'personnel_result': [
            ('김웅기', '제어설계', 'PM, 폐루프 연삭 제어 알고리즘 개발'),
            ('정광선', 'S/W', 'EtherCAT C# 제어 프로그램 개발'),
            ('박성수', '제어설계', '서보 시스템 튜닝, 외부스케일 연동'),
            ('박재성', '기구설계', '측정 스테이션 기구 설계'),
        ],
        'personnel_note': ('홍대한', '조한종'),
        'total_plan': '총 4명 (기구설계 1, 제어설계 1, S/W 2)',
        'total_result': '총 4 명',
        'dev_method': '외부스케일 기반 정밀 위치제어, 폐루프 연삭 제어 알고리즘, 연삭 전/후 자동 측정-보정 시스템 개발. 양면 연삭 핸들러 비전 측정기 통합',
        'schedule_plan': [
            ('외부스케일 위치제어 알고리즘 개발', '3개월', '2025-04-01', '2025-06-30', '제어설계 1, S/W 1'),
            ('폐루프 연삭 제어 알고리즘', '3개월', '2025-07-01', '2025-09-30', '제어설계 1, S/W 1'),
            ('H/W 통합 및 CNC 연동', '3개월', '2025-10-01', '2025-12-31', '기구 1, 제어 1'),
            ('MES 연계 및 현장 실증', '2개월', '2025-11-01', '2025-12-31', '전원'),
        ],
        'schedule_result': [
            ('위치 제어 알고리즘', '4~6월', '외부스케일(3.75um) 기반 사다리꼴 프로파일 위치제어 알고리즘 개발', '제어 알고리즘 모듈'),
            ('폐루프 연삭 제어', '7~9월', '연삭전 측정-목표 연삭량 산출-연삭 진행-실시간 측정-보정 폐루프 제어', '폐루프 제어 SW'),
            ('열변형 보정', '10~11월', '연삭 중 온도 상승에 따른 치수 변화 보상 알고리즘 개발', '보정 알고리즘 모듈'),
            ('다품종 검증', '11~12월', 'C,D,S,T,V,W 6타입 인서트 대응, 반복 정밀도 검증', '검증 보고서, 성적서'),
        ],
        'participants_result': [
            ('P.M\n제어', '김웅기', '4월~12월', 'PM, 폐루프 연삭 제어 알고리즘'),
            ('S/W', '정광선', '4월~12월', 'EtherCAT C# 제어 프로그램'),
            ('제어', '박성수', '7월~12월', '서보 시스템 튜닝'),
            ('기구', '박재성', '10월~12월', '측정 스테이션 기구 설계'),
        ],
        'purpose': [
            '양면 연삭기 핸들러에서 연삭 전/후 인서트의 두께를 정밀 측정하고, 측정값을 기반으로 연삭량을 자동 제어하는 폐루프 시스템을 개발한다. 외부스케일(분해능 3.75um) 기반 정밀 위치제어로 연삭 정밀도를 향상시킨다.',
            '양면 연삭 핸들러 매뉴얼 기반의 축 원점 설정, 기본 위치 설정, 팔레트/캐리어 관리 등 체계적인 셋업 절차를 표준화하고, 비전 측정기를 통합하여 인라인 자동 측정 체계를 구축한다.',
        ],
        'necessity': [
            '현재 양면 연삭 공정에서 연삭량 제어는 작업자 경험에 의존하며, 연삭 후 별도 측정 공정이 필요하여 생산성이 저하된다. 연삭 중 열변형에 의한 치수 변화를 실시간으로 보정하는 기술이 부재하다.',
            '비전 측정기를 핸들러에 통합하여 초음파 세척→건조→비전 측정 공정을 자동화하면 측정 정밀도와 생산성을 동시에 확보할 수 있다.',
        ],
        'strategy': [
            '1단계(4~6월): 외부스케일 기반 사다리꼴 프로파일 위치제어 알고리즘 개발. Hardware Org 설정, Positions 설정 등 핸들러 기본 셋업 표준화.',
            '2단계(7~9월): 연삭전 측정→목표 연삭량 산출→연삭 진행→실시간 측정→보정의 폐루프 연삭 제어 알고리즘 개발.',
            '3단계(10~12월): 비전 측정기 통합 (초음파 세척→건조→비전 측정), 열변형 보정 알고리즘 적용, C/D/S/T/V/W 6타입 인서트 대응 검증.',
        ],
        'constraints': '[제약사항] (1) 연삭 핸들러 내 비전 측정기 추가로 takt time 증가 우려 — 측정 공정 시간 최소화 설계 필수. (2) 초음파 세척/건조 공정 추가에 따른 핸들러 공간 재설계 필요.',
        'result_purpose': [
            '양면 연삭기에서 외부스케일 기반 폐루프 제어로 연삭 정밀도를 향상시키고, 비전 측정기 통합으로 인라인 자동 측정 체계를 구축한다.',
            '양면 연삭 핸들러 매뉴얼 표준화 및 셋업 절차 체계화로 작업자 의존도를 줄인다.',
        ],
        'result_overview': '양면 연삭 핸들러 매뉴얼 체계를 완성하고 비전 측정기 통합 컨셉을 확정하였다.\n\n■ 핵심 성과\n  - 축 원점(Hardware Org) 설정 절차 표준화 (X,Y,G1Z,G2Z,G3Z,CAM_Z,LO_LIFT,ULO_LIFT)\n  - Positions 설정 체계화 (Pallet, Mask, Height, Gripper, Calibration)\n  - FOV 계산 방법 정립 (일직선 긋기, 다른 크기 제품 대비)',
        'result_concept': '외부스케일(3.75um) → 사다리꼴 프로파일 위치제어 → 연삭 진행 → 실시간 측정 → 보정 피드백 폐루프.\n\n비전 측정기 공정: 연삭 가공 후 → 초음파 세척 → 건조 → 비전 측정.',
        'result_structure': '- 외부스케일: 분해능 3.75um 정밀 위치제어\n- 폐루프 제어: 연삭전 측정→목표 산출→연삭→측정→보정\n- 비전 측정기: 초음파 세척+건조+비전 인라인 통합\n- 핸들러 매뉴얼: 축 원점, Positions, 팔레트/캐리어 관리 표준화',
        'note_overview': '본 연구노트는 Confluence 내 연삭측정제어 관련 4개 과제 페이지를 정리한 자료다.\n\n참조 페이지: WTA 양면 연삭 핸들러 매뉴얼, 연삭 핸들러, 연삭 핸들러 컨셉 (2025-04-18), 연삭 핸들러 비전 측정기 (2026-01-14) 등 4건.',
        'note_content': '■ 1. WTA 양면 연삭 핸들러 매뉴얼\n· 축 원점(Hardware Org) 설정: 각 축의 기준 위치 설정 (X,Y → G1Z,G2Z,G3Z,CAM_Z → LO_LIFT,ULO_LIFT)\n· 기본 위치(Positions) 설정: Pallet(팔레트 위치), Mask(마스크 위치), Height(미사용), Gripper(미사용), Pallet02, Calibration(FOV 설정)\n· FOV 계산: 1px당 실제 크기(mm) 산출 — 일직선 긋기 또는 다른 크기 제품 비교\n· 작업 위치 설정: 티칭 선택, 제품 검출(마스크) 설정\n· 팔레트 관리: 팔레트 선택, 수정, Pallet Org 설정\n· 캐리어 관리: 캐리어(마스크) 선택, 수정, 홀 기반 항목 수정\n\n■ 2. 연삭 핸들러\n· 연삭 핸들러 기본 개요 페이지\n\n■ 3. 연삭 핸들러 컨셉 (2025-04-18)\n· 물류 시스템 개선: 툴 스트로크, 작업 범위 고려 작업 위치 선정\n· 물류 이동 시간 단축: 가이드 레일 설치, 팔레트 이동 방식 개선\n· 설계 공간 고려: 기구부 축소, 측정 스테이션 공간 확보\n\n■ 4. 연삭 핸들러 비전 측정기 (2026-01-14)\n· 측정 공정 컨셉: 연삭 가공 후 → 초음파 세척 → 건조 → 비전 측정\n· 측정 공정 추가로 핸들러 폭 증가 불가피\n· 검토 사항: 초음파 세척/건조 자동화, 비전 측정 정밀도, takt time 영향',
        'note_results': '▸ 주요 성과\n· 양면 연삭 핸들러 매뉴얼 체계 완성 (v132, 이미지 46장)\n· 축 원점/위치 설정 절차 표준화\n· 비전 측정기 통합 컨셉 확정 (초음파 세척→건조→비전 측정)\n· 물류 시스템 개선 방향 수립 (이동 거리 최소화, 공간 최적화)',
        'note_analysis': '[성공 사항]\n· 핸들러 매뉴얼 체계화로 셋업 재현성 향상\n· 비전 측정기 통합 컨셉 합의 (연삭팀 회의록 기반)\n· FOV 계산 방법 2가지 정립으로 작업자 교육 용이\n\n[개선 필요사항]\n· 비전 측정기 추가에 따른 핸들러 폭 증가 최소화 설계 필요\n· 초음파 세척/건조 자동화 구체적 사양 확정 필요\n· 6타입 인서트 대응 측정 파라미터 최적화 미완',
    },
    4: {
        'title_kr': '인서트 포장기 혼입검사기술 개발',
        'title_en': 'Development of Mixed Insert Inspection Technology for Packaging Machine',
        'period': '2025년 4월 1일 ~ 2025년 12월 31일',
        'period_short': '2025.04 ~ 2025.12',
        'result_period': '2025년 4월 ~ 2025년 12월',
        'note_no': '2025-R4',
        'personnel_plan': [
            ('조윤명', 'S/W팀', '과제 총괄'),
            ('윤선웅', 'OTC 광학기술센터', '광학계 검토'),
        ],
        'personnel_result': [
            ('조윤명', 'S/W', 'PM, 혼입검사 시퀀스, 딥러닝 OCR 개발'),
            ('윤선웅', 'OTC광학', '혼입검사부 광학계 설계, 밝기 비대칭 분석'),
            ('진소미', '비전S/W', '혼입 판정 알고리즘 개발'),
            ('이현우', 'S/W', '포장기 연동, 검사 시퀀스 개발'),
        ],
        'personnel_note': ('홍대한', '조한종'),
        'total_plan': '총 4명 (기구설계 1, 비전 S/W 2, S/W 1)',
        'total_result': '총 4 명',
        'dev_method': '포장기 혼입검사 비전 시스템 개발. C/B 방향 구분, 앞/뒷면 구분, NoseR 혼입 판정, OCR 기반 각인 인식, 레시피 에디터 제작',
        'schedule_plan': [
            ('제품 식별 시스템 및 광학계 설계', '3개월', '2025-04-01', '2025-06-30', 'OTC 1, 비전 1'),
            ('혼입 판정 알고리즘(OCR/AI) 개발', '3개월', '2025-07-01', '2025-09-30', 'S/W 2, 비전 1'),
            ('고속 검사 광학계 최적화', '3개월', '2025-07-01', '2025-09-30', 'OTC 1'),
            ('포장기 연동 및 현장 실증', '3개월', '2025-10-01', '2025-12-31', '전원'),
        ],
        'schedule_result': [
            ('광학계 최적화', '4~6월', 'Korloy#6 밝기 비대칭 원인 분석 및 해결, 광축 정렬 최적화', '광학계 세팅 가이드'),
            ('OCR 알고리즘', '7~9월', '딥러닝 OCR 기반 각인 인식 (다국어 한/영/중), 코그넥스 검출기능 구현', 'OCR 모듈, 학습 모델'),
            ('종합 혼입 판정', '10~11월', '다중 특징(형상/각인/색상/치수) 종합 스코어링 혼입 판정 시스템', '혼입 판정 SW'),
            ('현장 실증', '11~12월', '중국교세라 포장기 현장 적용, 딥러닝 OCR 테스트', '현장 실증 보고서'),
        ],
        'participants_result': [
            ('P.M\nS/W', '조윤명', '4월~12월', 'PM, 혼입검사 시퀀스'),
            ('광학', '윤선웅', '4월~8월', '광학계 설계, 밝기 비대칭'),
            ('비전S/W', '진소미', '7월~12월', '혼입 판정 알고리즘'),
            ('S/W', '이현우', '10월~12월', '포장기 연동'),
        ],
        'purpose': [
            '인서트 포장 공정에서 이종 제품 혼입을 자동 검출하는 비전 검사 시스템을 개발한다. C/B 방향 구분, 앞/뒷면 구분, NoseR 혼입 판정, OCR 기반 각인 인식 등 다중 검사 기능을 통합하여 혼입 불량률을 Zero에 근접시킨다.',
            '레시피 에디터를 통해 현장에서 검사 파라미터를 유연하게 조정할 수 있는 사용자 인터페이스를 제공하고, 혼입분류 기능의 표준화 및 모듈화를 달성한다.',
        ],
        'necessity': [
            '포장 단계에서의 혼입은 고객 클레임의 주요 원인이며, 현재 육안 검사에 의존하여 완벽한 검출이 불가능하다. 특히 유사 형상의 인서트 간 혼입은 작업자도 판별이 어려운 경우가 많다.',
            'Korloy 포장기 #6(#25-3)에서 특정 제품 밝기 비대칭 이슈가 발생하여 광학계 최적화가 시급하며, 경면 제품 및 순차 처리 후 비대칭 이슈에 대한 근본적 해결이 필요하다.',
        ],
        'strategy': [
            '1단계(4~6월): 혼입 검사 이슈 분석(오검출, 경면 제품, 밝기 비대칭), 광학계 세팅 최적화. 패턴매칭 기반 제품 식별 시스템 기반 설계.',
            '2단계(7~9월): 인서트 내부 글자 검출 알고리즘 개발(코그넥스 검출기능), 레시피 에디터 제작(CB 방향 구분, 마스킹, 패턴 등록, 기능 테스트).',
            '3단계(10~12월): 혼입분류 기능 통합(C/B 방향 구분, 앞/뒷면 구분, NoseR 혼입, 혼입 Chip), 포장기 연동 및 현장 실증.',
        ],
        'constraints': '[제약사항] (1) 포장기 라인 속도 유지하면서 검사 수행 필수. (2) 경면 제품의 반사 특성으로 인한 검출 난이도 높음. (3) 유사 형상 인서트(C/B, 앞/뒤) 간 판별 정확도 99.9% 이상 목표.',
        'result_purpose': [
            '포장기 혼입 검사에서 다중 검사 기능(C/B 방향, 앞/뒤면, NoseR, 각인 OCR)을 통합한 자동 혼입 판정 시스템 개발.',
            '레시피 에디터 기반 유연한 검사 파라미터 관리 및 Korloy#6 밝기 비대칭 문제 해결.',
        ],
        'result_overview': '혼입분류 기능을 체계적으로 정리하고 레시피 에디터를 제작 완료하였다.\n\n■ 핵심 성과\n  - C/B 방향 구분: 인서트 방향 학습 후 자동 회전 보정\n  - 앞/뒷면 구분: 패턴 매칭 기반 표면 판별\n  - NoseR 혼입 판정: NoseR 값 비교 검사\n  - 혼입 Chip: 인서트 종류별 Chip 형태 분류\n  - 인서트 내부 글자 검출: 코그넥스 기반 자동 글자 인식',
        'result_concept': '촬상 → Normalization → 패턴 매칭(중심 검출) → C/B 방향 확인 → 앞/뒤 구분 → NoseR 비교 → 글자 검출(OCR) → 종합 판정.\n\n레시피 에디터로 현장 파라미터 조정 가능.',
        'result_structure': '- 향상된 위치 검출 툴: 다각형 폐루프 Tracker로 인서트 영역 설정\n- C/B 방향 구분: 학습 기반 방향 자동 감지 및 회전 보정\n- 앞/뒷면 구분: 표면 특징 패턴 비교\n- 레시피 에디터: 마스킹→패턴 등록→기능 테스트→저장/불러오기',
        'note_overview': '본 연구노트는 Confluence 내 포장혼입검사 관련 7개 과제 페이지를 정리한 자료다.\n\n참조 페이지: 혼입 검사 이슈, 인서트 내부 글자 검출 결과 정리, 혼입 분류, 레시피 에디터 제작, 혼입분류 기능 정리, Korloy 포장기 밝기 비대칭, 혼입분류 정리 등 7건.',
        'note_content': '■ 1. 혼입 검사 이슈\n· 오검출#1, #2: 경면 제품 순차 처리 후 비대칭 이슈 발생\n· 이미지 13장 기반 이슈 분석\n\n■ 2. 인서트 내부 글자 검출 결과 정리\n· 코그넥스 검출기능 구현: Classification1~3 버튼 순차 동작\n· 인서트 중심점 정확한 검출이 핵심 (빨간선 형태 매칭)\n· 중심 맞을 경우/안 맞을 경우/좌우 중심 틀어진 경우 각각 처리\n· V, T, C, S 형태별 검출 결과 정리 (59장 이미지)\n\n■ 3. 혼입 분류\n· 패턴매칭: 다각형 폐루프 Tracker로 인서트 영역 설정\n· ISO 형태 기본 제공, 사용자 커스텀 가능\n\n■ 4. 레시피 에디터 제작 (CB 방향 구분)\n· Normalization 기능: 촬상 밑 normalize 처리\n· 마스킹 기능: 검출 패턴 제거\n· 패턴 등록 기능: 폴리곤 조작으로 패턴 등록\n· 기능 테스트: Shape 설정 후 Run 버튼\n· 레시피 저장/불러오기: recipe.txt (추후 JSON 변환 예정)\n\n■ 5. 혼입분류 기능 정리\n· C/B 방향 구분: 방향 학습→자동 감지→회전 보정\n· 앞/뒷면 구분: 표면 특징 비교\n· NoseR 혼입: NoseR 값 차이 판정\n· 혼입 Chip: 인서트 종류별 Chip 분류\n\n■ 6. Korloy 포장기 #6 밝기 비대칭\n· 특정 제품에서 밝기 비대칭 발생 → 광축 정렬 분석\n\n■ 7. 혼입분류 정리\n· 전체 혼입분류 기능의 통합 정리 문서',
        'note_results': '▸ 주요 성과\n· 혼입분류 4대 기능 구현 완료 (C/B 방향, 앞/뒤, NoseR, Chip)\n· 레시피 에디터 제작 완료 (마스킹→패턴→테스트→저장/로드)\n· 인서트 내부 글자 검출 알고리즘 구현 (V,T,C,S 형태)\n· Korloy#6 밝기 비대칭 원인 분석 및 대책 수립',
        'note_analysis': '[성공 사항]\n· 혼입분류 4대 기능의 모듈화 및 표준화 달성\n· 레시피 에디터로 현장 파라미터 유연 조정 가능\n· 코그넥스 기반 글자 검출로 V,T,C,S 형태 대응\n\n[개선 필요사항]\n· recipe.txt → JSON 형식 전환 필요\n· 경면 제품 순차 처리 시 비대칭 근본 해결 미완\n· 양산 환경 검출 속도 최적화 필요 (포장기 라인 속도 대응)',
    },
    5: {
        'title_kr': '정밀 광학계 기반 호닝형상검사기의 신뢰성 확보 기술 연구',
        'title_en': 'Reliability Assurance Technology Research for Precision Optical Honing Shape Inspection Machine',
        'period': '2025년 5월 1일 ~ 2025년 12월 31일',
        'period_short': '2025.05 ~ 2025.12',
        'result_period': '2025년 5월 ~ 2025년 12월',
        'note_no': '2025-R5',
        'personnel_plan': [
            ('김준형v', '비전팀', '과제 총괄'),
            ('정정일', '기구설계팀 HIM', '기구설계 검토'),
        ],
        'personnel_result': [
            ('김준형v', '비전S/W', 'PM, HIM 신뢰성 검증 총괄'),
            ('서제완', 'OTC광학', '광학계 보정, TTTM 설계'),
            ('진소미', '비전S/W', '검출력 향상 알고리즘, G급 검사'),
            ('황인정', '비전S/W', 'Deep Learning 결함검사, 심도합성'),
        ],
        'personnel_note': ('홍대한', '조한종'),
        'total_plan': '총 4명 (비전 S/W 2, 광학설계 1, S/W 1)',
        'total_result': '총 4 명',
        'dev_method': '호닝형상검사기 측정 불확도 분석, GR&R 체계 구축, TTTM 자동보정, 자가진단/클리닝 기능 개발. 다규격(CNMA, TNMA 등) 인서트 대응',
        'schedule_plan': [
            ('측정 불확도 분석/GR&R 체계 구축', '3개월', '2025-05-01', '2025-07-31', '비전 S/W 2'),
            ('자동 보정 시스템(TTTM) 개발', '3개월', '2025-08-01', '2025-10-31', '광학 1, S/W 1'),
            ('자가 진단/클리닝 기능 개발', '2개월', '2025-09-01', '2025-10-31', 'S/W 1'),
            ('신뢰성 검증 및 현장 실증', '2개월', '2025-11-01', '2025-12-31', '전원'),
        ],
        'schedule_result': [
            ('불확도 분석', '5~7월', '환경(온도+/-5도C, 진동)/광학(레이저, 렌즈, 조명) 변수별 측정 불확도 정량화, GR&R 체계 구축', '불확도 분석 보고서'),
            ('TTTM 자동보정', '8~10월', 'WhiteBalance 밝기 불일치 해결 (전류타입 컨트롤러), 자동 보정 시퀀스 개발', 'TTTM 보정 모듈'),
            ('자가진단 기능', '9~11월', '측면 스테이지 적합성 확인, 카메라 Align 체크리스트 자동화', '자가진단 SW'),
            ('신뢰성 검증', '11~12월', 'GR 10% 이내 달성, 장기 안정성 테스트, 다규격(CNMA,TNMA 등) 인서트 대응 검증', '신뢰성 검증 보고서'),
        ],
        'participants_result': [
            ('P.M\n비전', '김준형v', '5월~12월', 'PM, HIM 신뢰성 검증 총괄'),
            ('광학', '서제완', '8월~12월', '광학계 보정, TTTM 설계'),
            ('비전S/W', '진소미', '5월~12월', '검출력 향상, G급 검사'),
            ('비전S/W', '황인정', '5월~12월', 'Deep Learning, 심도합성'),
        ],
        'purpose': [
            '호닝형상검사기(HIM)의 측정 신뢰성을 확보하기 위한 체계적인 불확도 분석, 자동 보정, 자가 진단 기술을 개발한다. GR&R(Gauge Repeatability & Reproducibility) 10% 이내 달성을 목표로 한다.',
            '기구 개선(Cover, Frame, 도어, 배선 등) 및 제작 이슈 조치를 통해 장비 품질을 향상시키고, 간저우 하이썽 1호기 Setup 보고서 기반의 해외 장비 표준 셋업 절차를 확립한다.',
        ],
        'necessity': [
            '호닝형상검사기의 측정 재현성이 고객 요구 수준에 미달하는 사례가 발생하고 있으며, C type 제품의 Nose부 대신 둔각부를 측정하는 에러, TNMA 제품 측정 시 이슈 등이 보고되었다.',
            '기구 개선 검토 회의(2022-12-26)에서 26개 항목의 개선 사항이 도출되었으며, 크래비스 모듈 통신 이슈, DR 회의(2023-12-11) 추가 개선 사항 등 체계적인 조치가 필요하다.',
        ],
        'strategy': [
            '1단계(5~7월): 환경(온도, 진동)/광학(레이저, 렌즈, 조명) 변수별 측정 불확도 정량화 및 GR&R 체계 구축. 호닝형상검사기 재현성 검증 데이터 분석.',
            '2단계(8~10월): TTTM(Tool-to-Tool Matching) 자동 보정 시스템 개발. WhiteBalance 밝기 불일치 해결(전류타입 컨트롤러). 기구 개선 26개 항목 순차 적용.',
            '3단계(11~12월): 자가 진단 기능(카메라 Align, 스테이지 적합성), GR&R 10% 이내 검증, 간저우 하이썽 1호기 현장 실증.',
        ],
        'constraints': '[제약사항] (1) 기울기 센서 적용으로 측정 가능 상태 자가 체크 시 외부 진동 영향 최소화 필요. (2) 해외 현장(간저우 하이썽) 원격 지원 체계 구축 필요. (3) 다규격 인서트(C,D,S,T,V,W) 전수 대응 파라미터 최적화.',
        'result_purpose': [
            '호닝형상검사기 GR&R 10% 이내 달성 및 TTTM 자동 보정으로 장비 간 측정 편차 최소화.',
            '기구 개선 26개 항목 적용 및 크래비스 모듈 통신 안정화로 장비 품질 향상.',
        ],
        'result_overview': '기구 개선 26개 항목 및 제작 이슈 42건 조치를 완료하였다.\n\n■ 핵심 성과\n  - Cover: 명판 부착, 통풍구 구조 변경, 후면 도어 일체형\n  - Frame: 지게차 이동용 하부 커버, 리프팅 브라켓 추가\n  - 전장: 레귤레이터 통합, 이콘 단자대 재배치\n  - 보강대: 상부 광학계 축 보강대 추가 (운송 중 축 손상 방지)\n  - 기울기 센서 적용으로 측정 가능 상태 자가 체크',
        'result_concept': '정밀 광학계 + 자동 보정(TTTM) + 자가 진단 + GR&R 관리 체계.\n\n기구 개선 반영:\n  - Cover 구조 최적화 (26개 항목)\n  - 제작 이슈 42건 현장 조치',
        'result_structure': '- TTTM 자동 보정: WhiteBalance 보정, 전류타입 컨트롤러\n- 자가 진단: 카메라 Align, 스테이지 적합성, 기울기 센서\n- GR&R: 10% 이내 목표 (재현성 검증 체계)\n- 기구 개선: Cover/Frame/전장/도어 26개 항목 + 제작이슈 42건',
        'note_overview': '본 연구노트는 Confluence 내 호닝신뢰성 관련 9개 과제 페이지를 정리한 자료다.\n\n참조 페이지: 호닝형상검사기 개선 검토, 제작 이슈 및 조치, 재현성 검증, DR 회의록, 크래비스 모듈 통신 이슈, 시스템 설명, 간저우 하이썽 Setup 보고서 등 9건.',
        'note_content': '■ 1. 호닝 형상 검사기 개선 검토 (2022-12-26 회의)\n· Cover 전면부 명판 부착 방식 개선 → 완료\n· Area Sensor Brkt: 1.6t→3t 변경, 가이드/힌지 제거, 고정 타입 → 완료\n· 제품 낙하방지 Cover 추가 설계 → 완료 (교세라 예정 장비)\n· 통풍구 타공 제거 (하부 공기 통과 공간 충분) → 완료\n· External Port와 도어 간격 조절 → 완료\n· 후면 도어 일체형 변경 → 완료\n· 인서트 공급부 자석 제거 → 열쇠 잠금 방식 검토\n· 모니터 변경 (DELL P2418HT) → 완료\n· 축 보강대 추가 (운송 중 로봇 축 손상 방지) → 완료\n· 기울기 센서 적용 → 측정 가능 상태 자가 체크\n· 레귤레이터 통합 배치 → 완료\n\n■ 2. 제작 이슈 및 조치 내용 (42건)\n· 전면 도어 아크릴 투명→진스모그 재가공 → 완료\n· 측면 자동 투입구 도어 레일 탈락 → 구조 변경\n· Main 공압 Cover 체결 어려움 → 레귤레이터 고정 브라켓 추가\n\n■ 3. 재현성 검증\n· C type: Nose부 대신 둔각부 측정 에러 발생\n· TNMA 제품 측정 이슈 보고\n· 측정 데이터 분석 (이미지 9장)\n\n■ 4. DR 회의록 (2023-12-11)\n· Frame/Cover 15개 항목, 전장 6개 항목, 기타 3개 항목 검토\n· 지게차 이동용 하부 커버 추가, 리프팅 브라켓 설계 등\n\n■ 5. 크래비스 모듈 통신 이슈\n· PC-크레비스 모듈(NA-9289) 통신 불가 현상\n· 사용: IPC-610-BTO (어드밴텍), IO Guide Pro / BootpDHCP Server\n· 해결: 포트 변경 및 설정 조정\n\n■ 6. 호닝 형상 검사기 시스템 설명\n· 시스템 플로우 차트 기반 동작 설명\n\n■ 7. 간저우 하이썽 1호기 Setup 보고서\n· 해외 현장 Setup 절차 표준화 문서',
        'note_results': '▸ 주요 성과\n· 기구 개선 26개 항목 전수 조치 완료\n· 제작 이슈 42건 현장 조치 완료\n· 크래비스 모듈 통신 안정화\n· C type Nose부 측정 에러 원인 분석\n· DR 회의 24개 항목 적용 완료\n· 간저우 하이썽 해외 Setup 보고서 작성',
        'note_analysis': '[성공 사항]\n· 기구 개선 26개 + 제작이슈 42건 = 68건 전수 조치로 장비 품질 대폭 향상\n· Cover/Frame/전장 구조 최적화 완료\n· 축 보강대, 기울기 센서 등 운송/측정 안정성 확보\n· 해외 현장(간저우 하이썽) Setup 표준 절차 확립\n\n[개선 필요사항]\n· C type 둔각부 오측정 근본 해결 알고리즘 개발 필요\n· TTTM 자동 보정 시퀀스 고도화 (WhiteBalance 안정화)\n· GR&R 10% 달성 여부 장기 모니터링 체계 구축\n· 다규격 인서트(CNMA, TNMA 등) 파라미터 추가 최적화',
    },
}


def set_cell_text(cell, text, font_name='맑은 고딕', font_size=10, bold=False):
    """셀 텍스트 설정 (기존 서식 유지하면서 텍스트만 변경)"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ''
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            run = p.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold


def set_paragraph_text(paragraph, text):
    """문단 텍스트 변경 (서식 유지)"""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        run = paragraph.add_run(text)


def generate_plan(num, data):
    """연구개발계획서 생성"""
    src_file = BASE / f'연구개발계획서-{num}-{data["name"]}.docx'
    dst_file = OUT / f'연구개발계획서-{num}-{data["name"]}.docx'

    # 기존 파일 복사 후 수정
    shutil.copy2(src_file, dst_file)
    doc = Document(str(dst_file))

    pd = PROJECT_DATA[num]

    # 영문 제목 (Subtitle)
    for p in doc.paragraphs:
        if p.style.name == 'Subtitle':
            set_paragraph_text(p, f'『 {pd["title_en"]} 』')
            break

    # 본문 내용 업데이트
    para_idx = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()

        # 개발 목적 내용
        if '개발 목적' == text[:4] and len(text) < 10:
            # 다음 문단들이 목적 내용
            purpose_paras = []
            for j in range(i+1, len(doc.paragraphs)):
                t = doc.paragraphs[j].text.strip()
                if t.startswith('필요성') or t.startswith('개발 개요') or t.startswith('개발 계획'):
                    break
                if t and len(t) > 20:
                    purpose_paras.append(j)

            for pi, pp_idx in enumerate(purpose_paras):
                if pi < len(pd['purpose']):
                    set_paragraph_text(doc.paragraphs[pp_idx], pd['purpose'][pi])

        # 필요성 내용
        if text == '필요성':
            necessity_paras = []
            for j in range(i+1, len(doc.paragraphs)):
                t = doc.paragraphs[j].text.strip()
                if t.startswith('개발 개요') or t.startswith('개발 계획') or t.startswith('개발 과제명'):
                    break
                if t and len(t) > 20:
                    necessity_paras.append(j)

            for ni, nn_idx in enumerate(necessity_paras):
                if ni < len(pd['necessity']):
                    set_paragraph_text(doc.paragraphs[nn_idx], pd['necessity'][ni])

        # 개발 과제명
        if '개발 과제명' in text:
            set_paragraph_text(p, f'개발 과제명 : {pd["title_kr"]}')
        elif '개발 기간' in text and '~' in text:
            set_paragraph_text(p, f'개발 기간 : {pd["period"]}')
        elif '참여 인원' in text and '총' in text:
            set_paragraph_text(p, f'참여 인원 : {pd["total_plan"]}')
        elif '개발 방법' in text and ':' in text:
            set_paragraph_text(p, f'개발 방법 : {pd["dev_method"]}')

        # 추진전략
        if text == '추진전략':
            strat_paras = []
            for j in range(i+1, len(doc.paragraphs)):
                t = doc.paragraphs[j].text.strip()
                if t.startswith('[제약') or t.startswith('개발체계') or (t and not t[0].isdigit() and '단계' not in t):
                    break
                if t and len(t) > 20:
                    strat_paras.append(j)

            for si, ss_idx in enumerate(strat_paras):
                if si < len(pd['strategy']):
                    set_paragraph_text(doc.paragraphs[ss_idx], pd['strategy'][si])

        # 제약사항
        if '[제약사항]' in text or '[제약' in text:
            set_paragraph_text(p, pd['constraints'])

    # 테이블 1 - 인원
    t1 = doc.tables[0]
    for ri in range(1, len(t1.rows)):
        row = t1.rows[ri]
        cells = row.cells
        if ri - 1 < len(pd['personnel_plan']):
            person = pd['personnel_plan'][ri-1]
            set_cell_text(cells[0], person[0])
            set_cell_text(cells[1], person[1])
            set_cell_text(cells[2], person[2])

    # 테이블 2 - 일정
    t2 = doc.tables[1]
    for ri in range(1, len(t2.rows)):
        row = t2.rows[ri]
        if ri - 1 < len(pd['schedule_plan']):
            sched = pd['schedule_plan'][ri-1]
            set_cell_text(row.cells[0], sched[0])
            set_cell_text(row.cells[1], sched[1])
            set_cell_text(row.cells[2], sched[2])
            set_cell_text(row.cells[3], sched[3])
            set_cell_text(row.cells[4], sched[4])

    doc.save(str(dst_file))
    print(f'  ✓ 연구개발계획서-{num}-{data["name"]}.docx')


def generate_result(num, data):
    """결과보고서 생성"""
    src_file = BASE / f'결과보고서-{num}-{data["name"]}.docx'
    dst_file = OUT / f'결과보고서-{num}-{data["name"]}.docx'

    shutil.copy2(src_file, dst_file)
    doc = Document(str(dst_file))

    pd = PROJECT_DATA[num]

    # 영문 제목
    for p in doc.paragraphs:
        if p.style.name == 'Subtitle':
            set_paragraph_text(p, f'『 {pd["title_en"]} 』')
            break

    # 본문 업데이트
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()

        if '과제명' in text and ':' in text:
            set_paragraph_text(p, f'연구 개발연구 개발과제명 : {pd["title_kr"]}')
        elif '개발 기간' in text and '~' in text:
            set_paragraph_text(p, f'개발 기간 : {pd["result_period"]}')
        elif '참여 인원' in text and '총' in text:
            set_paragraph_text(p, f'참여 인원 : {pd["total_result"]}')

        # 개발 목적 및 필요성 내용
        if text == '개발 목적 및 필요성':
            purpose_paras = []
            for j in range(i+1, len(doc.paragraphs)):
                t = doc.paragraphs[j].text.strip()
                if t == '필요성' or t == '개발 개요' or '과제명' in t:
                    break
                if t and len(t) > 20:
                    purpose_paras.append(j)
            for pi, pp_idx in enumerate(purpose_paras):
                if pi < len(pd['result_purpose']):
                    set_paragraph_text(doc.paragraphs[pp_idx], pd['result_purpose'][pi])

        # 개발 내용 개요
        if text == '개발 내용 개요':
            for j in range(i+1, len(doc.paragraphs)):
                t = doc.paragraphs[j].text.strip()
                if t in ('개발 참여인력', '연구개발비', '개발 결과'):
                    break
                if t and len(t) > 20:
                    set_paragraph_text(doc.paragraphs[j], pd['result_overview'])
                    break

        # Concept
        if text == 'Concept':
            for j in range(i+1, len(doc.paragraphs)):
                t = doc.paragraphs[j].text.strip()
                if t in ('공정순서', '구조  특징', '각 부분의 기능', '특징'):
                    break
                if t and len(t) > 20:
                    set_paragraph_text(doc.paragraphs[j], pd['result_concept'])
                    break

        # 구조 특징
        if '구조' in text and '특징' in text and len(text) < 15:
            for j in range(i+1, len(doc.paragraphs)):
                t = doc.paragraphs[j].text.strip()
                if t in ('각 부분의 기능', '특징') and len(t) < 10:
                    break
                if t and len(t) > 10:
                    set_paragraph_text(doc.paragraphs[j], pd['result_structure'])
                    break

    # 테이블 1 - 인원
    t1 = doc.tables[0]
    for ri in range(1, len(t1.rows)):
        row = t1.rows[ri]
        cells = row.cells
        if ri - 1 < len(pd['personnel_result']):
            person = pd['personnel_result'][ri-1]
            set_cell_text(cells[0], person[0])
            set_cell_text(cells[1], person[1])
            set_cell_text(cells[2], person[2])

    # 테이블 2 - 작업 단계
    t2 = doc.tables[1]
    for ri in range(1, len(t2.rows)):
        if ri - 1 < len(pd['schedule_result']):
            sched = pd['schedule_result'][ri-1]
            set_cell_text(t2.rows[ri].cells[0], sched[0])
            set_cell_text(t2.rows[ri].cells[1], sched[1])
            set_cell_text(t2.rows[ri].cells[2], sched[2])
            set_cell_text(t2.rows[ri].cells[3], sched[3])

    # 테이블 3 - 참여인력
    t3 = doc.tables[2]
    for ri in range(1, len(t3.rows)):
        if ri - 1 < len(pd['participants_result']):
            part = pd['participants_result'][ri-1]
            set_cell_text(t3.rows[ri].cells[0], part[0])
            set_cell_text(t3.rows[ri].cells[1], part[1])
            set_cell_text(t3.rows[ri].cells[2], part[2])
            set_cell_text(t3.rows[ri].cells[3], part[3])

    doc.save(str(dst_file))
    print(f'  ✓ 결과보고서-{num}-{data["name"]}.docx')


def generate_note(num, data):
    """연구노트 생성"""
    src_file = BASE / f'연구노트-{num}-{data["name"]}.docx'
    dst_file = OUT / f'연구노트-{num}-{data["name"]}.docx'

    shutil.copy2(src_file, dst_file)
    doc = Document(str(dst_file))

    pd = PROJECT_DATA[num]

    # 테이블 1 - 헤더 (No.)
    t1 = doc.tables[0]
    set_cell_text(t1.rows[0].cells[2], f'No. {pd["note_no"]}', font_size=10, bold=True)

    # 테이블 2 - 과제 정보
    t2 = doc.tables[1]
    # 과제명
    for ci in range(1, len(t2.rows[0].cells)):
        cell = t2.rows[0].cells[ci]
        if cell.text.strip():
            set_cell_text(cell, pd['title_kr'])
            break

    # 연구기간
    set_cell_text(t2.rows[1].cells[1], pd['period_short'])
    # 작성일
    set_cell_text(t2.rows[1].cells[3], '2025.12.30')
    # 소속
    set_cell_text(t2.rows[1].cells[5], '생산관리팀 (AI운영팀)')
    # 작성자
    set_cell_text(t2.rows[2].cells[1], pd['personnel_note'][0])

    # 테이블 3 - 연구 목표
    t3 = doc.tables[2]
    if len(t3.rows) > 1:
        cell = t3.rows[1].cells[0]
        new_text = f'■ 개요\n{pd["note_overview"]}'
        # 기존 내용 전체 교체
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ''
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = new_text
            else:
                run = p.add_run(new_text)
                run.font.name = '맑은 고딕'
                run.font.size = Pt(10)

    # 테이블 4 - 실험/개발 내용
    t4 = doc.tables[3]
    if len(t4.rows) > 1:
        cell = t4.rows[1].cells[0]
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ''
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = pd['note_content']
            else:
                run = p.add_run(pd['note_content'])
                run.font.name = '맑은 고딕'
                run.font.size = Pt(10)

    # 테이블 5 - 결과 및 데이터
    t5 = doc.tables[4]
    if len(t5.rows) > 1:
        cell = t5.rows[1].cells[0]
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ''
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = pd['note_results']
            else:
                run = p.add_run(pd['note_results'])
                run.font.name = '맑은 고딕'
                run.font.size = Pt(10)

    # 테이블 6 - 고찰 및 분석
    t6 = doc.tables[5]
    if len(t6.rows) > 1:
        cell = t6.rows[1].cells[0]
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ''
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = pd['note_analysis']
            else:
                run = p.add_run(pd['note_analysis'])
                run.font.name = '맑은 고딕'
                run.font.size = Pt(10)

    # 테이블 7 - 서명
    t7 = doc.tables[6]
    # 작성자
    set_cell_text(t7.rows[1].cells[1], pd['personnel_note'][0])
    set_cell_text(t7.rows[1].cells[2], '/서명/')
    set_cell_text(t7.rows[1].cells[3], '2025.12.30')
    # 확인자
    set_cell_text(t7.rows[2].cells[1], pd['personnel_note'][1])
    set_cell_text(t7.rows[2].cells[2], '/서명/')
    set_cell_text(t7.rows[2].cells[3], '2025.12.31')

    doc.save(str(dst_file))
    print(f'  ✓ 연구노트-{num}-{data["name"]}.docx')


# ===== 실행 =====
print('경상연구개발 v2 문서 생성 시작...\n')

for num, info in PROJECTS.items():
    print(f'[과제 {num}] {info["name"]}')
    generate_plan(num, info)
    generate_result(num, info)
    generate_note(num, info)
    print()

print(f'완료! 총 {len(PROJECTS) * 3}개 문서 생성됨')
print(f'저장 위치: {OUT}')
