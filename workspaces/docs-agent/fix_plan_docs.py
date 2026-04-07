"""연구개발 계획서 5건 수정 스크립트
1. 표지 텍스트박스: 개발과제명 → 실제 과제명
2. 표지 영문 표기명 업데이트
3. 헤더 머리글 표: 과제명, 작성일, 작성자 업데이트
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from pathlib import Path
from docx import Document
from lxml import etree

OUT_DIR = Path(r"C:\MES\wta-agents\reports\MAX")

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}

PROJECTS = [
    {
        "id": 1,
        "file": "연구개발계획서-1-장비물류.docx",
        "title_kr": "장비 무인화운영을 위한 장비 물류 개발",
        "title_en": "Development of Equipment Logistics for Unmanned Equipment Operation",
        "period": "2025.03 ~ 2025.12",
        "members": "5명",
    },
    {
        "id": 2,
        "file": "연구개발계획서-2-분말검사.docx",
        "title_kr": "프레스성형 품질향상을 위한 분말성형체 검사기술 개발",
        "title_en": "Development of Powder Compact Inspection Technology for Press Forming Quality Improvement",
        "period": "2025.03 ~ 2025.12",
        "members": "5명",
    },
    {
        "id": 3,
        "file": "연구개발계획서-3-연삭측정제어.docx",
        "title_kr": "연삭체의 정밀 연삭 가공을 위한 측정 제어장치 및 그 제어방법",
        "title_en": "Measurement Control Device and Method for Precision Grinding of Ground Bodies",
        "period": "2025.04 ~ 2025.12",
        "members": "4명",
    },
    {
        "id": 4,
        "file": "연구개발계획서-4-포장혼입검사.docx",
        "title_kr": "인서트 포장기 혼입검사기술 개발",
        "title_en": "Development of Cross-Contamination Inspection Technology for Insert Packaging Machine",
        "period": "2025.04 ~ 2025.12",
        "members": "4명",
    },
    {
        "id": 5,
        "file": "연구개발계획서-5-호닝신뢰성.docx",
        "title_kr": "정밀 광학계 기반 호닝형상검사기의 신뢰성 확보 기술 연구",
        "title_en": "Research on Reliability Assurance Technology for Precision Optics-Based Honing Shape Inspection Machine",
        "period": "2025.05 ~ 2025.12",
        "members": "4명",
    },
]


def replace_textbox_text(body, old_text, new_text):
    """텍스트박스 내의 특정 텍스트를 교체"""
    count = 0
    for elem in body.iter():
        tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ''
        if tag == 'txbxContent':
            for t_elem in elem.iter():
                if isinstance(t_elem.tag, str) and etree.QName(t_elem.tag).localname == 't':
                    if t_elem.text and old_text in t_elem.text:
                        t_elem.text = t_elem.text.replace(old_text, new_text)
                        count += 1
    return count


def update_header_table(doc, proj):
    """헤더의 머리글 표 업데이트"""
    for section in doc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                cells = row.cells
                for ci, cell in enumerate(cells):
                    text = cell.text.strip()
                    # 프로젝트명 셀의 다음 셀이 "개발과제명"이면 교체
                    if text == "개발과제명":
                        # 셀 내부 paragraph의 run 텍스트 교체
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "개발과제명" in run.text:
                                    run.text = run.text.replace("개발과제명", proj["title_kr"])
                    # 작성일 교체
                    if text == "2020.00.00":
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "2020.00.00" in run.text:
                                    run.text = run.text.replace("2020.00.00", "2025.03.01")
                    # 작성자 교체
                    if text == "사원 홍길동":
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "사원 홍길동" in run.text:
                                    run.text = run.text.replace("사원 홍길동", "생산관리팀")


def update_cover_english(doc, eng_title):
    """표지 영문 표기명 업데이트"""
    for p in doc.paragraphs:
        if "『" in p.text and "』" in p.text:
            # 현재 영문 제목을 새 것으로 교체
            runs = p.runs
            if len(runs) >= 3:
                # 가운데 run들에 영문 제목 설정
                # run[0]='『', run[-1]='』', 나머지가 내용
                content_runs = runs[1:-1]
                if content_runs:
                    content_runs[0].text = f" {eng_title} "
                    for r in content_runs[1:]:
                        r.text = ""
            break


def fix_doc(proj):
    """하나의 문서 수정"""
    filepath = OUT_DIR / proj["file"]
    doc = Document(str(filepath))

    # 1. 표지 텍스트박스: "개발과제명" → 실제 과제명
    body = doc.element.body
    count = replace_textbox_text(body, "개발과제명", proj["title_kr"])
    print(f"  텍스트박스 교체: {count}건")

    # 2. 표지 영문 표기명
    update_cover_english(doc, proj["title_en"])
    print(f"  영문 표기명 업데이트")

    # 3. 헤더 머리글 표
    update_header_table(doc, proj)
    print(f"  헤더 머리글 표 업데이트")

    doc.save(str(filepath))
    print(f"  저장 완료: {filepath.name}")


if __name__ == "__main__":
    print("=" * 60)
    print("연구개발 계획서 5건 수정")
    print("=" * 60)

    for proj in PROJECTS:
        print(f"\n#{proj['id']} {proj['file']}")
        try:
            fix_doc(proj)
        except Exception as e:
            print(f"  [오류] {e}")
            import traceback
            traceback.print_exc()

    print("\n" + "=" * 60)
    print("수정 완료. 검증 중...")
    print("=" * 60)

    # 검증
    for proj in PROJECTS:
        filepath = OUT_DIR / proj["file"]
        doc = Document(str(filepath))

        # 텍스트박스 확인
        found_title = False
        for elem in doc.element.body.iter():
            tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ''
            if tag == 'txbxContent':
                for t_elem in elem.iter():
                    if isinstance(t_elem.tag, str) and etree.QName(t_elem.tag).localname == 't':
                        if t_elem.text and proj["title_kr"] in t_elem.text:
                            found_title = True

        # 헤더 확인
        header_ok = False
        for section in doc.sections:
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if proj["title_kr"] in cell.text:
                            header_ok = True

        # 영문 확인
        eng_ok = False
        for p in doc.paragraphs:
            if proj["title_en"] in p.text:
                eng_ok = True

        print(f"  #{proj['id']}: 표지={found_title}, 헤더={header_ok}, 영문={eng_ok}")

    print("\n완료")
